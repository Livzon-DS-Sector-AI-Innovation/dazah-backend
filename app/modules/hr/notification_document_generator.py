"""培训通知 Word 文档生成器."""

from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from pydantic import BaseModel


class TrainingNotificationInput(BaseModel):
    department: str
    training_date: date
    subject: str
    training_time_start: str | None = None
    training_time_end: str | None = None
    location: str | None = None
    trainer: str | None = None
    content: str | None = None
    trainee_names: list[str] = []
    issuer_department: str | None = None
    issue_date: date | None = None


def _find_template(factory: str = "old") -> Path:
    """Locate the docx template by factory."""
    if factory == "new":
        template_name = "SOP-GN-2002 Q 培训通知.docx"
        dir_name = "新厂人员培训管理规程"
    else:
        template_name = "培训通知.docx"
        dir_name = "旧厂员工培训教育管理规程"
    base = Path(__file__).resolve().parent.parent.parent.parent
    candidates = [
        Path(dir_name) / template_name,
        Path("..") / dir_name / template_name,
        base / dir_name / template_name,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {dir_name}/{template_name}")


def _find_underlined_groups(paragraph) -> list[list[int]]:
    """按连续段分组返回下划线 run 索引组."""
    groups: list[list[int]] = []
    current: list[int] = []
    for i, r in enumerate(paragraph.runs):
        if r.font.underline:
            current.append(i)
        else:
            if current:
                groups.append(current)
                current = []
    if current:
        groups.append(current)
    return groups


def _remove_empty_space_runs(paragraph) -> None:
    """删除纯空格且无下划线的尾随空 run，避免干扰下划线显示."""
    for r in paragraph.runs[:]:
        if not r.font.underline and r.text and r.text.strip() == "":
            paragraph._p.remove(r._r)


def _fill_first_underlined_run(paragraph, text: str) -> None:
    """在下划线区域内居中填入文字，保留原下划线总宽度。

    若文字长度 >= 下划线总宽：文字直接填入（自然延展下划线）。
    若文字长度 < 下划线总宽：左右填充空格使文字居中，保持原下划线宽度不变。
    """
    underlined = [r for r in paragraph.runs if r.font.underline]
    if not underlined:
        if not paragraph.runs:
            run = paragraph.add_run(text)
        else:
            run = paragraph.runs[0]
            run.text = text
        run.font.underline = True
        return

    total_width = sum(len(r.text) for r in underlined)
    text_len = len(text) if text else 0

    first = underlined[0]
    if text_len >= total_width:
        first.text = text or ""
    else:
        total_pad = total_width - text_len
        left_pad = total_pad // 2
        right_pad = total_pad - left_pad
        first.text = " " * left_pad + (text or "") + " " * right_pad
    first.font.underline = True

    # 删除其余下划线 runs
    for r in underlined[1:]:
        paragraph._p.remove(r._r)

    # 删除纯空格且无下划线的尾随空 run
    _remove_empty_space_runs(paragraph)


def _center_fill_group(paragraph, group_indices: list[int], text: str) -> None:
    """对 paragraph 中指定索引的一组下划线 run 居中填入文字。

    计算该组所有 run 的总字符宽度，将文字居中填入第一个下划线 run，
    不足部分用空格填充，保持原下划线总宽度不变。超长时直接填入。
    """
    total_width = sum(len(paragraph.runs[i].text) for i in group_indices)
    text_len = len(text) if text else 0

    first = paragraph.runs[group_indices[0]]
    if text_len >= total_width:
        first.text = text or ""
    else:
        total_pad = total_width - text_len
        left_pad = total_pad // 2
        right_pad = total_pad - left_pad
        first.text = " " * left_pad + (text or "") + " " * right_pad
    first.font.underline = True

    # 删除该组其余下划线 runs（从后往前避免索引错位）
    for idx in reversed(group_indices[1:]):
        paragraph._p.remove(paragraph.runs[idx]._r)


def generate_training_notification(data: TrainingNotificationInput, factory: str = "old") -> BytesIO:
    """根据填写的培训信息生成培训通知 Word 文档."""
    template_path = _find_template(factory)
    doc = Document(str(template_path))

    # ── P2: 部门 将于 年 月 日 举行 主题 的培训 ──
    p2 = doc.paragraphs[2]
    groups2 = _find_underlined_groups(p2)
    # groups2 = [[0,1], [3,4,5], [7,8,9], [11,12,13], [15,16,17], [19,20,21]]
    # 对应: 段首缩进 | 部门 | 年 | 月 | 日 | 主题
    if len(groups2) >= 6:
        # 段首缩进 runs[0-1] 不动

        # 居中填充各下划线组（从后往前处理避免索引错位）
        _center_fill_group(p2, groups2[5], data.subject or "")        # 主题
        _center_fill_group(p2, groups2[4], data.training_date.strftime("%d"))  # 日
        _center_fill_group(p2, groups2[3], data.training_date.strftime("%m"))  # 月
        _center_fill_group(p2, groups2[2], data.training_date.strftime("%Y"))  # 年
        _center_fill_group(p2, groups2[1], data.department)            # 部门

        # 修正固定文字（运行时索引已因 center_fill 删除 run 而改变）
        p2.runs[2].text = ""           # 清空原"主办部门"标签
        p2.runs[4].text = "部门将于"    # 原 run[6] "于"，前缀"部门"无下划线
        p2.runs[10].text = "日举行"      # 原 run[18] "举行"

    # ── P3: 培训时间 ──
    time_parts = []
    if data.training_time_start:
        time_parts.append(data.training_time_start)
    if data.training_time_end:
        time_parts.append(data.training_time_end)
    time_str = " ~ ".join(time_parts) if len(time_parts) == 2 else (time_parts[0] if time_parts else "")
    _fill_first_underlined_run(doc.paragraphs[3], time_str)

    # ── P4: 培训地点 ──
    _fill_first_underlined_run(doc.paragraphs[4], data.location or "")

    # ── P5: 培训师 ──
    _fill_first_underlined_run(doc.paragraphs[5], data.trainer or "")

    # ── P6: 培训内容 ──
    _fill_first_underlined_run(doc.paragraphs[6], data.content or "")

    # ── P7: 培训人员 ──
    people_str = "、".join(data.trainee_names) if data.trainee_names else ""
    _fill_first_underlined_run(doc.paragraphs[7], people_str)

    # ── P8-P10: 备注 — 固定内容，不动 ──
    # 模板中备注是固定的，不根据表单 remarks 字段覆盖

    # ── P12: 部门（落款） ──
    issuer = data.issuer_department or data.department or ""
    _fill_first_underlined_run(doc.paragraphs[12], issuer)

    # ── P13: 日期（落款） ──
    p13 = doc.paragraphs[13]
    groups13 = _find_underlined_groups(p13)
    # groups13 = [[1,2,3], [5,6,7], [9,10]] 对应 年 | 月 | 日
    # 固定文字 runs[4]="年"/[8]="月"/[11]="日" 保留原样不动
    issue_date = data.issue_date or data.training_date
    if len(groups13) >= 3:
        _center_fill_group(p13, groups13[2], issue_date.strftime("%d"))  # 日
        _center_fill_group(p13, groups13[1], issue_date.strftime("%m"))  # 月
        _center_fill_group(p13, groups13[0], issue_date.strftime("%Y"))  # 年

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
