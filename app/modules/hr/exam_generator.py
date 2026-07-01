"""AI 出题：根据文件内容生成题目，并导出试卷 Word 文档.

Moved from app/platform/ai/exam_generator.py to comply with the rule:
"业务功能代码（prompt、业务逻辑、错误处理）必须在模块内部，不得放到全局层"
"""

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pydantic import BaseModel

# ─── Pydantic schemas (previously in app/platform/ai/schemas.py) ───


class ChoiceOption(BaseModel):
    label: str
    text: str


class ChoiceQuestion(BaseModel):
    number: int
    question: str
    options: list[ChoiceOption]
    answer: str | None = None


class TrueFalseQuestion(BaseModel):
    number: int
    question: str
    answer: str | None = None


class ExamExportRequest(BaseModel):
    title: str
    examiner: str = ""
    exam_date: str = ""
    assessment_date: str = ""
    choice_questions: list[ChoiceQuestion] = []
    true_false_questions: list[TrueFalseQuestion] = []


# ─── Prompt template ───

_PROMPT_TEMPLATE = """你是一位专业的培训考核出题专家。请根据以下文件内容，生成一份新员工入职培训考核试卷。

要求：
1. 生成 5 道选择题，每题 10 分，共 50 分。每道题有 A、B、C、D 四个选项。
2. 生成 5 道判断题，每题 10 分，共 50 分。
3. 题目必须严格基于文件内容，不能编造文件中没有的信息。
4. 题目难度适中，适合新员工入职培训考核。
5. 每道题都要给出正确答案。

请严格按照以下 JSON 格式返回，不要包含任何其他文字：

{
  "choice_questions": [
    {
      "number": 1,
      "question": "题目内容",
      "options": [
        {"label": "A", "text": "选项内容"},
        {"label": "B", "text": "选项内容"},
        {"label": "C", "text": "选项内容"},
        {"label": "D", "text": "选项内容"}
      ],
      "answer": "A"
    }
  ],
  "true_false_questions": [
    {
      "number": 1,
      "question": "题目内容",
      "answer": "√"
    }
  ]
}

文件内容如下：

{content}
"""


def build_generate_prompt(file_content: str) -> str:
    """构建让 AI 根据文件内容出题的 prompt."""
    return _PROMPT_TEMPLATE.replace("{content}", file_content)


# ─── 字号常量 ───
SIZE_S4 = Pt(12)   # 小四
SIZE_4 = Pt(14)    # 四号
SIZE_X3 = Pt(15)   # 小三号


def _set_run_font(run, western: str, east_asian: str, size: Pt, bold: bool = False) -> None:
    """设置 run 的字体：西文字体 + 中文字体 + 字号 + 加粗."""
    run.font.name = western
    run.font.size = size
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), east_asian)


def _apply_para_format(
    para,
    space_before: Pt = Pt(0),
    space_after: Pt = Pt(0),
    line_spacing: float = 1.0,
    alignment=None,
    first_line_indent: Pt = Pt(0),
) -> None:
    """统一设置段落格式."""
    pf = para.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    pf.first_line_indent = first_line_indent
    if alignment is not None:
        para.alignment = alignment


def _add_para(
    doc,
    text: str,
    western: str = "Times New Roman",
    east_asian: str = "宋体",
    size: Pt = SIZE_S4,
    bold: bool = False,
    alignment=None,
    **fmt_kwargs,
) -> None:
    """添加一个段落并统一设置字体和段落格式."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, western, east_asian, size, bold)
    _apply_para_format(para, alignment=alignment, **fmt_kwargs)
    return para


def generate_exam_docx(data: ExamExportRequest) -> BytesIO:
    """根据试卷数据生成 Word 文档."""
    doc = Document()

    # ── 试卷标题（黑体 四号 居中 加粗）──
    _add_para(
        doc,
        data.title,
        western="Times New Roman",
        east_asian="黑体",
        size=SIZE_4,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # ── 表头：姓名、部门、分数 ──
    _add_para(
        doc,
        "姓 名：                部 门：                    分 数：        ",
        east_asian="宋体",
        size=SIZE_S4,
    )

    # ── 表头：出卷人、出卷时间、考核时间 ──
    _add_para(
        doc,
        f"出卷人：{data.examiner}    出卷时间：{data.exam_date}    考核时间：{data.assessment_date}",
        east_asian="宋体",
        size=SIZE_S4,
    )

    # 表头结束后空一行
    _add_para(doc, "", size=SIZE_S4, space_after=Pt(0))

    # ── 一、选择题 ──
    _add_para(
        doc,
        "一、选择题：（共 50 分，每题 10 分）",
        east_asian="宋体",
        size=SIZE_X3,
        bold=True,
    )

    for q in data.choice_questions:
        answer = q.answer or ""

        # 题干段落
        stem_para = _add_para(
            doc,
            f"{q.number}. {q.question}（{answer}  ）",
            east_asian="宋体",
            size=SIZE_S4,
        )
        # 题干段落后无额外间距，与选项紧密衔接
        _apply_para_format(stem_para, space_after=Pt(0))

        # 选项段落（四个选项横向排列）
        opt_para = doc.add_paragraph()
        _apply_para_format(opt_para, space_after=Pt(0))

        # 设置制表位，使选项均匀分布（约每 4cm 一个）
        tab_stops = opt_para.paragraph_format.tab_stops
        for i in range(1, len(q.options)):
            tab_stops.add_tab_stop(Cm(i * 4), alignment=WD_ALIGN_PARAGRAPH.LEFT)

        # 用制表符连接选项
        opt_text = "\t".join([f"{opt.label}. {opt.text}" for opt in q.options])
        opt_run = opt_para.add_run(opt_text)
        _set_run_font(opt_run, "Times New Roman", "宋体", SIZE_S4)

    # 选择题结束后空一行
    _add_para(doc, "", size=SIZE_S4, space_after=Pt(0))

    # ── 二、判断题 ──
    _add_para(
        doc,
        "二、判断题：（共 50 分，每题 10 分）",
        east_asian="宋体",
        size=SIZE_X3,
        bold=True,
    )

    for q in data.true_false_questions:
        answer = q.answer or ""
        _add_para(
            doc,
            f"{q.number}. {q.question}（{answer}  ）",
            east_asian="宋体",
            size=SIZE_S4,
        )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
