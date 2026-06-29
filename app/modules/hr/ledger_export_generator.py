"""培训台账导出文档生成器 — 基于模板填入数据."""

from datetime import date
from io import BytesIO
from pathlib import Path

import openpyxl
from openpyxl.styles import Alignment, Font, Border, Side
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn


def _find_template(dir_name: str, file_name: str) -> Path:
    base = Path(__file__).resolve().parent.parent.parent.parent
    candidates = [
        Path(dir_name) / file_name,
        Path("..") / dir_name / file_name,
        base / dir_name / file_name,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板未找到: {dir_name}/{file_name}")


def _fmt_date(value) -> str:
    if not value:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%Y.%m.%d")
    if isinstance(value, str):
        from datetime import datetime
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
            try:
                return datetime.strptime(value, fmt).strftime("%Y.%m.%d")
            except ValueError:
                continue
    return str(value)


def _set_docx_cell(cell, text: str) -> None:
    """Fill a docx cell with 宋体小四 centered text."""
    tc = cell._tc
    for p in tc.findall(qn('w:p')):
        tc.remove(p)
    safe = str(text or "")
    escaped = safe.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    p_xml = (
        f'<w:p {nsdecls("w")}>'
        f'<w:pPr><w:jc w:val="center"/></w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/>'
        f'<w:sz w:val="24"/>'
        f'<w:szCs w:val="24"/>'
        f'</w:rPr>'
        f'<w:t>{escaped}</w:t>'
        f'</w:r>'
        f'</w:p>'
    )
    p = parse_xml(p_xml)
    tc.append(p)


def _fill_old_event_ledger(employee: dict, records: list[dict]) -> BytesIO:
    """旧厂事件台账：7.9员工培训台账（非文件）.xlsx"""
    wb = openpyxl.load_workbook(str(_find_template("员工培训教育管理规程", "7.9员工培训台账（非文件）.xlsx")))
    ws = wb.active

    # 姓名 (A4:C4 merged → write to A4)
    ws["A4"] = employee.get("name", "")
    # 性别 (G4:I4 merged)
    ws["G4"] = employee.get("gender", "")
    # 工作卡号 (J4:M4 merged)
    ws["J4"] = employee.get("employee_number", "")
    # 部门 (A5:C5 merged → write to A5)
    ws["A5"] = employee.get("department", "")
    # 岗位/职务 (G5:I5 merged)
    ws["G5"] = employee.get("position", "")
    # 入厂时间 (L5:M5 merged)
    ws["L5"] = employee.get("factory_entry_date") or employee.get("hire_date", "")
    # 岗位变动 (D6:M7 merged)
    ws["D6"] = employee.get("transfer_history", "")
    # 岗位变动 (D6:M7 merged)
    ws["D6"] = employee.get("transfer_history", "")

    # Training records starting from row 8
    for idx, record in enumerate(records):
        row = 10 + idx  # data rows start at row 10 (after header rows 8-9)
        # Year/Month (A, B)
        train_date = record.get("training_date", "")
        if train_date:
            parts = str(train_date).split("-")
            if len(parts) >= 2:
                ws.cell(row=row, column=1).value = parts[0]  # year
                ws.cell(row=row, column=2).value = parts[1]  # month
            if len(parts) >= 3:
                ws.cell(row=row, column=2).value = f"{parts[1]}.{parts[2]}"  # month.day
        # 培训课程 (C8:G9 merged → write to C)
        ws.cell(row=row, column=3).value = record.get("training_subject", "")
        # 培训方式 (H)
        ws.cell(row=row, column=8).value = record.get("training_method", "")
        # 课时 (I)
        ws.cell(row=row, column=9).value = record.get("duration_hours", "")
        # 培训单位/培训师 (K)
        ws.cell(row=row, column=11).value = record.get("trainer", "")
        # 考核成绩 (M)
        ws.cell(row=row, column=13).value = record.get("assessment_result", "")

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _fill_old_sop_ledger(employee: dict, records: list[dict]) -> BytesIO:
    """旧厂SOP台账：7.10员工文件培训台账.xlsx"""
    wb = openpyxl.load_workbook(str(_find_template("员工培训教育管理规程", "7.10员工文件培训台账.xlsx")))
    ws = wb.active

    # 姓名 (A4:B4 merged)
    ws["A4"] = employee.get("name", "")
    # 部门 (C4:D4 merged)
    ws["C4"] = employee.get("department", "")
    # 职务 (E4:F4 merged)
    ws["E4"] = employee.get("position", "")
    # 工作卡号 (G4:H4 merged)
    ws["G4"] = employee.get("employee_number", "")

    # Records starting from row 6
    for idx, record in enumerate(records):
        row = 6 + idx
        # 培训日期 (A)
        ws.cell(row=row, column=1).value = _fmt_date(record.get("training_date", ""))
        # 文件名 (B:C merged) — use training_subject
        ws.cell(row=row, column=2).value = record.get("training_subject", "")
        # 文件编码 (D:E merged)
        ws.cell(row=row, column=4).value = record.get("source_id", "")
        # 生效日期 (F)
        ws.cell(row=row, column=6).value = ""
        # 培训师 (G)
        ws.cell(row=row, column=7).value = record.get("trainer", "")
        # 培训方式 (H)
        ws.cell(row=row, column=8).value = record.get("training_method", "")
        # 备注 (I)
        ws.cell(row=row, column=9).value = record.get("remarks", "")

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _fill_new_sop_ledger(employee: dict, records: list[dict]) -> BytesIO:
    """新厂SOP台账：R-GN-2002 I 员工SOP培训台账.docx"""
    doc = Document(str(_find_template("新厂人员培训管理规程", "R-GN-2002 I 员工SOP培训台账.docx")))
    table = doc.tables[0]

    # Row 0: employee info
    _set_docx_cell(table.rows[0].cells[1], employee.get("name", ""))
    _set_docx_cell(table.rows[0].cells[3], employee.get("department", ""))
    _set_docx_cell(table.rows[0].cells[5], employee.get("position", ""))
    _set_docx_cell(table.rows[0].cells[7], employee.get("employee_number", ""))

    # Records starting from row 2
    for idx, record in enumerate(records):
        row_idx = 2 + idx
        if row_idx >= len(table.rows):
            break
        row = table.rows[row_idx]
        _set_docx_cell(row.cells[0], record.get("source_id", ""))
        _set_docx_cell(row.cells[1], record.get("training_subject", ""))
        _set_docx_cell(row.cells[2], "")
        _set_docx_cell(row.cells[3], "")
        _set_docx_cell(row.cells[4], record.get("training_method", ""))
        _set_docx_cell(row.cells[5], _fmt_date(record.get("training_date", "")))
        _set_docx_cell(row.cells[6], record.get("trainer", ""))
        _set_docx_cell(row.cells[7], record.get("remarks", ""))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _fill_new_event_ledger(employee: dict, records: list[dict]) -> BytesIO:
    """新厂事件台账：R-GN-2002 J 员工事件培训台账.docx"""
    doc = Document(str(_find_template("新厂人员培训管理规程", "R-GN-2002 J 员工事件培训台账.docx")))
    table = doc.tables[0]

    # Row 0: 姓名, 性别, 岗位/职务
    _set_docx_cell(table.rows[0].cells[2], employee.get("name", ""))
    _set_docx_cell(table.rows[0].cells[4], employee.get("gender", ""))
    _set_docx_cell(table.rows[0].cells[9], employee.get("position", ""))
    # Row 1: 部门, 工作卡号, 报到日期
    _set_docx_cell(table.rows[1].cells[2], employee.get("department", ""))
    _set_docx_cell(table.rows[1].cells[4], employee.get("employee_number", ""))
    _set_docx_cell(table.rows[1].cells[9], employee.get("hire_date", ""))
    # Row 2: 岗位变动
    _set_docx_cell(table.rows[2].cells[2], employee.get("transfer_history", ""))

    # Records starting from row 5 (after headers at rows 3-4)
    for idx, record in enumerate(records):
        row_idx = 5 + idx
        if row_idx >= len(table.rows):
            break
        row = table.rows[row_idx]
        train_date = str(record.get("training_date", ""))
        parts = train_date.split("-") if train_date else []
        _set_docx_cell(row.cells[0], parts[0] if len(parts) > 0 else "")
        _set_docx_cell(row.cells[1], parts[1] if len(parts) > 1 else "")
        _set_docx_cell(row.cells[2], record.get("training_subject", ""))
        _set_docx_cell(row.cells[5], record.get("training_method", ""))
        _set_docx_cell(row.cells[7], str(record.get("duration_hours", "")))
        _set_docx_cell(row.cells[9], record.get("trainer", ""))
        _set_docx_cell(row.cells[10], record.get("assessment_result", ""))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_ledger_export(employee: dict, records: list[dict], factory: str = "old", ledger_type: str = "event") -> BytesIO:
    """根据厂区和台账类型选择模板并填入数据."""
    if factory == "new":
        if ledger_type == "sop":
            return _fill_new_sop_ledger(employee, records)
        return _fill_new_event_ledger(employee, records)
    else:
        if ledger_type == "sop":
            return _fill_old_sop_ledger(employee, records)
        return _fill_old_event_ledger(employee, records)
