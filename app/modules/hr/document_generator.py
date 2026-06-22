"""Generate onboarding training record documents from templates."""

from io import BytesIO
from pathlib import Path

from docx import Document

from app.modules.hr.models import Employee


OLD_TEMPLATE_NAME = "7.3新员工入职培训记录.docx"
NEW_TEMPLATE_NAME = "R-GN-2002 B 新员工入职培训记录.docx"


def _find_template(factory: str = "old") -> Path:
    """Locate the docx template, trying several path candidates."""
    if factory == "new":
        candidates = [
            Path("新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
            Path("../新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
            Path(__file__).resolve().parent.parent.parent.parent
            / "新厂人员培训管理规程"
            / NEW_TEMPLATE_NAME,
        ]
    else:
        candidates = [
            Path("员工培训教育管理规程") / OLD_TEMPLATE_NAME,
            Path("../员工培训教育管理规程") / OLD_TEMPLATE_NAME,
            Path(__file__).resolve().parent.parent.parent.parent
            / "员工培训教育管理规程"
            / OLD_TEMPLATE_NAME,
        ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {NEW_TEMPLATE_NAME if factory == 'new' else OLD_TEMPLATE_NAME}")


def generate_onboarding_training_record(employee: Employee, factory: str = "old") -> BytesIO:
    """Fill the onboarding training record template with employee data.

    Returns a BytesIO buffer containing the generated docx.
    """
    template_path = _find_template(factory)
    doc = Document(str(template_path))

    if not doc.tables:
        raise ValueError("模板中未找到表格")

    table = doc.tables[0]

    if factory == "new":
        # New factory template: 24 rows x 6 cols
        # Row 1: [姓名] [] [性别] [] [工作卡号] []
        table.rows[1].cells[1].text = employee.name or ""
        table.rows[1].cells[3].text = employee.gender or ""
        table.rows[1].cells[5].text = employee.employee_number or ""

        # Row 2: [部门] [] [] [拟定岗位] [] []
        table.rows[2].cells[1].text = employee.department or ""
        table.rows[2].cells[4].text = employee.position or ""

        # Row 3: [报到日期] [] [] [拟定转正日期] [] []
        hire_date_str = str(employee.hire_date) if employee.hire_date else ""
        table.rows[3].cells[1].text = hire_date_str
        table.rows[3].cells[4].text = ""
    else:
        # Old factory template: 21 rows x 12 cols
        # Row 1: 姓名 | 01 | 性别 | 02... | 工作卡号 | 03...
        table.rows[1].cells[1].text = employee.name or ""
        for idx in (3, 4, 5):
            table.rows[1].cells[idx].text = employee.gender or ""
        for idx in (9, 10, 11):
            table.rows[1].cells[idx].text = employee.employee_number or ""

        # Row 2: 部门 | 04... | 拟定岗位 | 05...
        for idx in (1, 2):
            table.rows[2].cells[idx].text = employee.department or ""
        for idx in (6, 7, 8, 9, 10, 11):
            table.rows[2].cells[idx].text = employee.position or ""

        # Row 3: 报到日期 | 06... | 转正日期 | 07...
        hire_date_str = str(employee.hire_date) if employee.hire_date else ""
        for idx in (1, 2):
            table.rows[3].cells[idx].text = hire_date_str
        # 转正日期数据库中无对应字段，留空
        for idx in (6, 7, 8, 9, 10, 11):
            table.rows[3].cells[idx].text = ""

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
