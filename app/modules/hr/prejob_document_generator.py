"""Generate pre-job training plan documents from templates."""

from io import BytesIO
from pathlib import Path

import openpyxl
from docx import Document

from app.modules.hr.models import Employee

OLD_TEMPLATE_NAME = "7.4岗前培训计划.xlsx"
NEW_TEMPLATE_NAME = "R-GN-2002 C 岗前培训计划.docx"


def _find_old_template() -> Path:
    """Locate the old factory xlsx template."""
    candidates = [
        Path("员工培训教育管理规程") / OLD_TEMPLATE_NAME,
        Path("../员工培训教育管理规程") / OLD_TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / OLD_TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {OLD_TEMPLATE_NAME}")


def _find_new_template() -> Path:
    """Locate the new factory docx template."""
    candidates = [
        Path("新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
        Path("../新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent.parent
        / "新厂人员培训管理规程"
        / NEW_TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {NEW_TEMPLATE_NAME}")


DEPT_CONTENT_MAP: dict[str, list[str]] = {
    "人事行政部": [
        "公司级公用文件(详见附件一)",
        "部门级公用文件(详见附件二)",
        "人事行政部人事行政专员岗位文件(详见附件三)",
        "人事行政专员岗位职责(QP.PM.053)",
        "生产安全知识",
        "岗前培训计划",
    ],
}


def _generate_old(employee: Employee, items: list[dict] | None = None) -> BytesIO:
    """Fill the old factory pre-job training plan xlsx template."""
    template_path = _find_old_template()
    wb = openpyxl.load_workbook(str(template_path))
    ws = wb.active

    # Part 1: Employee overview
    ws["C5"] = employee.name or ""
    ws["I5"] = employee.department or ""
    ws["C6"] = employee.employee_number or ""
    ws["I6"] = str(employee.hire_date) if employee.hire_date else ""
    ws["C7"] = employee.position or ""

    # Part 2: Training content — use provided items or fallback to department map
    if items:
        for item in items:
            seq = item.get("seq", 0)
            if 1 <= seq <= 10:
                # Only fill content column for old factory (B column)
                row = 10 + seq  # seq 1 → row 11, seq 10 → row 20
                ws[f"B{row}"] = item.get("content", "")
    else:
        content_list = DEPT_CONTENT_MAP.get(employee.department or "", [])
        for i, content in enumerate(content_list):
            row = 11 + i
            if row <= 20:
                ws[f"B{row}"] = content

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def _generate_new(employee: Employee, items: list[dict] | None = None) -> BytesIO:
    """Fill the new factory pre-job training plan docx template."""
    template_path = _find_new_template()
    doc = Document(str(template_path))

    if not doc.tables:
        raise ValueError("模板中未找到表格")

    table = doc.tables[0]

    # Row 1: [姓名] [姓名] [] [] [部门] []
    table.rows[1].cells[2].text = employee.name or ""
    table.rows[1].cells[5].text = employee.department or ""

    # Row 2: [学历] [学历] [] [] [毕业院校] []
    table.rows[2].cells[2].text = employee.education or ""
    table.rows[2].cells[5].text = employee.school or ""

    # Row 3: [毕业时间] [毕业时间] [] [] [工作卡号] []
    grad_date_str = str(employee.graduation_date) if employee.graduation_date else ""
    table.rows[3].cells[2].text = grad_date_str
    table.rows[3].cells[5].text = employee.employee_number or ""

    # Row 4: [报到日期] [报到日期] [] [] [拟定岗位] []
    hire_date_str = str(employee.hire_date) if employee.hire_date else ""
    table.rows[4].cells[2].text = hire_date_str
    table.rows[4].cells[5].text = employee.position or ""

    # Part 2: Training plan rows (rows 6-15 in the template table, 10 rows)
    if items:
        for item in items:
            seq = item.get("seq", 0)
            if 1 <= seq <= 10:
                row_idx = 5 + seq  # seq 1 → row 6
                row = table.rows[row_idx]
                # Assuming cell layout: [seq, content, deadline, trainer]
                if len(row.cells) >= 4:
                    row.cells[1].text = item.get("content", "")
                    row.cells[2].text = item.get("deadline", "")
                    row.cells[3].text = item.get("trainer", "")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_prejob_training_plan(
    employee: Employee, factory: str = "old", items: list[dict] | None = None
) -> BytesIO:
    """Fill the pre-job training plan template with employee data.

    Args:
        employee: Employee record to pre-fill.
        factory: 'old' or 'new'.
        items: Optional list of {seq, content, deadline, trainer} dicts.
               When provided, fills the training plan table from these items.
               When None, falls back to DEPT_CONTENT_MAP.

    Returns a BytesIO buffer containing the generated document.
    """
    if factory == "new":
        return _generate_new(employee, items)
    return _generate_old(employee, items)
