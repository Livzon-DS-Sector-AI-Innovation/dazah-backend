"""根据 CDE 通知函解析结果生成发补回复 Word 文档"""

import io
import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)


def _set_run_font(run, size=12, bold=False, color=None):
    """设置字体"""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if color:
        run.font.color.rgb = color


def _add_para(doc, text, size=12, bold=False, color=None, align=None, indent=0):
    """添加段落"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, size=size, bold=bold, color=color)
    if align:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def _add_cover_page(doc, drug_info: dict):
    """生成封面页"""
    for _ in range(6):
        doc.add_paragraph()

    _add_para(doc, "原料药登记资料", size=26, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    _add_para(doc, f"资料类型：{drug_info.get('doc_type', '补充资料')}",
              size=16, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc, f"产品名称：{drug_info.get('drug_name', '')}",
              size=16, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc, f"登 记 号：{drug_info.get('registration_number', '')}",
              size=16, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc, f"受 理 号：{drug_info.get('acceptance_number', '')}",
              size=16, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc, f"申请人：{drug_info.get('company_name', '')}",
              size=16, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc, f"申请事项：{drug_info.get('application_type', '首次登记')}",
              size=16, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc, f"联 系 人：{drug_info.get('contact', '魏永红')}",
              size=16, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc, f"联系电话：{drug_info.get('phone', '0756-8686208')}",
              size=16, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    _add_para(doc,
              f"注册地址：{drug_info.get('address', '中国，广东省，珠海市，珠海保税区联峰路22号')}",
              size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc,
              f"邮政编码：{drug_info.get('zipcode', '519030')}",
              size=14, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    _add_para(doc, f"关联制剂受理号：{drug_info.get('related_no', '/')}",
              size=16, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_para(doc,
              f"Email：{drug_info.get('email', 'lzsyntpharmzhucebu@livzon.cn')}",
              size=16, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    doc.add_paragraph()
    _add_para(doc, drug_info.get("drug_name", ""),
              size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "补充资料研究报告",
              size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_toc_page(doc, questions: list[dict]):
    """生成目录页"""
    _add_para(doc, "目  录", size=14, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    for q in questions:
        if q["level"] == 1:
            _add_para(doc, q["title"], size=12, bold=True)
        elif q["level"] == 2:
            _add_para(doc, q["title"], size=12, bold=True,
                      color=RGBColor(0, 112, 192), indent=0.74)
        elif q["level"] == 3:
            _add_para(doc, q["title"], size=12, bold=True,
                      color=RGBColor(0, 112, 192), indent=1.48)


def _add_body_content(doc, questions: list[dict]):
    """生成正文内容"""
    for q in questions:
        if q["level"] == 1:
            _add_para(doc, q["title"], size=12, bold=True)
            continue

        # 问题原文（蓝色加粗）
        _add_para(doc, q["title"], size=12, bold=True,
                  color=RGBColor(0, 0, 255))
        doc.add_paragraph()

        # 答复
        _add_para(doc, "【答复】", size=12, bold=True)
        _add_para(doc, "（请在此处填写回复内容）", size=12, indent=0.74)
        doc.add_paragraph()

        # 回复建议
        _add_para(doc, "【回复建议】", size=12, bold=True)
        _add_para(doc, "• 回复思路：请分析CDE关注点，制定回复策略", size=12,
                  indent=0.74)
        _add_para(doc, "• 回复模板：请参照CDE通知函原文进行回复", size=12,
                  indent=0.74)
        _add_para(doc, "• 补充实验：视具体情况而定", size=12, indent=0.74)
        _add_para(doc, "• 法规依据：ICH Q系列；中国药典；相关技术指导原则",
                  size=12, indent=0.74)
        doc.add_paragraph()


def generate_reply_document(
    drug_info: dict,
    questions: list[dict],
) -> bytes:
    """
    生成发补回复 Word 文档，返回文件二进制内容。

    Args:
        drug_info: 药品信息（drug_name, registration_number, acceptance_number 等）
        questions: 问题列表（从 PDF 解析得到）
    """
    doc = Document()

    # 页面设置
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    # 封面页
    _add_cover_page(doc, drug_info)
    doc.add_page_break()

    # 目录页
    _add_toc_page(doc, questions)
    doc.add_page_break()

    # 正文
    _add_body_content(doc, questions)

    # 输出为 bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
