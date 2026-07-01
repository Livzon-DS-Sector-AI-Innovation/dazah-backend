#!/usr/bin/env python3
"""
根据COA数据生成对照物质说明表Word文档
"""

import logging
import sys

logger = logging.getLogger(__name__)
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt

# 固定信息
FIXED_INFO = {
    '提供单位': '珠海保税区丽珠合成制药有限公司',
    '经办人': '魏永红',
    '联系方式': '13570680132',
}

# 模板路径
TEMPLATE_PATH = 'assets/对照物质说明表模板.docx'


def set_cell_format(cell, text, align_center=True, is_formula=False):
    """设置单元格格式
    is_formula: 是否为分子式，需要处理下标数字
    """
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 清空单元格
    cell.text = ''
    paragraph = cell.paragraphs[0]

    if is_formula and text:
        # 处理分子式，数字转为下标
        import re
        # 匹配字母和数字，数字部分设为下标
        parts = re.split(r'(\d+)', text)
        for part in parts:
            run = paragraph.add_run(part)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 如果是数字，设为下标
            if part.isdigit():
                run.font.subscript = True
    else:
        paragraph.text = text
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置对齐方式
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def generate_document(coa_data, output_path):
    """生成对照物质说明表"""

    # 打开模板
    doc = Document(TEMPLATE_PATH)

    # 获取三个表格
    table1 = doc.tables[0]  # 基本信息
    table2 = doc.tables[1]  # 详细信息
    table3 = doc.tables[2]  # 签章

    # ===== 表格1：基本信息 =====
    # R1C2-4: 药品名称
    set_cell_format(table1.rows[0].cells[1], coa_data.get('药品名称', ''))

    # R2C2: 对照物质名称
    set_cell_format(table1.rows[1].cells[1], coa_data.get('对照物质名称', ''))

    # R2C4: 批号
    set_cell_format(table1.rows[1].cells[3], coa_data.get('批号', ''))

    # R3C2-4: 来源（COA生产厂家）
    manufacturer = coa_data.get('生产厂家', '')
    set_cell_format(table1.rows[2].cells[1], f"☑以下生产厂家或研究单位研制：{manufacturer}", align_center=False)

    # ===== 表格2：详细信息 =====
    # R1C2: 英文名
    set_cell_format(table2.rows[0].cells[1], coa_data.get('英文名', ''))

    # R1C4: 有效期
    set_cell_format(table2.rows[0].cells[3], coa_data.get('有效期', ''))

    # R2C2: 分子式（数字下标）
    set_cell_format(table2.rows[1].cells[1], coa_data.get('分子式', ''), is_formula=True)

    # R2C4: 分子量
    set_cell_format(table2.rows[1].cells[3], coa_data.get('分子量', ''))

    # R5C2-4: 使用范围（勾选含量测定）
    set_cell_format(table2.rows[4].cells[1], "□鉴别          □检查         ☑含量测定", align_center=False)

    # R6C2-4: 含量
    content = f"含量＝ {coa_data.get('含量', '')} ％      RSD％＝\n（测定方法——                                     ）"
    set_cell_format(table2.rows[5].cells[1], content, align_center=False)

    # R7C2-4: 贮存条件
    storage = coa_data.get('贮存条件', '冷藏')
    if '冷藏' in storage or '2-8' in storage:
        set_cell_format(table2.rows[6].cells[1], "□避光   □常温   □阴凉   ☑冷藏   □冷冻   □其他", align_center=False)
    elif '冷冻' in storage:
        set_cell_format(table2.rows[6].cells[1], "□避光   □常温   □阴凉   □冷藏   ☑冷冻   □其他", align_center=False)
    else:
        set_cell_format(table2.rows[6].cells[1], "□避光   □常温   □阴凉   □冷藏   □冷冻   □其他", align_center=False)

    # R8C2-4: 使用方法（勾选直接折算）
    set_cell_format(table2.rows[7].cells[1], "□使用前干燥    ☑直接折算   □其他", align_center=False)

    # ===== 表格3：签章 =====
    # R2C2-4: 提供单位（固定信息）
    provider_text = f"{FIXED_INFO['提供单位']}\n\n(公章)"
    set_cell_format(table3.rows[1].cells[1], provider_text)

    # R3C2: 经办人签名（固定信息）
    set_cell_format(table3.rows[2].cells[1], FIXED_INFO['经办人'])

    # R3C4: 联系方式（固定信息）
    set_cell_format(table3.rows[2].cells[3], FIXED_INFO['联系方式'])

    # 保存文档
    doc.save(output_path)
    logger.info(f"文档已生成: {output_path}")


if __name__ == '__main__':
    # 测试数据
    test_data = {
        '药品名称': 'Paliperidone Palmitate',
        '对照物质名称': 'Paliperidone Palmitate',
        '批号': 'U4O-1601280-04',
        '生产厂家': 'OST Research Chemicals Inc.',
        '英文名': 'Paliperidone Palmitate',
        '有效期': '2026-10-11',
        '分子式': 'C39H57FN4O4',
        '分子量': '664.91',
        '含量': '99.03',
        '贮存条件': '2-8℃',
    }

    output = sys.argv[1] if len(sys.argv) > 1 else 'output.docx'
    generate_document(test_data, output)
