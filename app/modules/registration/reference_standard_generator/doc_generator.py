"""根据 COA 数据生成对照物质说明表 Word 文档"""

import io
import logging
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt

logger = logging.getLogger(__name__)

# 固定信息
FIXED_INFO = {
    "provider": "珠海保税区丽珠合成制药有限公司",
    "handler": "魏永红",
    "contact": "13570680132",
}


def _set_cell_format(cell, text, align_center=True, is_formula=False):
    """设置单元格格式，is_formula 为 True 时分子式数字转为下标"""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 清空单元格
    cell.text = ""
    paragraph = cell.paragraphs[0]

    if is_formula and text:
        parts = re.split(r"(\d+)", text)
        for part in parts:
            run = paragraph.add_run(part)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if part.isdigit():
                run.font.subscript = True
    else:
        paragraph.text = text
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def generate_reference_standard_document(
    coa_data: dict,
    template_data: bytes,
) -> bytes:
    """
    根据 COA 数据和模板生成对照物质说明表，返回 Word 文件二进制内容。

    Args:
        coa_data: 从 COA 提取的数据字典
        template_data: Word 模板文件二进制内容
    """
    doc = Document(io.BytesIO(template_data))

    tables = doc.tables
    if len(tables) < 3:
        raise ValueError("模板格式不正确，至少需要3个表格")

    table1 = tables[0]  # 基本信息
    table2 = tables[1]  # 详细信息
    table3 = tables[2]  # 签章

    # ===== 表格1：基本信息 =====
    # R1C2-4: 药品名称
    if len(table1.rows) > 0 and len(table1.rows[0].cells) > 1:
        _set_cell_format(
            table1.rows[0].cells[1], coa_data.get("药品名称", "")
        )

    # R2C2: 对照物质名称
    if len(table1.rows) > 1 and len(table1.rows[1].cells) > 1:
        _set_cell_format(
            table1.rows[1].cells[1], coa_data.get("对照物质名称", "")
        )

    # R2C4: 批号
    if len(table1.rows) > 1 and len(table1.rows[1].cells) > 3:
        _set_cell_format(
            table1.rows[1].cells[3], coa_data.get("批号", "")
        )

    # R3C2-4: 来源（COA生产厂家）
    manufacturer = coa_data.get("生产厂家", "")
    if len(table1.rows) > 2 and len(table1.rows[2].cells) > 1:
        _set_cell_format(
            table1.rows[2].cells[1],
            f"☑以下生产厂家或研究单位研制：{manufacturer}",
            align_center=False,
        )

    # ===== 表格2：详细信息 =====
    # R1C2: 英文名
    if len(table2.rows) > 0 and len(table2.rows[0].cells) > 1:
        _set_cell_format(
            table2.rows[0].cells[1], coa_data.get("英文名", "")
        )

    # R1C4: 有效期
    if len(table2.rows) > 0 and len(table2.rows[0].cells) > 3:
        _set_cell_format(
            table2.rows[0].cells[3], coa_data.get("有效期", "")
        )

    # R2C2: 分子式（数字下标）
    if len(table2.rows) > 1 and len(table2.rows[1].cells) > 1:
        _set_cell_format(
            table2.rows[1].cells[1],
            coa_data.get("分子式", ""),
            is_formula=True,
        )

    # R2C4: 分子量
    if len(table2.rows) > 1 and len(table2.rows[1].cells) > 3:
        _set_cell_format(
            table2.rows[1].cells[3], coa_data.get("分子量", "")
        )

    # R5C2-4: 使用范围（勾选含量测定）
    if len(table2.rows) > 4 and len(table2.rows[4].cells) > 1:
        _set_cell_format(
            table2.rows[4].cells[1],
            "□鉴别          □检查         ☑含量测定",
            align_center=False,
        )

    # R6C2-4: 含量和 RSD
    if len(table2.rows) > 5 and len(table2.rows[5].cells) > 1:
        content_value = coa_data.get('含量', '')
        rsd_value = coa_data.get('RSD', '')
        moisture_value = coa_data.get('水分/干燥失重', '')

        content_text = f"含量＝ {content_value} ％      RSD％＝ {rsd_value}\n"
        if moisture_value:
            content_text += f"水分/干燥失重＝ {moisture_value} ％\n"
        content_text += "（测定方法——                                     ）"

        _set_cell_format(
            table2.rows[5].cells[1], content_text, align_center=False
        )

    # R7C2-4: 贮存条件
    if len(table2.rows) > 6 and len(table2.rows[6].cells) > 1:
        storage = coa_data.get("贮存条件", "")
        if "冷藏" in storage or "2-8" in storage:
            storage_text = "□避光   □常温   □阴凉   ☑冷藏   □冷冻   □其他"
        elif "冷冻" in storage:
            storage_text = "□避光   □常温   □阴凉   □冷藏   ☑冷冻   □其他"
        elif "常温" in storage or "室温" in storage:
            storage_text = "□避光   ☑常温   □阴凉   □冷藏   □冷冻   □其他"
        elif "阴凉" in storage:
            storage_text = "□避光   □常温   ☑阴凉   □冷藏   □冷冻   □其他"
        else:
            storage_text = "□避光   □常温   □阴凉   □冷藏   □冷冻   □其他"
        _set_cell_format(
            table2.rows[6].cells[1], storage_text, align_center=False
        )

    # R8C2-4: 使用方法（勾选直接折算）
    if len(table2.rows) > 7 and len(table2.rows[7].cells) > 1:
        _set_cell_format(
            table2.rows[7].cells[1],
            "□使用前干燥    ☑直接折算   □其他",
            align_center=False,
        )

    # ===== 表格3：签章 =====
    # R2C2-4: 提供单位
    if len(table3.rows) > 1 and len(table3.rows[1].cells) > 1:
        provider_text = f"{FIXED_INFO['provider']}\n\n(公章)"
        _set_cell_format(table3.rows[1].cells[1], provider_text)

    # R3C2: 经办人签名
    if len(table3.rows) > 2 and len(table3.rows[2].cells) > 1:
        _set_cell_format(table3.rows[2].cells[1], FIXED_INFO["handler"])

    # R3C4: 联系方式
    if len(table3.rows) > 2 and len(table3.rows[2].cells) > 3:
        _set_cell_format(table3.rows[2].cells[3], FIXED_INFO["contact"])

    # 输出为 bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
