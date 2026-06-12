#!/usr/bin/env python3
"""
模板解析器 - 从公司模板Word文档中提取格式规范
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

class TemplateParser:
    def __init__(self, template_path):
        self.doc = Document(template_path)
        self.format_specs = {
            'header': {},
            'fonts': {},
            'paragraphs': {},
            'headings': {},
            'cover': {}
        }
    
    def parse_header(self):
        """解析页眉格式"""
        section = self.doc.sections[0]
        if section.header:
            header = section.header
            header_texts = []
            
            for para in header.paragraphs:
                if para.text.strip():
                    header_texts.append(para.text.strip())
                    # 分析第一个有文字的段落
                    if para.runs:
                        run = para.runs[0]
                        self.format_specs['header'] = {
                            'font_name': run.font.name,
                            'font_size': run.font.size.pt if run.font.size else 12,
                            'bold': run.font.bold,
                            'texts': header_texts
                        }
        
        return self.format_specs['header']
    
    def parse_fonts(self):
        """解析字体设置"""
        fonts_found = {}
        
        for para in self.doc.paragraphs[:50]:  # 检查前50段
            for run in para.runs:
                if run.text.strip():
                    font_name = run.font.name
                    font_size = run.font.size.pt if run.font.size else None
                    bold = run.font.bold
                    
                    # 检测蓝色文字（问题文本）
                    if run.font.color and run.font.color.rgb:
                        rgb = run.font.color.rgb
                        if rgb[2] > rgb[0] and rgb[2] > rgb[1]:  # 蓝色分量最大
                            fonts_found['question_color'] = rgb
                    
                    if font_name and font_size:
                        key = f"{font_name}_{font_size}_{bold}"
                        if key not in fonts_found:
                            fonts_found[key] = {
                                'font_name': font_name,
                                'font_size': font_size,
                                'bold': bold,
                                'sample': run.text[:20]
                            }
        
        self.format_specs['fonts'] = fonts_found
        return fonts_found
    
    def parse_paragraph_format(self):
        """解析段落格式"""
        for para in self.doc.paragraphs[:30]:
            if para.text.strip():
                self.format_specs['paragraphs'] = {
                    'line_spacing': para.paragraph_format.line_spacing,
                    'space_before': para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0,
                    'space_after': para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0,
                    'first_line_indent': para.paragraph_format.first_line_indent.cm if para.paragraph_format.first_line_indent else 0,
                    'alignment': para.alignment
                }
                break
        
        return self.format_specs['paragraphs']
    
    def parse_headings(self):
        """解析标题样式"""
        heading_patterns = {
            'level1': [],  # 1、 2、
            'level2': [],  # （1）（2）
            'level3': []   # ① ②
        }
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检测一级标题: 1、 2、
            if re.match(r'^\d+[、\.]', text):
                if para.runs:
                    run = para.runs[0]
                    heading_patterns['level1'] = {
                        'prefix': 'number_dun',
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else 12,
                        'bold': run.font.bold,
                        'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                    }
            
            # 检测二级标题: （1）（2）
            elif re.match(r'^（\d+）', text):
                if para.runs:
                    run = para.runs[0]
                    heading_patterns['level2'] = {
                        'prefix': 'bracket_number',
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else 12,
                        'bold': run.font.bold,
                        'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                    }
            
            # 检测标签: 【答复】【回复建议】
            elif '【' in text and '】' in text:
                if para.runs:
                    run = para.runs[0]
                    heading_patterns['label'] = {
                        'font_name': run.font.name,
                        'font_size': run.font.size.pt if run.font.size else 12,
                        'bold': run.font.bold
                    }
        
        self.format_specs['headings'] = heading_patterns
        return heading_patterns
    
    def parse_cover_page(self):
        """解析封面页"""
        cover_info = {
            'title': None,
            'subtitle': None,
            'fields': []
        }
        
        # 检查前20段（通常在封面）
        for para in self.doc.paragraphs[:20]:
            text = para.text.strip()
            if not text:
                continue
            
            # 检测大标题
            if '登记资料' in text and len(text) < 20:
                cover_info['title'] = {
                    'text': text,
                    'alignment': para.alignment,
                    'font_size': para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else 16
                }
            
            # 检测信息字段
            elif any(keyword in text for keyword in ['产品名称', '登记号', '受理号', '申请人']):
                cover_info['fields'].append({
                    'text': text,
                    'alignment': para.alignment
                })
        
        self.format_specs['cover'] = cover_info
        return cover_info
    
    def get_all_specs(self):
        """获取所有格式规范"""
        self.parse_header()
        self.parse_fonts()
        self.parse_paragraph_format()
        self.parse_headings()
        self.parse_cover_page()
        
        return self.format_specs
    
    def apply_to_document(self, new_doc, drug_name=""):
        """将解析的格式应用到新文档"""
        specs = self.format_specs
        
        # 应用段落格式
        if 'paragraphs' in specs and specs['paragraphs']:
            para_fmt = specs['paragraphs']
            for para in new_doc.paragraphs:
                if para_fmt.get('line_spacing'):
                    para.paragraph_format.line_spacing = para_fmt['line_spacing']
                if para_fmt.get('space_before'):
                    para.paragraph_format.space_before = Pt(para_fmt['space_before'])
                if para_fmt.get('space_after'):
                    para.paragraph_format.space_after = Pt(para_fmt['space_after'])
                if para_fmt.get('first_line_indent'):
                    para.paragraph_format.first_line_indent = para_fmt['first_line_indent']


def extract_template_format(template_path):
    """便捷函数：提取模板格式"""
    parser = TemplateParser(template_path)
    return parser.get_all_specs()


if __name__ == '__main__':
    import json
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python3 template_parser.py <template_word_path>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    specs = extract_template_format(template_path)
    
    print(json.dumps(specs, indent=2, default=str))
