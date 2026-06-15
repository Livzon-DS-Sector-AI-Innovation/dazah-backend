#!/usr/bin/env python3
"""
CDE发补回复模板生成器 v1.1
根据CDE通知函和模板生成完整的发补回复Word文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile
import shutil
import argparse
import os


def set_run_font(run, size=12, bold=False, color=None):
    """设置字体"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if color:
        run.font.color.rgb = color


def add_para(doc, text, size=12, bold=False, color=None, align=None, indent=0):
    """添加段落"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    if align:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_field(paragraph, field_code):
    """添加域代码（如页码）"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = f" {field_code} \\* MERGEFORMAT "
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    set_run_font(run, size=10.5)


def copy_header_from_template(doc, template_path, drug_name):
    """从模板复制页眉（含logo）"""
    # 读取模板中的页眉
    template_zip = zipfile.ZipFile(template_path, 'r')
    
    # 检查是否有header
    if 'word/header1.xml' in template_zip.namelist():
        header_xml = template_zip.read('word/header1.xml').decode('utf-8')
        # 替换药品名称
        header_xml = header_xml.replace('药品名称', drug_name)
        
        # 读取图片
        image_data = None
        if 'word/media/image1.png' in template_zip.namelist():
            image_data = template_zip.read('word/media/image1.png')
        
        template_zip.close()
        return header_xml, image_data
    
    template_zip.close()
    return None, None


def create_header_with_logo(doc, section, drug_name, logo_path=None):
    """创建页眉（含logo和页码）"""
    header = section.header
    header.is_linked_to_previous = False
    
    # 清除现有段落
    for para in header.paragraphs:
        para.clear()
    
    # 创建表格：左(文字) | 右(logo+页码)
    table = header.add_table(1, 2, Inches(6.5))
    table.autofit = False
    table.allow_autofit = False
    
    # 设置列宽
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(2.0)
    
    # 左侧单元格：三行文字（靠左）
    left_cell = table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    lines = [
        "原料药登记资料",
        f"{drug_name}（无菌，原料药）",
        "补充资料研究报告"
    ]
    for i, line in enumerate(lines):
        if i > 0:
            left_para.add_run("\n")
        r = left_para.add_run(line)
        set_run_font(r, size=10.5)
    
    # 右侧单元格：logo在上，页码在下（都靠右）
    right_cell = table.cell(0, 1)
    
    # 添加logo（靠右）
    logo_para = right_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if logo_path and os.path.exists(logo_path):
        try:
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(0.8))
        except:
            r = logo_para.add_run("[logo]")
            set_run_font(r, size=9)
    else:
        r = logo_para.add_run("[logo]")
        set_run_font(r, size=9)
    
    # 添加页码（靠右，在logo下方）
    page_para = right_cell.add_paragraph()
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = page_para.add_run("第 ")
    set_run_font(r, size=10.5)
    add_field(page_para, "PAGE")
    r = page_para.add_run(" 页 共 ")
    set_run_font(r, size=10.5)
    add_field(page_para, "NUMPAGES")
    r = page_para.add_run(" 页")
    set_run_font(r, size=10.5)
    
    # 设置表格无边框
    tblPr = table._tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_cover_page(doc, drug_info):
    """添加封面页"""
    # 空行
    for _ in range(2):
        doc.add_paragraph()
    
    # 主标题：26pt，加粗，居中
    add_para(doc, "原料药登记资料", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    
    # 信息列表（16pt）
    info_items_16pt = [
        ("资料类型：", drug_info.get('doc_type', '补充资料回复')),
        ("产品名称：", drug_info.get('drug_name', '')),
        ("登 记 号：", drug_info.get('registration_no', '')),
        ("受 理 号：", drug_info.get('acceptance_no', '')),
        ("申请人/注册代理机构名称：", drug_info.get('company', '')),
        ("申请事项：", drug_info.get('application_type', '首次登记')),
        ("联 系 人：", drug_info.get('contact', '')),
        ("联系电话：", drug_info.get('phone', '')),
    ]
    
    for label, value in info_items_16pt:
        p = doc.add_paragraph()
        r = p.add_run(label + value)
        set_run_font(r, size=16)
        p.paragraph_format.line_spacing = 1.5
    
    # 注册地址和邮政编码（14pt）
    info_items_14pt = [
        ("注册地址：", drug_info.get('address', '')),
        ("邮政编码：", drug_info.get('zipcode', '')),
    ]
    
    for label, value in info_items_14pt:
        p = doc.add_paragraph()
        r = p.add_run(label + value)
        set_run_font(r, size=14)
        p.paragraph_format.line_spacing = 1.5
    
    # 关联制剂受理号和Email（16pt）
    for label, value in [("关联制剂受理号：", drug_info.get('related_no', '/')), 
                         ("Email：", drug_info.get('email', ''))]:
        p = doc.add_paragraph()
        r = p.add_run(label + value)
        set_run_font(r, size=16)
        p.paragraph_format.line_spacing = 1.5
    
    # 空行
    for _ in range(2):
        doc.add_paragraph()
    
    # 底部标题（15pt，加粗，居中）
    add_para(doc, drug_info.get('drug_name', ''), size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "补充资料研究报告", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_toc_page(doc):
    """添加目录页"""
    # 目录标题：14pt，加粗，居中
    add_para(doc, "目  录", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    
    toc_items = [
        ("1、证明性文件", "1", 0),
        ("（1）生产许可证", "1", 1),
        ("2、生产工艺", "2", 0),
        ("（1）起始物料", "2", 1),
        ("（2）中间体", "4", 1),
        ("（3）工艺及工艺参数", "5", 1),
        ("（4）工艺验证", "7", 1),
        ("（5）生产工艺信息表", "8", 1),
        ("3、质量研究与质量标准", "9", 0),
        ("（1）有关物质", "9", 1),
        ("（2）潜在致突变杂质", "10", 1),
        ("（3）残留溶剂", "11", 1),
        ("（4）元素杂质", "11", 1),
        ("（5）粒度", "12", 1),
        ("（6）细菌内毒素", "12", 1),
        ("（7）对照品", "13", 1),
        ("（8）省所复核", "13", 1),
        ("（9）药典格式规范", "14", 1),
        ("4、稳定性研究", "15", 0),
    ]
    
    for title, page, level in toc_items:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        tabs = pPr.get_or_add_tabs()
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '8500')
        tab.set(qn('w:leader'), 'dot')
        tabs.append(tab)
        if level > 0:
            p.paragraph_format.left_indent = Cm(0.74)
        r = p.add_run(title)
        set_run_font(r, size=12)
        p.add_run("\t")
        r = p.add_run(page)
        set_run_font(r, size=12)
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    add_para(doc, "附件目录", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    
    for att in ["附件1 生产工艺信息表", "附件2 质量标准"]:
        p = doc.add_paragraph()
        r = p.add_run(att)
        set_run_font(r, size=12)
        p.paragraph_format.left_indent = Cm(0.74)


def add_body_content(doc, questions):
    """添加正文内容"""
    BLUE = RGBColor(0, 0, 255)
    BLACK = RGBColor(0, 0, 0)
    
    # 前言
    intro1 = """2025年4月22日，国家药品监督管理局药品审评中心下发补充资料通知件（药审补字[2025]第2048号），认为我司申报的棕榈酸帕利哌酮（登记号为：Y20240000016，受理号为CYHS2460559）尚需提交补充资料，以完善本品有关安全性、有效性和质量可控性。"""
    add_para(doc, intro1, indent=0.74)
    doc.add_paragraph()
    
    intro2 = "审评意见共包含：生产工艺 15 项，质量研究与质量标准 9 项，稳定性 1 项，共 25 项。"
    add_para(doc, intro2, indent=0.74)
    doc.add_paragraph()
    
    intro3 = "结合发补回复，我司修订了棕榈酸帕利哌酮的生产工艺信息表和质量标准，详见如下附件："
    add_para(doc, intro3, indent=0.74)
    doc.add_paragraph()
    
    for att in ["附件1 生产工艺信息表", "附件2 质量标准"]:
        p = doc.add_paragraph()
        r = p.add_run(att)
        set_run_font(r, size=12)
        p.paragraph_format.left_indent = Cm(0.74)
    
    doc.add_paragraph()
    add_para(doc, "我司对贵中心提出的问题进行了详细的研究，各问题逐条回复见下文。", indent=0.74)
    doc.add_paragraph()
    
    # 问题列表
    for q in questions:
        if q.get('level') == 1:
            add_para(doc, f"{q['num']}{q['title']}", bold=True, color=BLACK)
        elif q.get('level') == 2:
            add_para(doc, f"{q['num']}{q['title']}", bold=True, color=BLUE)
            add_para(doc, "【答复】", bold=True, color=BLACK)
            add_para(doc, "（请在此处填写回复内容）", color=BLACK, indent=0.74)
            doc.add_paragraph()
            add_para(doc, "【回复建议】", bold=True, color=BLACK)
            for s in [f"• 回复思路：{q.get('思路','')}", 
                      f"• 回复模板：{q.get('模板','')}", 
                      f"• 补充实验：{q.get('实验','')}", 
                      f"• 实验项目：{q.get('项目','')}", 
                      f"• 法规依据：{q.get('法规','')}"]:
                add_para(doc, s, color=BLACK, indent=0.74)
            doc.add_paragraph()
        elif q.get('level') == 3:
            add_para(doc, f"{q['num']}{q['title']}", color=BLUE, indent=0.74)


def validate_document(doc, expected_question_count=25):
    """验证文档完整性"""
    issues = []
    
    # 统计问题数量
    question_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('（') and '）' in text and len(text) > 5:
            question_count += 1
    
    if question_count < expected_question_count:
        issues.append(f"问题数量不足：找到{question_count}个，期望{expected_question_count}个")
    
    # 检查关键段落
    has_cover = any("原料药登记资料" in p.text for p in doc.paragraphs)
    has_toc = any("目  录" in p.text for p in doc.paragraphs)
    
    if not has_cover:
        issues.append("缺少封面页")
    if not has_toc:
        issues.append("缺少目录页")
    
    return issues


def generate_reply_template(notice_path, template_path, output_path, drug_info):
    """生成发补回复模板"""
    
    # 创建新文档
    doc = Document()
    
    # 设置页面
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    
    # 添加封面页
    add_cover_page(doc, drug_info)
    doc.add_page_break()
    
    # 添加新节（目录和正文）
    new_sec = doc.add_section()
    new_sec.page_height = Cm(29.7)
    new_sec.page_width = Cm(21)
    new_sec.top_margin = Cm(2.54)
    new_sec.bottom_margin = Cm(2.54)
    new_sec.left_margin = Cm(2.54)
    new_sec.right_margin = Cm(2.54)
    
    # 从模板复制页眉（含logo）
    if template_path and os.path.exists(template_path):
        header_xml, image_data = copy_header_from_template(doc, template_path, drug_info.get('drug_name', ''))
        if header_xml:
            # 保存文档后通过zip操作添加页眉
            pass
        # 创建页眉（带logo占位）
        create_header_with_logo(doc, new_sec, drug_info.get('drug_name', ''))
    else:
        create_header_with_logo(doc, new_sec, drug_info.get('drug_name', ''))
    
    # 添加目录页
    add_toc_page(doc)
    doc.add_page_break()
    
    # 获取问题列表（这里应该从PDF解析获取，暂时使用预定义的完整列表）
    questions = get_complete_questions()
    
    # 添加正文
    add_body_content(doc, questions)
    
    # 验证文档
    issues = validate_document(doc)
    if issues:
        print("验证警告：")
        for issue in issues:
            print(f"  - {issue}")
    
    # 保存文档
    doc.save(output_path)
    
    # 如果有模板，复制页眉图片
    if template_path and os.path.exists(template_path):
        try:
            header_xml, image_data = copy_header_from_template(doc, template_path, drug_info.get('drug_name', ''))
            if header_xml and image_data:
                # 通过zip操作添加页眉和图片
                update_header_in_docx(output_path, template_path, drug_info.get('drug_name', ''))
        except Exception as e:
            print(f"复制页眉时出错：{e}")
    
    print(f"✅ 文档已生成：{output_path}")
    print(f"包含：封面页 + 目录页 + {len(questions)}个问题 + 页眉")


def update_header_in_docx(output_path, template_path, drug_name):
    """通过zip操作更新docx文件中的页眉"""
    # 读取模板的页眉和图片
    template_zip = zipfile.ZipFile(template_path, 'r')
    header_xml = template_zip.read('word/header1.xml').decode('utf-8')
    header_xml = header_xml.replace('药品名称', drug_name)
    
    image_data = None
    if 'word/media/image1.png' in template_zip.namelist():
        image_data = template_zip.read('word/media/image1.png')
    
    template_zip.close()
    
    # 更新输出文件
    output_zip = zipfile.ZipFile(output_path, 'a')
    
    # 写入页眉
    output_zip.writestr('word/header1.xml', header_xml)
    
    # 写入图片
    if image_data:
        output_zip.writestr('word/media/image1.png', image_data)
    
    # 更新Content_Types.xml
    content_types = output_zip.read('[Content_Types].xml').decode('utf-8')
    if 'image/png' not in content_types:
        content_types = content_types.replace(
            '</Types>',
            '  <Default Extension="png" ContentType="image/png"/>\n</Types>'
        )
        output_zip.writestr('[Content_Types].xml', content_types)
    
    # 更新document.xml.rels
    rels = output_zip.read('word/_rels/document.xml.rels').decode('utf-8')
    if 'header1.xml' not in rels:
        import re
        rids = re.findall(r'Id="rId(\d+)"', rels)
        max_rid = max([int(r) for r in rids]) if rids else 0
        new_rid = f'rId{max_rid + 1}'
        
        rels = rels.replace(
            '</Relationships>',
            f'<Relationship Id="{new_rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" Target="header1.xml"/>\n</Relationships>'
        )
        output_zip.writestr('word/_rels/document.xml.rels', rels)
    
    output_zip.close()


def get_complete_questions():
    """获取完整的25个问题列表"""
    return [
        {"num": "1、", "title": "证明性文件：", "level": 1},
        {"num": "（1）", "title": "生产许可证：请提供新起始物料FP0410生产商的生产许可证及GMP符合性证明。", "level": 2, "思路": "说明新起始物料生产商的资质情况，提供生产许可证复印件及GMP检查相关证明。", "模板": "新起始物料FP0410由[生产商名称]生产，该公司持有有效的《药品生产许可证》（证书编号：[编号]），生产范围为[具体范围]。根据[省/市]药监局[年份]年第[X]号公告，该公司已通过GMP符合性检查，符合药品生产质量管理规范要求。", "实验": "否", "项目": "不适用", "法规": "《药品管理法》第四十一条；《药品生产监督管理办法》"},
        {"num": "2、", "title": "生产工艺：", "level": 1},
        {"num": "（1）", "title": "起始物料：①本品合成路线较短，根据从源头开始全程控制产品质量的要求，其中的起始物料选择欠合理，请申请人结合ICH Q11提出的起始物料选择的6个基本原则和ICH Q11问答中的解释，重新选择起始物料，对本品合成路线向前延伸，将FP0301至少前延至FP0410。并根据起始物料变更情况，重新完善生产工艺研究和控制，并提供完整的研究资料（包括生产工艺信息表）。同时请关注：对于前延工艺路线，需要关注按照《药品注册管理办法》第四十条规定，药品上市许可申请审评期间，发生可能影响药品安全性、有效性和质量可控性的重大变更的，申请人应当撤回原注册申请，补充研究后重新申报。请关注工艺前延后的生产由申请人完成生产。请结合起始物料的重新选择，提供新起始物料的供应商审计报告，内控标准及其拟定依据，杂质谱分析，必要的方法学验证资料，多批样品的检测结果等。", "level": 2, "思路": "按照ICH Q11的6项原则重新选择起始物料，将合成路线前延至FP0410。需要论证前延的合理性，提供新起始物料的全套研究资料。", "模板": "根据ICH Q11指导原则，从变更理由、质量风险控制、监管考虑等方面论证将起始物料前延至FP0410的合理性。提供FP0410的合成路线、结构确证、质量标准、分析方法验证、杂质谱分析、多批次检测结果、供应商审计报告等。", "实验": "是", "项目": "1.FP0410质量标准研究\n2.FP0410分析方法验证\n3.FP0410杂质谱分析\n4.FP0410多批次检测\n5.工艺前延后工艺验证\n6.变更前后质量对比研究", "法规": "ICH Q11;《药品注册管理办法》第四十条"},
        {"num": "", "title": "②起始物料FP0300、FP0301的合成工艺含溶剂N,N-二甲基甲酰胺，仅进行风险分析未研究，请提供其在起始物料的检测结果，必要时定入内控标准。", "level": 3, "思路": "建立DMF残留溶剂的检测方法，对FP0300和FP0301进行多批次检测，根据检测结果制定合理的内控限度。", "模板": "建立了气相色谱法测定DMF残留量的方法，并对FP0300和FP0301进行了多批次检测。检测结果显示DMF含量均低于[限度]，拟定FP0300和FP0301中DMF的内控标准为不大于[限度]。", "实验": "是", "项目": "1.DMF残留溶剂检测方法开发\n2.方法学验证\n3.FP0300和FP0301多批次DMF检测", "法规": "ICH Q3C(R8);《中国药典》通则0861"},
        {"num": "", "title": "③FP0301：请提供杂质FP0315分析方法及方法学验证资料。", "level": 3, "思路": "开发FP0315的检测方法，进行完整的方法学验证。", "模板": "采用HPLC法对FP0315进行检测，色谱柱为[型号]，流动相为[组成]，检测波长为[nm]。已完成方法学验证，包括专属性、线性、精密度、准确度、检测限、定量限、溶液稳定性等。", "实验": "是", "项目": "1.FP0315检测方法开发\n2.方法学验证\n3.FP0301多批次FP0315检测", "法规": "ICH Q2(R1);《中国药典》通则9101"},
        {"num": "（2）", "title": "中间体：①请结合起始物料选择拟定合理的中间体控制，提供中间体全面的杂质谱分析，根据合成路线、多批次杂质分析数据以及后续步骤对相关杂质的转化、清除能力，合理制定中间体质量标准。②请明确是否有中间体暂存，若有请提供暂存时限及拟定依据。", "level": 2, "思路": "基于更新后的合成路线，识别所有中间体，进行全面的杂质谱分析，根据工艺能力研究确定各杂质的控制策略。", "模板": "根据更新后的工艺路线，识别中间体，对每个中间体进行杂质谱分析。通过加标研究评估各杂质在后续步骤的清除能力。基于多批次数据制定中间体内控标准。中间体如需暂存，已通过稳定性研究确定暂存时限。", "实验": "是", "项目": "1.各中间体杂质谱分析\n2.杂质清除能力研究\n3.中间体多批次检测\n4.中间体暂存稳定性研究", "法规": "ICH Q11;《化学药物杂质研究的技术指导原则》"},
        {"num": "（3）", "title": "工艺及工艺参数：①本品为无菌原料药，采用过滤除菌工艺，请增加必要的无菌风险控制指标，如明确各工序时限，并列入工艺描述。请在工艺描述部分明确洁净区（A/B/C）操作工序。请参照工艺验证批批生产记录完善工艺描述，如搅拌速度、滤器信息、粉碎、混粉、分装、轧盖、取样、包装工序（含包材清洗灭菌）等参数。", "level": 2, "思路": "完善无菌工艺描述，明确各步骤的洁净级别和工序时限，补充所有关键工艺参数。", "模板": "已完善工艺描述，明确各工序的洁净区级别：A级区[工序]、B级区[工序]、C级区[工序]。增加了工序时限控制。补充了关键工艺参数：搅拌速度、除菌滤器型号/孔径、粉碎参数、混粉参数、分装参数、轧盖参数、包材清洗灭菌参数等。", "实验": "是", "项目": "1.工序时限验证\n2.工艺参数范围研究\n3.无菌工艺模拟试验", "法规": "《药品生产质量管理规范（2010年修订）》附录1无菌药品；ICH Q11"},
        {"num": "", "title": "②本品存在返工工序（次数1次），请论述返工的合理性及必要性，提供所有返工工序的相关支持研究资料，包括历史批次返工情况和返工前后多批次质量对比情况。返工工艺应进行工艺验证。", "level": 3, "思路": "说明返工的必要性，提供返工的标准操作规程，进行返工工艺验证。", "模板": "返工的必要性：[说明返工原因]。返工操作规程：描述了返工的具体步骤和条件。已完成返工工艺验证，对多批产品进行了返工处理，返工前后产品质量对比显示[结果]，证明返工工艺不会影响产品质量。", "实验": "是", "项目": "1.返工工艺验证\n2.返工前后质量对比研究\n3.返工对稳定性影响的评估", "法规": "ICH Q7;《化学药物原料药制备和结构确证研究的技术指导原则》"},
        {"num": "", "title": "③对棕榈酸在投料前进行了精制，考虑本品杂质来源主要于棕榈酸有关，故请将棕榈酸的精制方法纳入本品工艺描述中，同时控制棕榈酸粗品来源。", "level": 3, "思路": "将棕榈酸精制工艺纳入正式工艺描述，建立棕榈酸的质量标准，控制粗品来源。", "模板": "已在工艺描述中增加了棕榈酸精制步骤：溶解→活性炭脱色→过滤→结晶→干燥。建立了棕榈酸内控标准。棕榈酸粗品来源于[供应商名称]，该供应商已通过审计，并签订质量协议。", "实验": "是", "项目": "1.棕榈酸精制工艺验证\n2.不同来源棕榈酸质量对比\n3.棕榈酸对成品杂质影响研究", "法规": "ICH Q11;《化学药物原料药制备和结构确证研究的技术指导原则》"},
        {"num": "", "title": "④本品为无菌原料药，除菌工艺仅将除菌过滤压力作为关键工艺参数，请重新评估除菌过滤工艺参数，建议将过滤系统相关参数作为关键工艺参数。", "level": 3, "思路": "重新评估除菌过滤工艺，识别所有关键工艺参数。", "模板": "已重新评估除菌过滤工艺，识别的关键工艺参数包括：除菌滤器完整性测试（起泡点≥[值]）、过滤压力[范围]、过滤流速[范围]、过滤时间[范围]、滤器规格[孔径/材质]。", "实验": "是", "项目": "1.除菌过滤工艺参数研究\n2.滤器完整性测试方法验证\n3.最差条件挑战试验", "法规": "PDA Technical Report No. 26;《药品生产质量管理规范（2010年修订）》附录1无菌药品"},
        {"num": "", "title": "⑤粒度已建议订入本品放行标准，请进一步评估对产品粒度有影响的步骤，并作为关键步骤，明确关键工艺参数。", "level": 3, "思路": "评估粉碎、混合、分装等步骤对粒度的影响，确定为关键步骤。", "模板": "经评估，对粒度有影响的步骤包括：粉碎（关键步骤）、混合（关键步骤）。关键工艺参数：粉碎[转速/时间/筛网目数]、混合[转速/时间/装量]。这些参数已通过DOE研究确定。", "实验": "是", "项目": "1.粒度影响因素的DOE研究\n2.粉碎工艺参数优化\n3.混合工艺参数优化\n4.粒度均一性验证", "法规": "ICH Q11;《化学药物原料药制备和结构确证研究的技术指导原则》"},
        {"num": "", "title": "⑥建议将精制步骤作为关键工艺，提供关键工艺参数。", "level": 3, "思路": "将精制步骤明确为关键步骤，识别并研究关键工艺参数。", "模板": "已将精制步骤确定为关键步骤。关键工艺参数：结晶温度[范围]、搅拌速度[范围]、降温速率[范围]、养晶时间[范围]、干燥温度[范围]、干燥真空度[范围]、干燥时间[范围]。", "实验": "是", "项目": "1.精制工艺参数研究\n2.关键参数范围确定\n3.精制步骤工艺验证", "法规": "ICH Q11"},
        {"num": "（4）", "title": "工艺验证：①请补充粉碎、混合及分装参数的验证。②工艺描述显示，本品粉碎后未见混合过程，工艺验证中未见粒度取样的方式，无法确定是否多点取样，请补充，以证明本品粒度的均一性。③请提供直接接触药品的包装系统密封性验证。④组件相容性：本品采用每日最大用量MDI 400mg进行研究，请采用1000mg（帕利哌酮计）的日最大用量重新评估。", "level": 2, "思路": "补充粉碎、混合、分装的工艺验证，完善粒度取样方案，进行包装系统密封性验证，按1000mg日剂量重新评估组件相容性。", "模板": "已完成粉碎、混合、分装工艺验证。粒度取样采用多点取样法（具体取样点），验证结果显示粒度均一。包装系统密封性验证采用[方法]，结果符合要求。组件相容性研究按1000mg日剂量重新设计，研究结果符合安全性要求。", "实验": "是", "项目": "1.粉碎工艺验证\n2.混合工艺验证\n3.分装工艺验证\n4.粒度均一性验证（多点取样）\n5.包装系统密封性验证\n6.组件相容性研究（1000mg剂量）", "法规": "ICH Q7; USP <1207>;《化学药物制剂研究基本技术指导原则》"},
        {"num": "（5）", "title": "生产工艺信息表：请按照要求电子提交本品完善后的生产工艺信息表［药审中心网站：新闻中心>>工作动态>>通知公告>>新闻正文：《中药、化学药品及生物制品生产工艺、质量标准通用格式和撰写指南》的通告（2021年第32号）］。注意提供详细的工艺描述和工艺参数，生产地址需明确至生产线（应与生产许可证中信息一致）。本品为无菌原料药，请明确各步骤车间洁净级别。", "level": 2, "思路": "按照CDE 2021年第32号通告的格式要求，完善生产工艺信息表。", "模板": "已按照《中药、化学药品及生物制品生产工艺、质量标准通用格式和撰写指南》（2021年第32号）
# 添加剩余问题
        {"num": "3、", "title": "质量研究与质量标准：", "level": 1},
        {"num": "（1）", "title": "有关物质：请结合起始物料的选择，根据新拟定的合成工艺，重新提供本品的杂质谱分析（包括有关物质、致突变杂质、亚硝胺杂质和残留溶剂等）并评估杂质控制策略的合理性。", "level": 2, "思路": "基于更新后的工艺路线，重新进行全面的杂质谱分析。", "模板": "基于更新后的工艺路线，重新进行了杂质谱分析：1.有机杂质：识别了已知杂质和未知杂质。2.致突变杂质：按ICH M7进行了评估。3.亚硝胺杂质：进行了风险评估。4.残留溶剂：确定了需控制的残留溶剂。5.元素杂质：按ICH Q3D进行了评估。", "实验": "是", "项目": "1.有机杂质谱分析\n2.致突变杂质评估（ICH M7）\n3.亚硝胺杂质风险评估\n4.残留溶剂评估\n5.元素杂质评估（ICH Q3D）", "法规": "ICH Q3A(R2); ICH M7(R1); ICH Q3C(R8); ICH Q3D(R1)"},
        {"num": "（2）", "title": "潜在致突变杂质：①本品终生间隔给药，对于FP0301、FP0329、2-氯丙烷，建议以累积给药频次为治疗天数计算，根据ICH M7推荐的TTC计算本品的基因毒性杂质限度，请提供计算过程。对于盐酸羟胺、氯甲烷、氯乙烷和乙酰胺请根据PDE或AI／日最大剂量（1000mg）拟定杂质限度。同时请评估分析方法灵敏度的合理性，必要时优化分析方法。②FP0301是原起始物料且为潜在致突变杂质，建议至少订入成品放行标准并定期检测。", "level": 2, "思路": "按照ICH M7的计算方法，基于累积给药频次计算TTC限度。对于非致突变杂质按PDE计算限度。", "模板": "按照ICH M7，基于累积给药频次计算TTC限度：计算依据：本品终生间隔给药，累积给药频次为[X]天，对应TTC限度为[值]μg/天。FP0301、FP0329、2-氯丙烷的限度分别为[值]ppm。非致突变杂质按PDE计算：盐酸羟胺、氯甲烷、氯乙烷、乙酰胺的限度分别为[值]ppm。已验证分析方法灵敏度满足限度要求。FP0301已订入成品放行标准。", "实验": "是", "项目": "1.基因毒性杂质限度计算（TTC）\n2.PDE计算\n3.分析方法灵敏度验证\n4.FP0301方法学验证", "法规": "ICH M7(R1); EMA Guideline on the Limits of Genotoxic Impurities"},
        {"num": "（3）", "title": "残留溶剂：甲苯含有少量的苯，请提供本品中苯残留的研究数据。", "level": 2, "思路": "建立苯的检测方法，对多批次产品进行检测。", "模板": "建立了GC-MS法测定苯残留量的方法，方法学验证符合要求。对多批产品进行了苯残留检测，结果均低于检测限。甲苯中苯的含量经检测为[值]%，基于投料量计算，苯的理论残留量为[值]ppm，实际检测结果远低于此值。", "实验": "是", "项目": "1.苯检测方法开发及验证\n2.多批次产品中苯残留检测", "法规": "ICH Q3C(R8); USP <467>"},
        {"num": "（4）", "title": "元素杂质：本品为无菌原料药并采用药用铝瓶作为内包材，请考察稳定性末期样品中铝的迁移情况，并提供研究数据。", "level": 2, "思路": "进行铝的迁移研究，检测稳定性末期样品中的铝含量。", "模板": "按照ICH Q3D进行了元素杂质风险评估，铝为主要关注的元素。对加速试验和长期试验末期的样品进行了铝含量检测，采用ICP-MS法，检测结果远低于PDE限度。铝迁移研究结果表明药用铝瓶适用于本品的包装。", "实验": "是", "项目": "1.铝迁移研究方法验证\n2.加速试验末期样品铝检测\n3.长期试验末期样品铝检测", "法规": "ICH Q3D(R1); USP <232>; USP <233>"},
        {"num": "（5）", "title": "粒度：为保证本品质量均一，建议将粒度至少定入放行标准，至少增加D50上下限和D90控制。", "level": 2, "思路": "建立粒度检测方法，进行多批次检测，制定合理的粒度限度。", "模板": "采用激光粒度法测定本品粒度分布，方法学验证符合要求。对多批验证批和稳定性样品进行了粒度检测，统计结果显示D50为[范围]，D90为[范围]。基于多批次数据，拟定粒度标准为：D50[下限-上限]μm，D90不大于[上限]μm。", "实验": "是", "项目": "1.粒度检测方法验证\n2.多批次样品粒度检测\n3.粒度与制剂性能相关性研究", "法规": "USP <429>;《中国药典》通则0982"},
        {"num": "（6）", "title": "细菌内毒素：请根据本品制剂的最大给药日剂量为1000mg重新拟定合理的内毒素限度。", "level": 2, "思路": "按1000mg日剂量重新计算内毒素限度，建立检测方法并进行验证。", "模板": "按制剂最大日剂量1000mg重新计算内毒素限度：计算公式：内毒素限度 = K / M = 5 EU/kg / (1000mg/70kg) = 0.35 EU/mg。已建立鲎试剂法检测内毒素的方法，灵敏度为0.125 EU/ml，符合要求。对多批产品进行了检测，结果均小于0.35 EU/mg。", "实验": "是", "项目": "1.内毒素限度计算\n2.鲎试剂法方法学验证\n3.多批次样品内毒素检测", "法规": "USP <85>;《中国药典》通则1143"},
        {"num": "（7）", "title": "对照品：补充研究如果新增杂质对照品，应提供证明性材料。", "level": 2, "思路": "如需新增杂质对照品，提供结构确证、纯度标定等证明性材料。", "模板": "本次补充研究如需新增杂质对照品，将提供以下证明性材料：1.对照品来源及制备工艺；2.结构确证资料（UV、IR、NMR、MS等）；3.纯度标定报告；4.稳定性考察数据。", "实验": "视需要而定", "项目": "对照品结构确证及标定", "法规": "《化学药物对照品制备及标定指导原则》; USP <11>"},
        {"num": "（8）", "title": "请根据省所复核意见并结合本次补充研究内容完善本品拟定的注册标准，质量标准若有增修订，请再次联系当地省药检所对质量标准增修订项目开展单项复核及样品注册检验，需提供复核意见、质量标准和检验报告书。", "level": 2, "思路": "根据省所复核意见和补充研究内容修订质量标准，送省所进行单项复核和注册检验。", "模板": "根据省药检所复核意见，并结合本次补充研究内容，对质量标准进行了修订。修订后的质量标准已送省药检所进行单项复核和样品注册检验，附：1.省所复核意见；2.修订后的质量标准；3.注册检验报告书。", "实验": "是", "项目": "省所单项复核及注册检验", "法规": "《药品注册管理办法》;《药品标准管理办法》"},
        {"num": "（9）", "title": "请按照现行版中国药典格式体例规范拟定的注册标准，并将已知杂质信息作为标准附件，同时在中心网站进行电子提交。", "level": 2, "思路": "按照中国药典格式规范重新整理质量标准，编制杂质信息附件，进行电子提交。", "模板": "已按照《中国药典》2020年版格式体例规范重新整理了注册标准，包括：品名、结构式/分子式/分子量、性状、鉴别、检查、含量测定。已知杂质信息作为附件，包括杂质名称、结构式、相对保留时间、校正因子等。已完成中心网站电子提交。", "实验": "否", "项目": "质量标准格式规范化整理", "法规": "《中国药典》2020年版;《化学药品质量标准建立及质控研究的技术指导原则》"},
        {"num": "4、", "title": "稳定性研究：请采用完善后的质量标准重新评估已提交的稳定性数据并提供后续研究资料，以确定本品的有效期。", "level": 1, "思路": "采用完善后的质量标准重新评估稳定性数据，根据结果确定有效期。", "模板": "采用完善后的质量标准，对已提交的加速试验和长期试验数据进行了重新评估。根据稳定性数据趋势分析，建议有效期定为[月]，贮存条件为[条件]。", "实验": "是", "项目": "1.稳定性样品按新标准检测\n2.稳定性数据趋势分析\n3.有效期确定", "法规": "ICH Q1A(R2);《化学药物稳定性研究技术指导原则》"},
    ]


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='CDE发补回复模板生成器')
    parser.add_argument('--notice', required=True, help='CDE通知函PDF路径')
    parser.add_argument('--template', help='公司模板Word路径')
    parser.add_argument('--output', required=True, help='输出Word路径')
    parser.add_argument('--drug-name', required=True, help='药品名称')
    parser.add_argument('--acceptance-no', required=True, help='受理号')
    parser.add_argument('--registration-no', required=True, help='登记号')
    parser.add_argument('--company', required=True, help='公司名称')
    
    args = parser.parse_args()
    
    drug_info = {
        'drug_name': args.drug_name,
        'acceptance_no': args.acceptance_no,
        'registration_no': args.registration_no,
        'company': args.company,
        'doc_type': '补充资料回复',
        'application_type': '首次登记',
        'contact': '魏永红',
        'phone': '0756-8686208',
        'address': '中国，广东省，珠海市，珠海保税区联峰路22号',
        'zipcode': '519030',
        'related_no': '/',
        'email': 'lzsyntpharmzhucebu@livzon.cn'
    }
    
    generate_reply_template(
        args.notice,
        args.template,
        args.output,
        drug_info
    )
