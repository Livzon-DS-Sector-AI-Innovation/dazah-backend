#!/usr/bin/env python3
"""
CDE发补回复模板生成器 v2.0
根据CDE通知函和模板生成完整的发补回复Word文档
支持：自动解析CDE PDF通知函，提取所有问题和药品信息
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import zipfile
import shutil
import argparse
import os
import re
import json


def set_run_font(run, size=12, bold=False, color=None):
    """设置字体"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if color:
        run.font.color.rgb = color


def add_para(doc, text, size=12, bold=False, color=None, align=None, indent=0):
    """添加段落"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    if align:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def extract_from_pdf(notice_path):
    """从PDF提取药品信息和问题列表"""
    # 这里应该使用PDF解析，现在先用pdfplumber示例
    try:
        import pdfplumber
        with pdfplumber.open(notice_path) as pdf:
            full_text = "\n".join([page.extract_text() or "" for page in pdf.pages])
    except:
        # 如果pdfplumber不可用，返回空数据
        full_text = ""
    
    # 提取药品信息
    drug_info = {}
    
    # 提取药品名称
    name_match = re.search(r'(?:产品名称|药品名称|原料药名称)[：:]\s*([^\n]+)', full_text)
    if name_match:
        drug_info['drug_name'] = name_match.group(1).strip()
    
    # 提取登记号
    reg_match = re.search(r'(?:登记号|登记编号)[：:]\s*([A-Z]?\d+)', full_text)
    if reg_match:
        drug_info['registration_no'] = reg_match.group(1).strip()
    
    # 提取受理号
    acc_match = re.search(r'(?:受理号|受理编号)[：:]\s*(CY[A-Z]*\d+)', full_text)
    if acc_match:
        drug_info['acceptance_no'] = acc_match.group(1).strip()
    
    # 提取公司名称
    company_match = re.search(r'(?:申请人|公司名称)[：:]\s*([^；;]+)', full_text)
    if company_match:
        drug_info['company'] = company_match.group(1).strip()
    
    # 提取CDE通知日期和编号
    notice_match = re.search(r'(\d{4}年\d{1,2}月\d{1,2}日).*?药审补字\[(\d{4})\]第(\d+)号', full_text)
    if notice_match:
        drug_info['notice_date'] = notice_match.group(1)
        drug_info['notice_year'] = notice_match.group(2)
        drug_info['notice_no'] = notice_match.group(3)
    
    # 提取问题列表（简化版）
    questions = []
    # 按行分割，识别问题模式
    lines = full_text.split('\n')
    current_section = None
    current_subsection = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 一级标题：1、2、3...
        if re.match(r'^\d+[、.]\s*\S', line):
            current_section = line
            questions.append({
                'num': re.match(r'^(\d+[、.])', line).group(1),
                'title': line,
                'level': 1,
                'content': ''
            })
        # 二级标题：（1）（2）（3）...
        elif re.match(r'^（\d+）', line):
            current_subsection = line
            questions.append({
                'num': re.match(r'^（(\d+)）', line).group(0),
                'title': line[3:] if len(line) > 3 else line,
                'level': 2,
                'content': '',
                '思路': '分析CDE关注点，制定回复策略',
                '模板': '请参照CDE通知函原文进行回复',
                '实验': '视具体情况而定',
                '项目': '根据CDE要求确定',
                '法规': 'ICH Q系列；中国药典；相关技术指导原则'
            })
        # 三级标题：① ② ③...
        elif re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', line):
            questions.append({
                'num': line[0],
                'title': line[1:].strip() if len(line) > 1 else '',
                'level': 3,
                'content': line
            })
    
    return drug_info, questions


def copy_header_from_template(doc, template_path, drug_name):
    """从模板复制页眉（含logo）"""
    template_zip = zipfile.ZipFile(template_path, 'r')
    
    header_files = [f for f in template_zip.namelist() if 'header' in f and f.endswith('.xml')]
    header_data = {}
    
    for header_file in header_files:
        header_xml = template_zip.read(header_file).decode('utf-8')
        header_xml = header_xml.replace('药品名称', drug_name)
        header_xml = header_xml.replace('棕榈酸帕利哌酮', drug_name)
        header_xml = header_xml.replace('盐酸鲁拉西酮', drug_name)
        header_data[header_file] = header_xml
    
    # 读取图片
    image_files = {}
    for f in template_zip.namelist():
        if 'media' in f and f.endswith(('.png', '.jpg', '.jpeg')):
            image_files[f] = template_zip.read(f)
    
    template_zip.close()
    return header_data, image_files


def create_header_with_logo(doc, section, drug_name, logo_path=None):
    """创建页眉（含logo和页码）"""
    header = section.header
    header.is_linked_to_previous = False
    
    for para in header.paragraphs:
        para.clear()
    
    table = header.add_table(1, 2, Inches(6.5))
    table.autofit = False
    table.allow_autofit = False
    
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(2.0)
    
    left_cell = table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    lines = [
        "原料药登记资料",
        f"{drug_name}（无菌，原料药）",
        "补充资料研究报告"
    ]
    for i, line in enumerate(lines):
        if i > 0:
            left_para.add_run("\n")
        r = left_para.add_run(line)
        set_run_font(r, size=10.5)
    
    right_cell = table.cell(0, 1)
    
    logo_para = right_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if logo_path and os.path.exists(logo_path):
        try:
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(0.8))
        except:
            r = logo_para.add_run("[logo]")
            set_run_font(r, size=9)
    else:
        r = logo_para.add_run("[logo]")
        set_run_font(r, size=9)
    
    page_para = right_cell.add_paragraph()
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = page_para.add_run("第 ")
    set_run_font(r, size=10.5)
    
    run = page_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = " PAGE \\* MERGEFORMAT "
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    set_run_font(run, size=10.5)
    
    r = page_para.add_run(" 页 共 ")
    set_run_font(r, size=10.5)
    
    run = page_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = " NUMPAGES \\* MERGEFORMAT "
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    set_run_font(run, size=10.5)
    
    r = page_para.add_run(" 页")
    set_run_font(r, size=10.5)
    
    tblPr = table._tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_cover_page(doc, drug_info):
    """添加封面页"""
    for _ in range(2):
        doc.add_paragraph()
    
    add_para(doc, "原料药登记资料", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    
    info_items_16pt = [
        ("资料类型：", drug_info.get('doc_type', '补充资料')),
        ("产品名称：", drug_info.get('drug_name', '')),
        ("登 记 号：", drug_info.get('registration_no', '')),
        ("受 理 号：", drug_info.get('acceptance_no', '')),
        ("申请人/注册代理机构名称：", drug_info.get('company', '')),
        ("申请事项：", drug_info.get('application_type', '首次登记')),
        ("联 系 人：", drug_info.get('contact', '魏永红')),
        ("联系电话：", drug_info.get('phone', '0756-8686208')),
    ]
    
    for label, value in info_items_16pt:
        p = doc.add_paragraph()
        r = p.add_run(label + value)
        set_run_font(r, size=16)
        p.paragraph_format.line_spacing = 1.5
    
    info_items_14pt = [
        ("注册地址：", drug_info.get('address', '中国，广东省，珠海市，珠海保税区联峰路22号')),
        ("邮政编码：", drug_info.get('zipcode', '519030')),
    ]
    
    for label, value in info_items_14pt:
        p = doc.add_paragraph()
        r = p.add_run(label + value)
        set_run_font(r, size=14)
        p.paragraph_format.line_spacing = 1.5
    
    for label, value in [("关联制剂受理号：", drug_info.get('related_no', '/')), 
                         ("Email：", drug_info.get('email', 'lzsyntpharmzhucebu@livzon.cn'))]:
        p = doc.add_paragraph()
        r = p.add_run(label + value)
        set_run_font(r, size=16)
        p.paragraph_format.line_spacing = 1.5
    
    for _ in range(2):
        doc.add_paragraph()
    
    add_para(doc, drug_info.get('drug_name', ''), size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "补充资料研究报告", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_toc_page(doc, questions):
    """根据问题列表生成目录"""
    add_para(doc, "目  录", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    
    page_num = 1
    for q in questions:
        if q.get('level') == 1:
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            tabs = pPr.get_or_add_tabs()
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), '8500')
            tab.set(qn('w:leader'), 'dot')
            tabs.append(tab)
            
            r = p.add_run(q['title'])
            set_run_font(r, size=12)
            p.add_run("\t")
            r = p.add_run(str(page_num))
            set_run_font(r, size=12)
            p.paragraph_format.line_spacing = 1.5
            page_num += 1
        elif q.get('level') == 2:
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            tabs = pPr.get_or_add_tabs()
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), '8500')
            tab.set(qn('w:leader'), 'dot')
            tabs.append(tab)
            p.paragraph_format.left_indent = Cm(0.74)
            
            title_short = q['title'][:20] + '...' if len(q['title']) > 20 else q['title']
            r = p.add_run(f"{q['num']}{title_short}")
            set_run_font(r, size=12)
            p.add_run("\t")
            r = p.add_run(str(page_num))
            set_run_font(r, size=12)
            p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    add_para(doc, "附件目录", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    
    for att in ["附件1 生产工艺信息表", "附件2 质量标准"]:
        p = doc.add_paragraph()
        r = p.add_run(att)
        set_run_font(r, size=12)
        p.paragraph_format.left_indent = Cm(0.74)


def add_body_content(doc, questions, drug_info):
    """添加正文内容"""
    BLUE = RGBColor(0, 0, 255)
    BLACK = RGBColor(0, 0, 0)
    
    notice_date = drug_info.get('notice_date', '2024年9月29日')
    notice_year = drug_info.get('notice_year', '2024')
    notice_no = drug_info.get('notice_no', '5308')
    drug_name = drug_info.get('drug_name', '本品')
    reg_no = drug_info.get('registration_no', '')
    acc_no = drug_info.get('acceptance_no', '')
    
    intro1 = f"{notice_date}，国家药品监督管理局药品审评中心下发补充资料通知件（药审补字[{notice_year}]第{notice_no}号），认为我司申报的{drug_name}（登记号为：{reg_no}，受理号为{acc_no}）尚需提交补充资料，以完善本品有关安全性、有效性和质量可控性。"
    add_para(doc, intro1, indent=0.74)
    doc.add_paragraph()
    
    # 统计问题数量
    main_q_count = len([q for q in questions if q.get('level') == 1])
    sub_q_count = len([q for q in questions if q.get('level') == 2])
    
    intro2 = f"审评意见共包含：{main_q_count}个大项，{sub_q_count}个具体问题。"
    add_para(doc, intro2, indent=0.74)
    doc.add_paragraph()
    
    intro3 = f"结合发补回复，我司修订了{drug_name}的生产工艺信息表和质量标准，详见如下附件："
    add_para(doc, intro3, indent=0.74)
    doc.add_paragraph()
    
    for att in ["附件1 生产工艺信息表", "附件2 质量标准"]:
        p = doc.add_paragraph()
        r = p.add_run(att)
        set_run_font(r, size=12)
        p.paragraph_format.left_indent = Cm(0.74)
    
    doc.add_paragraph()
    add_para(doc, "我司对贵中心提出的问题进行了详细的研究，各问题逐条回复见下文。", indent=0.74)
    doc.add_paragraph()
    
    for q in questions:
        if q.get('level') == 1:
            add_para(doc, f"{q['num']}{q['title']}", bold=True, color=BLACK)
        elif q.get('level') == 2:
            add_para(doc, f"{q['num']}{q['title']}", bold=True, color=BLUE)
            add_para(doc, "【答复】", bold=True, color=BLACK)
            add_para(doc, "（请在此处填写回复内容）", color=BLACK, indent=0.74)
            doc.add_paragraph()
            add_para(doc, "【回复建议】", bold=True, color=BLACK)
            for s in [f"• 回复思路：{q.get('思路','')}", 
                      f"• 回复模板：{q.get('模板','')}", 
                      f"• 补充实验：{q.get('实验','')}", 
                      f"• 实验项目：{q.get('项目','')}", 
                      f"• 法规依据：{q.get('法规','')}"]:
                add_para(doc, s, color=BLACK, indent=0.74)
            doc.add_paragraph()
        elif q.get('level') == 3:
            add_para(doc, f"{q['num']}{q['title']}", color=BLUE, indent=0.74)


def update_header_in_docx(output_path, header_data, image_files):
    """通过zip操作更新docx文件中的页眉"""
    temp_path = output_path + ".tmp"
    
    with zipfile.ZipFile(output_path, 'r') as zip_in:
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
            for item in zip_in.namelist():
                if 'header' in item and item.endswith('.xml'):
                    continue
                if 'media' in item and item in image_files:
                    continue
                zip_out.writestr(item, zip_in.read(item))
            
            for header_file, header_xml in header_data.items():
                zip_out.writestr(header_file, header_xml)
            
            for image_file, image_data in image_files.items():
                zip_out.writestr(image_file, image_data)
            
            content_types = zip_in.read('[Content_Types].xml').decode('utf-8')
            if 'image/png' not in content_types and image_files:
                content_types = content_types.replace(
                    '</Types>',
                    '  <Default Extension="png" ContentType="image/png"/>\n</Types>'
                )
            zip_out.writestr('[Content_Types].xml', content_types)
    
    shutil.move(temp_path, output_path)


def generate_reply_template(notice_path, template_path, output_path, drug_info_override=None):
    """生成发补回复模板"""
    
    # 从PDF提取信息
    extracted_info, questions = extract_from_pdf(notice_path)
    
    # 合并覆盖的信息
    if drug_info_override:
        extracted_info.update(drug_info_override)
    
    drug_info = {
        'drug_name': extracted_info.get('drug_name', ''),
        'registration_no': extracted_info.get('registration_no', ''),
        'acceptance_no': extracted_info.get('acceptance_no', ''),
        'company': extracted_info.get('company', '珠海保税区丽珠合成制药有限公司'),
        'doc_type': '补充资料',
        'application_type': '首次登记',
        'contact': extracted_info.get('contact', '魏永红'),
        'phone': extracted_info.get('phone', '0756-8686208'),
        'address': extracted_info.get('address', '中国，广东省，珠海市，珠海保税区联峰路22号'),
        'zipcode': extracted_info.get('zipcode', '519030'),
        'related_no': '/',
        'email': extracted_info.get('email', 'lzsyntpharmzhucebu@livzon.cn'),
        'notice_date': extracted_info.get('notice_date', '2024年9月29日'),
        'notice_year': extracted_info.get('notice_year', '2024'),
        'notice_no': extracted_info.get('notice_no', '5308'),
    }
    
    if not questions:
        print("警告：未能从PDF提取问题列表，将使用示例问题列表")
        questions = get_default_questions()
    
    doc = Document()
    
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    
    add_cover_page(doc, drug_info)
    doc.add_page_break()
    
    new_sec = doc.add_section()
    new_sec.page_height = Cm(29.7)
    new_sec.page_width = Cm(21)
    new_sec.top_margin = Cm(2.54)
    new_sec.bottom_margin = Cm(2.54)
    new_sec.left_margin = Cm(2.54)
    new_sec.right_margin = Cm(2.54)
    
    header_data, image_files = {}, {}
    if template_path and os.path.exists(template_path):
        header_data, image_files = copy_header_from_template(doc, template_path, drug_info.get('drug_name', ''))
        create_header_with_logo(doc, new_sec, drug_info.get('drug_name', ''))
    else:
        create_header_with_logo(doc, new_sec, drug_info.get('drug_name', ''))
    
    add_toc_page(doc, questions)
    doc.add_page_break()
    
    add_body_content(doc, questions, drug_info)
    
    doc.save(output_path)
    
    if header_data and template_path:
        try:
            update_header_in_docx(output_path, header_data, image_files)
        except Exception as e:
            print(f"复制页眉时出错：{e}")
    
    print(f"✅ 文档已生成：{output_path}")
    print(f"包含：封面页 + 目录页 + {len(questions)}个问题 + 页眉")
    print(f"药品名称：{drug_info['drug_name']}")
    print(f"登记号：{drug_info['registration_no']}")
    print(f"受理号：{drug_info['acceptance_no']}")


def get_default_questions():
    """获取默认问题列表（当PDF解析失败时使用）"""
    return [
        {"num": "1、", "title": "证明性文件", "level": 1},
        {"num": "（1）", "title": "请提供相关证明性文件", "level": 2, 
         "思路": "按CDE要求提供", "模板": "已提供相关证明性文件", "实验": "否", "项目": "无", "法规": "药品管理法"},
        {"num": "2、", "title": "生产工艺", "level": 1},
        {"num": "（1）", "title": "起始物料问题", "level": 2,
         "思路": "完善起始物料研究", "模板": "参照CDE通知要求完善", "实验": "是", "项目": "起始物料研究", "法规": "ICH Q11"},
        {"num": "3、", "title": "质量研究与质量标准", "level": 1},
        {"num": "4、", "title": "稳定性研究", "level": 1},
    ]


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='CDE发补回复模板生成器 v2.0')
    parser.add_argument('--notice', required=True, help='CDE通知函PDF路径')
    parser.add_argument('--template', help='公司模板Word路径（可选）')
    parser.add_argument('--output', required=True, help='输出Word路径')
    parser.add_argument('--drug-name', help='药品名称（可选，默认从PDF提取）')
    parser.add_argument('--acceptance-no', help='受理号（可选，默认从PDF提取）')
    parser.add_argument('--registration-no', help='登记号（可选，默认从PDF提取）')
    parser.add_argument('--company', help='公司名称（可选，默认从PDF提取）')
    
    args = parser.parse_args()
    
    drug_info_override = {}
    if args.drug_name:
        drug_info_override['drug_name'] = args.drug_name
    if args.acceptance_no:
        drug_info_override['acceptance_no'] = args.acceptance_no
    if args.registration_no:
        drug_info_override['registration_no'] = args.registration_no
    if args.company:
        drug_info_override['company'] = args.company
    
    generate_reply_template(
        args.notice,
        args.template,
        args.output,
        drug_info_override if drug_info_override else None
    )
