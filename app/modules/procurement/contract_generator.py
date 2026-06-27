"""Generate procurement contracts from controlled Word templates."""

# ruff: noqa: E501

from __future__ import annotations

import copy
from collections.abc import Iterable
from dataclasses import dataclass
from decimal import ROUND_HALF_UP, Decimal
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

from app.modules.procurement.schemas import (
    ContractCategory,
    ContractGenerateRequest,
    ContractItemInput,
    ContractPartyInfo,
    ContractTemplateField,
    ContractTemplateMetadata,
)

TEMPLATE_DIR = Path(__file__).resolve().parent / "assets" / "contract_templates"

TEMPLATE_FILES = {
    ContractCategory.raw_materials: "raw-materials.docx",
    ContractCategory.consumables: "consumables.docx",
    ContractCategory.hardware: "hardware.docx",
    ContractCategory.fixed_assets: "fixed-assets.docx",
}

CATEGORY_LABELS = {
    ContractCategory.raw_materials: "原材料",
    ContractCategory.consumables: "耗材",
    ContractCategory.hardware: "五金",
    ContractCategory.fixed_assets: "固定资产",
}


@dataclass(frozen=True)
class AmountSummary:
    total: Decimal
    untaxed: Decimal
    tax: Decimal


COMMON_FIELDS = [
    ContractTemplateField(name="contract_number", label="合同编号", required=True),
    ContractTemplateField(
        name="contract_date",
        label="签订日期",
        input_type="date",
        required=True,
    ),
    ContractTemplateField(
        name="delivery_date", label="最迟交货日期", input_type="date"
    ),
    ContractTemplateField(name="delivery_terms", label="交货说明"),
    ContractTemplateField(name="payment_terms", label="付款期限/方式"),
    ContractTemplateField(
        name="tax_rate", label="税率", input_type="number", default_value="13"
    ),
    ContractTemplateField(name="seller.name", label="卖方名称", required=True),
    ContractTemplateField(name="seller.representative", label="卖方代表人"),
    ContractTemplateField(name="seller.address", label="卖方地址"),
    ContractTemplateField(name="seller.postal_code", label="邮编"),
    ContractTemplateField(name="seller.contact_person", label="卖方联系人"),
    ContractTemplateField(name="seller.contact_address", label="联系人地址"),
    ContractTemplateField(name="seller.contact_phone", label="联系人电话"),
    ContractTemplateField(name="seller.mobile", label="联系人手机"),
    ContractTemplateField(name="seller.phone", label="卖方电话"),
    ContractTemplateField(name="seller.bank_name", label="开户行"),
    ContractTemplateField(name="seller.bank_account", label="账号"),
    ContractTemplateField(name="seller.tax_id", label="税号/统一社会信用代码"),
]

CONSUMABLE_FIELDS = [
    ContractTemplateField(name="buyer_invoice_recipient", label="发票接收人"),
    ContractTemplateField(
        name="buyer_invoice_recipient_mobile", label="发票接收人手机"
    ),
    ContractTemplateField(name="buyer_receiver", label="收货人"),
    ContractTemplateField(name="buyer_receiver_mobile", label="收货人手机"),
    ContractTemplateField(name="buyer_receiver_phone", label="收货人电话"),
]

FIXED_ASSET_FIELDS = [
    ContractTemplateField(name="attached_documents", label="随货资料"),
    ContractTemplateField(name="seller.bank_line_number", label="银行行号"),
    ContractTemplateField(
        name="installation_days", label="安装调试天数", input_type="number"
    ),
    ContractTemplateField(
        name="warranty_months", label="质保期（月）", input_type="number"
    ),
    ContractTemplateField(
        name="response_hours", label="质保响应小时", input_type="number"
    ),
    ContractTemplateField(
        name="onsite_hours", label="到场处理小时", input_type="number"
    ),
    ContractTemplateField(
        name="maintenance_response_hours",
        label="质保期满维修响应小时",
        input_type="number",
    ),
    ContractTemplateField(
        name="overdue_days", label="逾期解除天数", input_type="number"
    ),
    ContractTemplateField(name="jurisdiction", label="争议管辖地"),
    ContractTemplateField(name="attachment_note", label="附件说明"),
    ContractTemplateField(name="copies", label="合同总份数", input_type="number"),
    ContractTemplateField(name="buyer_copies", label="买方执份数", input_type="number"),
    ContractTemplateField(name="arrival_payment_condition", label="到货款支付条件"),
    ContractTemplateField(name="arrival_payment_method", label="到货款支付方式"),
    ContractTemplateField(
        name="arrival_payment_ratio", label="到货款比例", input_type="number"
    ),
    ContractTemplateField(
        name="warranty_payment_ratio", label="质保金比例", input_type="number"
    ),
    ContractTemplateField(name="warranty_payment_method", label="质保金支付方式"),
]


def get_contract_template_metadata(
    category: ContractCategory,
) -> ContractTemplateMetadata:
    fields = list(COMMON_FIELDS)
    if category is ContractCategory.consumables:
        fields.extend(CONSUMABLE_FIELDS)
    if category is ContractCategory.fixed_assets:
        fields.extend(FIXED_ASSET_FIELDS)
    return ContractTemplateMetadata(
        category=category,
        label=CATEGORY_LABELS[category],
        fields=fields,
    )


def generate_contract(payload: ContractGenerateRequest) -> tuple[BytesIO, str, str]:
    """Generate a contract file and return buffer, filename, media type."""
    template_path = TEMPLATE_DIR / TEMPLATE_FILES[payload.category]
    if not template_path.exists():
        raise FileNotFoundError(f"合同模板不存在：{template_path.name}")

    doc = Document(str(template_path))
    if payload.category is ContractCategory.raw_materials:
        _fill_raw_material_contract(doc, payload)
    elif payload.category is ContractCategory.consumables:
        _fill_consumables_contract(doc, payload)
    elif payload.category is ContractCategory.hardware:
        _fill_hardware_contract(doc, payload)
    elif payload.category is ContractCategory.fixed_assets:
        _fill_fixed_asset_contract(doc, payload)

    _clear_remaining_red_runs(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    label = CATEGORY_LABELS[payload.category]
    safe_number = _safe_filename(payload.contract_number)
    filename = f"{label}合同_{safe_number}.docx"
    media_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    return buffer, filename, media_type


def _fill_raw_material_contract(
    doc: Document, payload: ContractGenerateRequest
) -> None:
    _replace_red_groups(
        _find_paragraph(doc, "签订日期"),
        [_date_cn(payload.contract_date), payload.contract_number],
    )
    _replace_red_groups(
        _find_paragraph(doc, "二、交货日期"),
        [payload.delivery_terms or _delivery_text(payload)],
    )
    _replace_red_groups(
        _find_paragraph(doc, "五、付款方式及期限"),
        [
            payload.payment_terms
            or "货到验收合格后，凭卖方开具的增值税发票30天内以银行承兑汇票支付货款"
        ],
    )

    summary = _amount_summary(payload)
    table = doc.tables[0]
    summary_start = _fit_detail_rows(
        table,
        template_row_index=1,
        first_detail_index=1,
        insert_before_index=2,
        count=len(payload.items),
    )
    for index, item in enumerate(payload.items, start=1):
        row = table.rows[index]
        _set_cells(
            row.cells,
            [
                item.item_code,
                item.name,
                item.specification,
                item.specification,
                item.unit,
                _decimal_text(item.quantity),
                f"{_money_text(item.unit_price)}元",
                f"{_money_text(_item_amount(item))}元",
                item.manufacturer or payload.seller.name,
                item.remarks or "/",
            ],
        )

    _set_cells(
        table.rows[summary_start].cells,
        [
            f"合同金额（含{_decimal_text(payload.tax_rate)}%增值税）：",
            f"合同金额（含{_decimal_text(payload.tax_rate)}%增值税）：",
            f"合同金额（含{_decimal_text(payload.tax_rate)}%增值税）：",
            f"大写（人民币）：{rmb_upper(summary.total)}",
            f"大写（人民币）：{rmb_upper(summary.total)}",
            f"大写（人民币）：{rmb_upper(summary.total)}",
            f"大写（人民币）：{rmb_upper(summary.total)}",
            f"大写（人民币）：{rmb_upper(summary.total)}",
            f"小写：￥{_money_text(summary.total)}",
            f"小写：￥{_money_text(summary.total)}",
        ],
    )
    _center_cells(table.rows[summary_start].cells[:3])
    _set_cells(
        table.rows[summary_start + 1].cells,
        [
            "不含税金额：",
            "不含税金额：",
            "不含税金额：",
            f"大写（人民币）：{rmb_upper(summary.untaxed)}",
            f"大写（人民币）：{rmb_upper(summary.untaxed)}",
            f"大写（人民币）：{rmb_upper(summary.untaxed)}",
            f"大写（人民币）：{rmb_upper(summary.untaxed)}",
            f"大写（人民币）：{rmb_upper(summary.untaxed)}",
            f"小写：￥{_money_text(summary.untaxed)}",
            f"小写：￥{_money_text(summary.untaxed)}",
        ],
    )
    _center_cells(table.rows[summary_start + 1].cells[:3])
    _set_cells(
        table.rows[summary_start + 2].cells,
        [
            "税额：",
            "税额：",
            "税额：",
            f"大写（人民币）：{rmb_upper(summary.tax)}",
            f"大写（人民币）：{rmb_upper(summary.tax)}",
            f"大写（人民币）：{rmb_upper(summary.tax)}",
            f"大写（人民币）：{rmb_upper(summary.tax)}",
            f"大写（人民币）：{rmb_upper(summary.tax)}",
            f"小写：￥{_money_text(summary.tax)}",
            f"小写：￥{_money_text(summary.tax)}",
        ],
    )
    _center_cells(table.rows[summary_start + 2].cells[:3])
    _fill_raw_material_party_table(doc.tables[1], payload.seller)


def _fill_consumables_contract(doc: Document, payload: ContractGenerateRequest) -> None:
    _replace_red_groups(
        _find_paragraph(doc, "签订日期"),
        [_date_plain_cn(payload.contract_date), payload.contract_number],
    )
    _replace_red_groups(
        _find_paragraph(doc, "交货日期"),
        _delivery_date_groups(payload),
    )
    _replace_red_groups(
        _find_paragraph(doc, "付款期限"),
        [payload.payment_terms or "买方收货且检验合格后30天以6个月承兑汇票方式付款"],
    )
    if payload.overdue_days is not None:
        _replace_red_groups(
            _find_paragraph(doc, "卖方逾期供货"), [f"{payload.overdue_days}天"]
        )

    summary = _amount_summary(payload)
    table = doc.tables[0]
    summary_row = len(table.rows) - 1
    summary_row = _fit_detail_rows(
        table,
        template_row_index=1,
        first_detail_index=1,
        insert_before_index=summary_row,
        count=len(payload.items),
    )
    for row_index in range(1, summary_row):
        item = payload.items[row_index - 1]
        values = [
            str(row_index),
            item.name,
            item.specification,
            _decimal_text(item.quantity),
            item.unit,
            _money_text(item.unit_price),
            _money_text(_item_amount(item)),
            item.department or item.remarks,
        ]
        _set_cells(table.rows[row_index].cells, values)
    _set_cells(
        table.rows[summary_row].cells,
        [_summary_text(summary, payload.tax_rate)] * 8,
    )
    _fill_simple_party_table(
        doc.tables[1],
        payload.seller,
        include_contact_phone=False,
        contact_person_label="代表人姓名（当货品有质量问题等联络用）",
        tax_label="税务登记号",
    )
    _fill_consumables_buyer_delivery_table(doc.tables[1], payload)


def _fill_hardware_contract(doc: Document, payload: ContractGenerateRequest) -> None:
    _replace_red_groups(
        _find_paragraph(doc, "签订日期"),
        [_date_plain_cn(payload.contract_date), payload.contract_number],
    )
    _replace_red_groups(
        _find_paragraph(doc, "交货日期"), [_delivery_single_date_text(payload)]
    )
    _replace_red_groups(
        _find_paragraph(doc, "付款期限"),
        [
            payload.payment_terms
            or "收到货并验收无误后，卖方开具全额增值税专用发票，买方收货后150日内以6个月承兑方式支付"
        ],
    )

    summary = _amount_summary(payload)
    table = doc.tables[0]
    summary_row = len(table.rows) - 1
    summary_row = _fit_detail_rows(
        table,
        template_row_index=1,
        first_detail_index=1,
        insert_before_index=summary_row,
        count=len(payload.items),
    )
    for row_index in range(1, summary_row):
        item = payload.items[row_index - 1]
        values = [
            str(row_index),
            item.item_code,
            item.name,
            item.department,
            _decimal_text(item.quantity),
            _money_text(item.unit_price),
            _money_text(_item_amount(item)),
            item.unit,
            item.remarks,
        ]
        _set_cells(table.rows[row_index].cells, values)
    _set_cells(
        table.rows[summary_row].cells,
        [_summary_text(summary, payload.tax_rate)] * 9,
    )
    _fill_simple_party_table(doc.tables[1], payload.seller, include_postal_code=True)


def _fill_fixed_asset_contract(doc: Document, payload: ContractGenerateRequest) -> None:
    _replace_red_groups(_find_paragraph(doc, "合同编号"), [payload.contract_number])
    _replace_red_groups(
        _find_paragraph(doc, "卖方在交货时应附上"),
        [payload.attached_documents or "A、B、C"],
    )
    _replace_red_groups(
        _find_paragraph(doc, "最迟交货日期"),
        [
            _date_cn(payload.delivery_date or payload.contract_date),
            payload.delivery_terms or "双方合同盖章回传后20天内交货",
        ],
    )
    _replace_red_groups(_find_paragraph(doc, "制造的全新"), [payload.seller.name])
    if payload.installation_days is not None:
        _replace_red_groups(
            _find_paragraph(doc, "安装调试及质量验收"), [str(payload.installation_days)]
        )
    _replace_red_groups(
        _find_paragraph(doc, "本合同生效后"),
        [
            payload.arrival_payment_condition or "货到验收完成后30",
            payload.arrival_payment_method or "6个月承兑汇票",
            _ratio_text(payload.arrival_payment_ratio, "90%"),
        ],
    )
    _replace_red_groups(
        _find_paragraph(doc, "余款为质保金"),
        [
            _ratio_text(payload.warranty_payment_ratio, "10%"),
            f"{payload.warranty_months or 12}个月",
            payload.warranty_payment_method or "承兑",
        ],
    )
    _replace_red_groups(
        _find_paragraph_after(doc, "5、卖方收款信息", "名称"),
        [f"名称：{payload.seller.name}"],
    )
    _replace_red_groups(
        _find_paragraph_after(doc, "5、卖方收款信息", "纳税人识别号"),
        [f"纳税人识别号：{payload.seller.tax_id}"],
    )
    _replace_red_groups(
        _find_paragraph_after(doc, "5、卖方收款信息", "地址："),
        [f"地址：{payload.seller.address}"],
    )
    _replace_red_groups(
        _find_paragraph_after(doc, "5、卖方收款信息", "电话："),
        [f"电话：{payload.seller.phone}"],
    )
    _replace_red_groups(
        _find_paragraph_after(doc, "5、卖方收款信息", "开户行"),
        [f"开户行：{payload.seller.bank_name}"],
    )
    _replace_red_groups(
        _find_paragraph_after(doc, "5、卖方收款信息", "账号："),
        [f"账号：{payload.seller.bank_account}"],
    )
    _replace_red_groups(
        _find_paragraph_after(doc, "5、卖方收款信息", "银行行号"),
        [f"银行行号：{payload.seller.bank_line_number}"],
    )
    if payload.warranty_months is not None:
        _replace_red_groups(
            _find_paragraph(doc, "设备验收合格之日起"), [str(payload.warranty_months)]
        )
    _replace_red_groups(
        _find_paragraph(doc, "接到买方电话"),
        [str(payload.response_hours or 24), str(payload.onsite_hours or 48)],
    )
    maintenance_hours = str(payload.maintenance_response_hours or 48)
    maintenance_paragraph = _find_paragraph(doc, "质保期满")
    _replace_red_groups(maintenance_paragraph, [maintenance_hours])
    _remove_placeholder_underscores(
        maintenance_paragraph,
        maintenance_hours,
        suffix="小时",
    )
    _replace_red_groups(
        _find_paragraph(doc, "卖方在最迟交货日期后逾期"),
        [str(payload.overdue_days or 2), "30%", "7", "30%"],
    )
    _replace_red_groups(_find_paragraph(doc, "产品质量缺陷"), ["15", "500"])
    _replace_red_groups(_find_paragraph(doc, "不按时开具"), ["100"])
    if payload.jurisdiction:
        _replace_red_groups(_find_paragraph(doc, "有管辖权"), [payload.jurisdiction])
    _replace_red_groups(
        _find_paragraph(doc, "附件 ①"), [payload.attachment_note or "/"]
    )
    _replace_red_groups(
        _find_paragraph(doc, "本合同一式"),
        [str(payload.copies or 2), str(payload.buyer_copies or 1)],
    )
    _replace_red_groups(
        _find_paragraph(doc, "签订日期"), [_date_cn(payload.contract_date)]
    )

    summary = _amount_summary(payload)
    table = doc.tables[0]
    summary_row = _fit_detail_rows(
        table,
        template_row_index=1,
        first_detail_index=1,
        insert_before_index=2,
        count=len(payload.items),
    )
    for index, item in enumerate(payload.items, start=1):
        _set_cells(
            table.rows[index].cells,
            [
                item.name,
                item.specification,
                item.manufacturer or payload.seller.name,
                f"{_decimal_text(item.quantity)}{item.unit}",
                _money_text(item.unit_price),
                _money_text(_item_amount(item)),
                item.remarks,
            ],
        )
    _set_cells(
        table.rows[summary_row].cells,
        [_summary_text(summary, payload.tax_rate)] * 7,
    )
    _fill_fixed_asset_party_table(doc.tables[1], payload.seller)


def _amount_summary(payload: ContractGenerateRequest) -> AmountSummary:
    total = sum((_item_amount(item) for item in payload.items), Decimal("0"))
    tax_base = Decimal("1") + payload.tax_rate / Decimal("100")
    untaxed = (total / tax_base).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    tax = (total - untaxed).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return AmountSummary(
        total=total.quantize(Decimal("0.01")), untaxed=untaxed, tax=tax
    )


def _item_amount(item: ContractItemInput) -> Decimal:
    if item.amount is not None:
        return item.amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return (item.quantity * item.unit_price).quantize(
        Decimal("0.01"), rounding=ROUND_HALF_UP
    )


def _summary_text(summary: AmountSummary, tax_rate: Decimal) -> str:
    return (
        f"总价：小写：￥{_money_text(summary.total)}    大写：{rmb_upper(summary.total)}"
        f"（含税价，含{_decimal_text(tax_rate)}%增值税）\n"
        f"小写：￥{_money_text(summary.untaxed)}    大写：{rmb_upper(summary.untaxed)}（无税价）\n"
        f"税额：￥{_money_text(summary.tax)}    大写：{rmb_upper(summary.tax)}"
    )


def _fill_raw_material_party_table(table, seller: ContractPartyInfo) -> None:
    _set_cell_text(table.rows[0].cells[0], f"卖方（必须盖章）：{seller.name}")
    _set_cell_text(table.rows[1].cells[0], f"地址：{seller.address}")
    _set_cell_text(table.rows[1].cells[1], f"开户行：{seller.bank_name}")
    _set_cell_text(
        table.rows[2].cells[0],
        f"邮    编：{seller.postal_code}\n电    话：{seller.phone}",
    )
    _set_cell_text(table.rows[2].cells[1], f"帐号：{seller.bank_account}")
    _set_cell_text(table.rows[3].cells[0], f"统一社会信用代码：{seller.tax_id}")
    _set_cell_text(table.rows[3].cells[1], f"统一社会信用代码：{seller.tax_id}")


def _fill_simple_party_table(
    table,
    seller: ContractPartyInfo,
    *,
    include_postal_code: bool = False,
    include_contact_phone: bool = True,
    contact_person_label: str = "代表人姓名",
    tax_label: str = "税号",
) -> None:
    _set_cell_text(table.rows[0].cells[1], seller.name)
    _set_cell_text(table.rows[1].cells[1], seller.representative)
    postal_line = f"\n邮编：{seller.postal_code}" if include_postal_code else ""
    contact_address = seller.contact_address or seller.address
    contact_phone = seller.contact_phone or seller.phone
    mobile = seller.mobile or seller.phone
    contact_phone_line = f"\n电话：{contact_phone}" if include_contact_phone else ""
    _set_cell_text(
        table.rows[2].cells[0],
        f"地址：{seller.address}\n开户行：{seller.bank_name}\n帐 号：{seller.bank_account}\n"
        f"电话：{seller.phone}{postal_line}\n{tax_label}：{seller.tax_id}",
    )
    _set_cell_text(
        table.rows[3].cells[0],
        f"{contact_person_label}：{seller.contact_person}\n地址：{contact_address}"
        f"{contact_phone_line}\n手机：{mobile}\n收合同邮箱：{seller.email}",
    )


def _fill_consumables_buyer_delivery_table(
    table, payload: ContractGenerateRequest
) -> None:
    address = "宁夏石嘴山市平罗太沙工业园丽珠制药"
    _set_cell_text(
        table.rows[3].cells[2],
        "发票接收人及地址（发票必须包含中文名称）\n"
        f"接收人：{payload.buyer_invoice_recipient}      "
        f"手机：{payload.buyer_invoice_recipient_mobile}\n"
        f"地址：{address}\n"
        f"收货人：{payload.buyer_receiver}（请在发货前与收货部门电话或邮件确认）\n"
        f"手机：{payload.buyer_receiver_mobile} 电话：{payload.buyer_receiver_phone}\n"
        f"地址：{address}",
    )


def _fill_fixed_asset_party_table(table, seller: ContractPartyInfo) -> None:
    _set_cell_text(
        table.rows[0].cells[1],
        f"卖方：{seller.name}\n"
        "     （必须盖公章或合同章）\n"
        f"签约代表：{seller.contact_person}\n"
        f"地址：{seller.address}\n"
        f"电话：{seller.phone}\n"
        f"开户银行：{seller.bank_name}\n"
        f"账号：{seller.bank_account}\n"
        f"纳税人识别号：{seller.tax_id}",
    )


def _replace_red_groups(paragraph, replacements: Iterable[str]) -> None:
    if paragraph is None:
        return
    replacement_iter = iter(replacements)
    active_run = None
    in_red_group = False
    for run in paragraph.runs:
        if _is_red_run(run):
            _set_run_black(run)
            if not in_red_group:
                value = next(replacement_iter, "")
                run.text = value
                active_run = run
                in_red_group = True
            elif active_run is not run:
                run.text = ""
        else:
            in_red_group = False
            active_run = None


def _is_red_run(run) -> bool:
    rpr = run._element.rPr
    if rpr is None:
        return False
    color = rpr.find(qn("w:color"))
    if color is None:
        return False
    value = (color.get(qn("w:val")) or "").upper()
    if len(value) == 6:
        try:
            red = int(value[0:2], 16)
            green = int(value[2:4], 16)
            blue = int(value[4:6], 16)
            return red >= 180 and green <= 100 and blue <= 100
        except ValueError:
            return False
    return value in {"FF0000", "C00000"}


def _set_run_black(run) -> None:
    run.font.color.rgb = RGBColor(0, 0, 0)


def _clear_remaining_red_runs(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if _is_red_run(run):
                _set_run_black(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if _is_red_run(run):
                            _set_run_black(run)


def _find_paragraph(doc: Document, contains: str):
    for paragraph in doc.paragraphs:
        if contains in paragraph.text:
            return paragraph
    return None


def _find_paragraph_after(doc: Document, after_contains: str, contains: str):
    matched_anchor = False
    for paragraph in doc.paragraphs:
        if matched_anchor and contains in paragraph.text:
            return paragraph
        if after_contains in paragraph.text:
            matched_anchor = True
    return None


def _remove_placeholder_underscores(paragraph, value: str, *, suffix: str) -> None:
    if paragraph is None:
        return
    text = paragraph.text
    cleaned = text.replace(f"_{value}__", f"{value}{suffix}")
    cleaned = cleaned.replace(f"_{value}_{suffix}", f"{value}{suffix}")
    cleaned = cleaned.replace(f"{value}{suffix}{suffix}", f"{value}{suffix}")
    if cleaned == text:
        return
    if not paragraph.runs:
        paragraph.add_run(cleaned)
        return
    paragraph.runs[0].text = cleaned
    _set_run_black(paragraph.runs[0])
    for run in paragraph.runs[1:]:
        run.text = ""


def _ensure_rows_before(
    table, *, template_row_index: int, insert_before_index: int, count: int
) -> None:
    existing_count = insert_before_index - template_row_index
    if count <= existing_count:
        return
    template_row = table.rows[template_row_index]._tr
    for _ in range(count - existing_count):
        new_row = copy.deepcopy(template_row)
        _insert_row_before(table, insert_before_index, new_row)
        insert_before_index += 1


def _fit_detail_rows(
    table,
    *,
    template_row_index: int,
    first_detail_index: int,
    insert_before_index: int,
    count: int,
) -> int:
    existing_count = insert_before_index - first_detail_index
    if existing_count < 1:
        raise ValueError("合同模板缺少可复制的明细行")

    template_row = table.rows[template_row_index]._tr
    if count > existing_count:
        for _ in range(count - existing_count):
            new_row = copy.deepcopy(template_row)
            _insert_row_before(table, insert_before_index, new_row)
            insert_before_index += 1
    elif count < existing_count:
        for _ in range(existing_count - count):
            row = table.rows[first_detail_index + count]._tr
            table._tbl.remove(row)

    return first_detail_index + count


def _insert_row_before(table, row_index: int, row_element) -> None:
    if row_index >= len(table.rows):
        table._tbl.append(row_element)
        return
    table.rows[row_index]._tr.addprevious(row_element)


def _set_cells(cells, values: list[str]) -> None:
    for cell, value in zip(cells, values, strict=False):
        _set_cell_text(cell, value)


def _set_cell_text(cell, value: str) -> None:
    cell.text = value or ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)


def _center_cells(cells) -> None:
    for cell in cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _delivery_date_groups(payload: ContractGenerateRequest) -> list[str]:
    if payload.delivery_terms and not payload.delivery_date:
        return [payload.delivery_terms]
    delivery_date = payload.delivery_date or payload.contract_date
    return [
        str(delivery_date.year),
        f"{delivery_date.month:02d}",
        f"{delivery_date.day:02d}",
    ]


def _delivery_single_date_text(payload: ContractGenerateRequest) -> str:
    if payload.delivery_terms and not payload.delivery_date:
        return payload.delivery_terms
    return _date_plain_cn(payload.delivery_date or payload.contract_date)


def _delivery_text(payload: ContractGenerateRequest) -> str:
    if payload.delivery_terms:
        return payload.delivery_terms
    if payload.delivery_date:
        return f"{_date_plain_cn(payload.delivery_date)}前"
    return "按买方电话通知送货"


def _date_cn(value) -> str:
    return f"{value.year} 年{value.month:02d}月{value.day:02d}日"


def _date_plain_cn(value) -> str:
    return f"{value.year}年{value.month:02d}月{value.day:02d}日"


def _money_text(value: Decimal) -> str:
    return f"{value.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP):.2f}"


def _decimal_text(value: Decimal) -> str:
    normalized = value.normalize()
    return format(normalized, "f")


def _ratio_text(value: Decimal | None, default: str) -> str:
    if value is None:
        return default
    return f"{_decimal_text(value)}%"


def _safe_filename(value: str) -> str:
    return "".join(
        char if char.isalnum() or char in ("-", "_") else "_" for char in value
    )


_CN_NUM = "零壹贰叁肆伍陆柒捌玖"
_CN_INT_UNITS = ["", "拾", "佰", "仟"]
_CN_SECTION_UNITS = ["", "万", "亿", "兆"]


def rmb_upper(value: Decimal) -> str:
    amount = value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    sign = "负" if amount < 0 else ""
    amount = abs(amount)
    integer = int(amount)
    fraction = int((amount - Decimal(integer)) * Decimal(100))

    integer_text = _integer_to_upper(integer)
    jiao = fraction // 10
    fen = fraction % 10
    if jiao == 0 and fen == 0:
        fraction_text = "整"
    elif jiao == 0:
        fraction_text = f"零{_CN_NUM[fen]}分"
    elif fen == 0:
        fraction_text = f"{_CN_NUM[jiao]}角整"
    else:
        fraction_text = f"{_CN_NUM[jiao]}角{_CN_NUM[fen]}分"
    return f"{sign}{integer_text}元{fraction_text}"


def _integer_to_upper(value: int) -> str:
    if value == 0:
        return "零"
    sections: list[int] = []
    while value:
        sections.append(value % 10000)
        value //= 10000

    parts: list[str] = []
    need_zero = False
    for index in range(len(sections) - 1, -1, -1):
        section = sections[index]
        if section == 0:
            need_zero = bool(parts)
            continue
        if need_zero or (parts and section < 1000):
            parts.append("零")
        parts.append(_section_to_upper(section) + _CN_SECTION_UNITS[index])
        need_zero = False
    return "".join(parts).rstrip("零")


def _section_to_upper(section: int) -> str:
    text = ""
    zero = False
    unit_index = 0
    while section:
        digit = section % 10
        if digit == 0:
            zero = bool(text)
        else:
            if zero:
                text = "零" + text
                zero = False
            text = _CN_NUM[digit] + _CN_INT_UNITS[unit_index] + text
        unit_index += 1
        section //= 10
    return text
