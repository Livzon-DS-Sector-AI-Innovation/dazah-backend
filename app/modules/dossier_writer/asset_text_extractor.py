"""素材文本提取器 - 将各种格式的素材统一转为纯文本，供 AI 解析"""
import logging
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Any

from docx import Document


class AssetTextExtractor:
    """素材文本提取器"""

    @staticmethod
    def extract(file_path: Path) -> Dict[str, Any]:
        """统一提取接口：根据文件类型选择提取方式

        Returns:
            {
                "text": "提取的全文",
                "paragraphs": [{"index": 0, "text": "..."}],
                "tables": [{"index": 0, "rows": N, "cols": M, "data": [[...]]}],
                "page_count": int (仅 PDF),
                "page_texts": [{"page": 1, "text": "..."}] (仅 PDF),
                "error": str (如有错误)
            }
        """
        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            return AssetTextExtractor._extract_docx(file_path)
        elif suffix == ".doc":
            return AssetTextExtractor._extract_doc(file_path)
        elif suffix == ".pdf":
            return AssetTextExtractor._extract_pdf(file_path)
        elif suffix in (".txt", ".csv"):
            return AssetTextExtractor._extract_text(file_path)
        else:
            return {"text": "", "error": f"不支持的文件类型: {suffix}"}

    @staticmethod
    def _extract_docx(file_path: Path) -> Dict[str, Any]:
        """从 docx 提取段落和表格"""
        try:
            doc = Document(str(file_path))

            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs.append({"index": i, "text": text})

            tables = []
            for t_idx, table in enumerate(doc.tables):
                rows_data = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_data.append(cells)
                tables.append({
                    "index": t_idx,
                    "rows": len(table.rows),
                    "cols": len(table.columns),
                    "data": rows_data,
                })

            full_text = "\n".join(p["text"] for p in paragraphs)
            return {"text": full_text, "paragraphs": paragraphs, "tables": tables}
        except Exception as e:
            return {"text": "", "error": f"docx 提取失败: {str(e)}"}

    @staticmethod
    def _extract_doc(file_path: Path) -> Dict[str, Any]:
        """从 .doc 提取：先转 docx 再提取"""
        # 查找同目录下是否有已转换的 .docx 版本
        docx_path = file_path.with_suffix(".docx")
        if docx_path.exists():
            return AssetTextExtractor._extract_docx(docx_path)

        # 尝试用 libreoffice 转换
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run(
                    [
                        "libreoffice", "--headless", "--convert-to", "docx",
                        "--outdir", tmpdir, str(file_path),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                if result.returncode == 0:
                    converted = Path(tmpdir) / (file_path.stem + ".docx")
                    if converted.exists():
                        return AssetTextExtractor._extract_docx(converted)
            return {"text": "", "error": "libreoffice 转换失败，请上传 .docx 格式"}
        except Exception as e:
            return {"text": "", "error": f"doc 转换失败: {str(e)}"}

    @staticmethod
    def _extract_pdf(file_path: Path) -> Dict[str, Any]:
        """从 PDF 提取：优先用 pdfplumber，失败则用 PaddleOCR PP-StructureV3"""
        # 尝试 pdfplumber（对文字型 PDF 更快更准）
        try:
            import pdfplumber
            with pdfplumber.open(str(file_path)) as pdf:
                page_texts = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    page_texts.append({"page": i + 1, "text": text.strip()})

                full_text = "\n\n".join(p["text"] for p in page_texts)
                if full_text.strip():
                    return {
                        "text": full_text,
                        "page_count": len(pdf.pages),
                        "page_texts": page_texts,
                    }
        except Exception:
            logger.warning("PDF extraction with pdfplumber failed, falling back to PaddleOCR")

        # 回退到 PaddleOCR PP-StructureV3（对扫描件更好，保持结构）
        try:
            from app.shared.ocr_service import get_ocr_service
            ocr_service = get_ocr_service()
            
            # 使用 PP-StructureV3 提取 Markdown（保持表格、公式等结构）
            markdown_text = ocr_service.extract_markdown(file_path)
            
            # 也可以获取结构化数据
            structure = ocr_service.extract_structure(file_path)
            
            return {
                "text": markdown_text,
                "page_count": 1,  # PP-StructureV3 不直接提供页数
                "page_texts": [{"page": 1, "text": markdown_text}],
                "structure": structure,  # 包含 layout、tables 等结构化信息
            }
        except Exception as e:
            return {"text": "", "error": f"PDF 提取失败: {str(e)}"}

    @staticmethod
    def _extract_text(file_path: Path) -> Dict[str, Any]:
        """从纯文本文件提取"""
        try:
            text = file_path.read_text(encoding="utf-8", errors="replace")
            paragraphs = [
                {"index": i, "text": line}
                for i, line in enumerate(text.splitlines())
                if line.strip()
            ]
            return {"text": text, "paragraphs": paragraphs, "tables": []}
        except Exception as e:
            return {"text": "", "error": f"文本提取失败: {str(e)}"}

    @staticmethod
    def pdf_page_to_image(file_path: Path, page_number: int, dpi: int = 200) -> Optional[Path]:
        """将 PDF 指定页转为图片，返回图片路径"""
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(
                str(file_path),
                dpi=dpi,
                first_page=page_number,
                last_page=page_number,
            )
            if not images:
                return None

            img_path = file_path.parent / f"{file_path.stem}_page{page_number}.png"
            images[0].save(str(img_path), "PNG")
            return img_path
        except Exception:
            return None
