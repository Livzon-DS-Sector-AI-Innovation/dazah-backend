"""Dossier Writer business workflows."""
import logging
import os
import shutil
import re
from pathlib import Path
from typing import Optional, List, Dict, Any
from uuid import UUID
from datetime import datetime

from docx import Document
# Removed invalid import
from sqlalchemy import select, and_
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import get_settings
from .models import ProductDossier, DossierTemplate, DossierChapter, ChapterAsset
from .repository import DossierRepository
from .m3_structure import M3_CHAPTERS, match_file_to_chapter
from .schemas import (
    ProductDossierCreate, ProductDossierUpdate,
    ChapterResponse, ChapterDetailResponse, AssetResponse,
)



def _chapter_sort_key(code: str):
    """将 chapter_code 转为可排序的整数元组, 如 '3.2.S.1.2' -> (3,2,100,1,2)"""
    if not code:
        return ()
    out = []
    for seg in code.split('.'):
        if seg.upper() == 'S':
            out.append(100)  # 字母段排在数字之后(不会与真实段号冲突)
        else:
            try:
                out.append(int(seg))
            except ValueError:
                out.append(999)
    return tuple(out)


class DossierService:
    """品种资料业务逻辑"""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.repo = DossierRepository(db)
        self.settings = get_settings()
        self.storage_root = Path(self.settings.STORAGE_ROOT) / "registration" / "dossier-writer"

    # ====== Product Dossier ======

    async def create_product_dossier(self, data: ProductDossierCreate) -> ProductDossier:
        """创建品种资料"""
        # 查重：检查是否已存在相同品种名称 + 生产商 + 无菌类型
        existing = await self.repo.check_duplicate(
            data.product_name, data.manufacturer, data.sterile_type
        )
        if existing:
            raise ValueError("该品种资料已存在，请勿重复创建")

        # 创建品种记录
        dossier = ProductDossier(
            product_name=data.product_name,
            sterile_type=data.sterile_type,
            manufacturer=data.manufacturer,
            template_original_product_name=data.template_original_product_name,
            template_original_manufacturer=data.template_original_manufacturer,
            status="draft",
            parse_status="pending",
        )
        dossier = await self.repo.create_product_dossier(dossier)

        # 创建存储目录
        dossier_id = str(dossier.id)
        paths = self._create_storage_dirs(dossier_id)
        
        dossier.source_templates_path = paths["source_templates"]
        dossier.working_path = paths["working"]
        dossier.assets_path = paths["assets"]
        dossier.outputs_path = paths["outputs"]
        
        # 创建 M3 标准目录树
        await self._create_m3_chapters(dossier.id)
        
        # 自动初始化 S.6 章节的 AI 配置（字段映射 + 素材分类）
        await self.init_chapter_ai_config("3.2.S.6")
        
        await self.db.commit()
        dossier_id = dossier.id
        result = await self.db.execute(
            select(ProductDossier).where(ProductDossier.id == dossier_id)
        )
        return result.scalar_one()

    async def get_product_dossier(self, dossier_id: UUID) -> Optional[ProductDossier]:
        """获取品种资料详情"""
        dossier = await self.repo.get_product_dossier(dossier_id)
        if dossier:
            chapter_count = await self.repo.count_chapters(dossier_id)
            dossier.chapter_count = chapter_count
        return dossier

    async def list_product_dossiers(
        self, skip: int = 0, limit: int = 100
    ) -> tuple[List[ProductDossier], int]:
        """获取品种资料列表"""
        items, total = await self.repo.list_product_dossiers(skip, limit)
        # 填充章节数量
        for item in items:
            item.chapter_count = await self.repo.count_chapters(item.id)
        return items, total

    async def update_product_dossier(
        self, dossier_id: UUID, data: ProductDossierUpdate
    ) -> Optional[ProductDossier]:
        """更新品种资料"""
        update_data = data.model_dump(exclude_unset=True)
        if update_data:
            dossier = await self.repo.update_product_dossier(dossier_id, **update_data)
            await self.db.commit()
            return dossier
        return await self.repo.get_product_dossier(dossier_id)

    async def delete_product_dossier(self, dossier_id: UUID) -> bool:
        """删除品种资料"""
        # 软删除数据库记录
        success = await self.repo.delete_product_dossier(dossier_id)
        if success:
            # 删除物理文件
            self._delete_storage_dirs(str(dossier_id))
            await self.db.commit()
        return success

    # ====== Template Management ======

    async def save_template_file(
        self, dossier_id: UUID, filename: str, file_content: bytes
    ) -> DossierTemplate:
        """保存模板文件（同名文件覆盖更新）"""
        dossier = await self.repo.get_product_dossier(dossier_id)
        if not dossier:
            raise ValueError(f"品种资料不存在: {dossier_id}")

        # 保存文件（覆盖同名文件）
        source_dir = Path(dossier.source_templates_path)
        source_dir.mkdir(parents=True, exist_ok=True)
        file_path = source_dir / filename
        file_path.write_bytes(file_content)

        # 查找是否已有同名模板
        existing = await self.repo.get_template_by_filename(dossier_id, filename)
        
        if existing:
            # 更新已有记录
            existing.file_path = str(file_path)
            existing.file_size = len(file_content)
            await self.db.flush()
            template = existing
        else:
            # 创建新记录
            template = DossierTemplate(
                product_dossier_id=dossier_id,
                original_filename=filename,
                file_path=str(file_path),
                file_size=len(file_content),
            )
            template = await self.repo.create_template(template)
        
        await self.db.commit()
        return template

    # ====== Template Parsing ======

    async def parse_templates(self, dossier_id: UUID) -> Dict[str, Any]:
        """解析模板并匹配到 M3 章节"""
        dossier = await self.repo.get_product_dossier(dossier_id)
        if not dossier:
            return {"success": False, "message": "品种资料不存在", "error": "NOT_FOUND"}

        try:
            # 更新状态
            await self.repo.update_product_dossier(
                dossier_id, parse_status="parsing", parse_error=None
            )
            await self.db.commit()

            # 获取模板文件
            templates = await self.repo.list_templates(dossier_id)
            if not templates:
                raise ValueError("没有上传模板文件")

            # 获取现有 M3 章节
            chapters = await self.repo.get_chapter_tree(dossier_id)
            if not chapters:
                # 如果没有章节，创建 M3 标准目录
                await self._create_m3_chapters(dossier_id)
                await self.db.commit()
                chapters = await self.repo.get_chapter_tree(dossier_id)

            # 匹配模板到章节
            matched_count = 0
            for template in templates:
                matched_code = match_file_to_chapter(template.original_filename)
                if matched_code:
                    # 找到对应章节
                    matched_chapter = None
                    for ch in chapters:
                        if ch.chapter_code == matched_code:
                            matched_chapter = ch
                            break
                    
                    if matched_chapter:
                        # 创建 working copy
                        await self._create_working_copy_for_chapter(dossier, template, matched_chapter)
                        matched_count += 1

            # 更新状态
            await self.repo.update_product_dossier(
                dossier_id, 
                parse_status="parsed", 
                status="active"
            )
            await self.db.commit()

            return {
                "success": True,
                "message": f"解析完成，共上传 {len(templates)} 个模板文件，成功匹配 {matched_count} 个",
                "chapter_count": len(chapters),
                "template_count": len(templates),
                "matched_count": matched_count,
            }

        except Exception as e:
            await self.repo.update_product_dossier(
                dossier_id, parse_status="failed", parse_error=str(e)
            )
            await self.db.commit()

            return {
                "success": False,
                "message": f"解析失败: {str(e)}",
                "error": str(e),
            }

    async def _parse_docx_template(
        self, dossier: ProductDossier, template: DossierTemplate
    ) -> List[DossierChapter]:
        """解析单个 DOCX 模板"""
        source_path = Path(template.file_path)
        working_dir = Path(dossier.working_path)
        working_dir.mkdir(parents=True, exist_ok=True)

        # 创建 working copy
        working_filename = self._generate_working_filename(
            template.original_filename, dossier
        )
        working_path = working_dir / working_filename
        shutil.copy2(source_path, working_path)

        # 执行基础信息替换
        self._replace_basic_info(working_path, dossier)

        # 解析章节结构
        doc = Document(str(working_path))
        chapters = []
        sort_order = 0

        for para_idx, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            
            if not text:
                continue

            # 判断是否是标题
            level = self._detect_heading_level(style_name, text)
            if level > 0:
                chapter_code = self._extract_chapter_code(text)
                chapter_title = self._clean_chapter_title(text)
                
                chapter = DossierChapter(
                    product_dossier_id=dossier.id,
                    chapter_code=chapter_code,
                    chapter_title=chapter_title,
                    level=level,
                    sort_order=sort_order,
                    source_file=template.original_filename,
                    working_file=working_filename,
                    paragraph_start=para_idx,
                    paragraph_end=para_idx,
                    has_content=True,
                )
                chapters.append(chapter)
                sort_order += 1

        return chapters

    def _detect_heading_level(self, style_name: str, text: str) -> int:
        """检测标题层级"""
        # 基于样式名称
        style_lower = style_name.lower()
        if "heading 1" in style_lower or style_lower == "heading1":
            return 1
        if "heading 2" in style_lower or style_lower == "heading2":
            return 2
        if "heading 3" in style_lower or style_lower == "heading3":
            return 3
        if "heading 4" in style_lower or style_lower == "heading4":
            return 4
        if "heading" in style_lower:
            return 2  # 默认二级

        # 基于编号规则（如 3.2.S.1.1）
        if re.match(r'^\d+(\.\d+)*\s', text):
            parts = text.split()[0].split('.')
            return min(len(parts), 4)

        return 0  # 不是标题

    def _extract_chapter_code(self, text: str) -> Optional[str]:
        """提取章节编号"""
        match = re.match(r'^(\d+(?:\.\d+)*|[A-Z](?:\.\d+)*)', text)
        if match:
            return match.group(1)
        return None

    def _clean_chapter_title(self, text: str) -> str:
        """清理章节标题"""
        # 移除编号前缀
        text = re.sub(r'^[\d\.]+[A-Z\.]*\s*', '', text)
        return text.strip()

    def _generate_working_filename(
        self, original_filename: str, dossier: ProductDossier
    ) -> str:
        """生成工作副本文件名"""
        name, ext = os.path.splitext(original_filename)
        
        # 替换旧品种名
        if dossier.template_original_product_name:
            name = name.replace(
                dossier.template_original_product_name,
                dossier.product_name
            )
        
        return f"{name}_working{ext}"

    def _replace_basic_info(self, file_path: Path, dossier: ProductDossier) -> None:
        """替换基础信息"""
        doc = Document(str(file_path))

        replacements = {}
        
        # 品种名称替换
        if dossier.template_original_product_name:
            replacements[dossier.template_original_product_name] = dossier.product_name
        
        # 生产商替换
        if dossier.template_original_manufacturer:
            replacements[dossier.template_original_manufacturer] = dossier.manufacturer

        # 无菌类型替换：根据用户选择的类型，替换模板中相反的值
        if dossier.sterile_type == "无菌":
            replacements["非无菌"] = "无菌"
        elif dossier.sterile_type == "非无菌":
            replacements["无菌"] = "非无菌"

        if not replacements:
            doc.save(str(file_path))
            return

        # 替换段落
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)

        # 替换表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)

        # 替换页眉页脚
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                try:
                    for para in header.paragraphs:
                        self._replace_in_paragraph(para, replacements)
                except Exception:
                    logger.warning("Header paragraph replacement failed")
            
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                try:
                    for para in footer.paragraphs:
                        self._replace_in_paragraph(para, replacements)
                except Exception:
                    logger.warning("Footer paragraph replacement failed")

        doc.save(str(file_path))

    def _replace_in_paragraph(self, para, replacements: Dict[str, str]) -> None:
        """替换段落中的文本"""
        for run in para.runs:
            for old_text, new_text in replacements.items():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

    # ====== Chapter Management ======

    async def get_chapter_tree(self, dossier_id: UUID) -> List[ChapterResponse]:
        """获取章节树"""
        chapters = await self.repo.get_chapter_tree(dossier_id)
        return self._build_chapter_tree(chapters)

    def _build_chapter_tree(self, chapters: List[DossierChapter]) -> List[ChapterResponse]:
        """构建章节树结构"""
        chapter_map: Dict[Optional[UUID], List[DossierChapter]] = {}
        
        for chapter in chapters:
            parent_id = chapter.parent_id
            if parent_id not in chapter_map:
                chapter_map[parent_id] = []
            chapter_map[parent_id].append(chapter)

        def build_node(chapter: DossierChapter) -> ChapterResponse:
            children = chapter_map.get(chapter.id, [])
            # 确保子节点按 sort_order 排序
            children = sorted(children, key=lambda c: _chapter_sort_key(c.chapter_code or ''))
            return ChapterResponse(
                id=chapter.id,
                parent_id=chapter.parent_id,
                chapter_code=chapter.chapter_code,
                chapter_title=chapter.chapter_title,
                level=chapter.level,
                sort_order=chapter.sort_order,
                has_content=chapter.has_content,
                has_assets=chapter.has_assets,
                asset_count=len(chapter.assets) if chapter.assets else 0,
                source_file=chapter.source_file,
                working_file=chapter.working_file,
                children=[build_node(c) for c in children],
            )

        # 根节点是 parent_id 为 None 的章节
        root_chapters = chapter_map.get(None, [])
        # 确保根节点也按 sort_order 排序
        root_chapters = sorted(root_chapters, key=lambda c: _chapter_sort_key(c.chapter_code or ''))
        return [build_node(c) for c in root_chapters]

    async def get_chapter_detail(self, chapter_id: UUID) -> Optional[ChapterDetailResponse]:
        """获取章节详情"""
        chapter = await self.repo.get_chapter(chapter_id)
        if not chapter:
            return None
        
        assets = [
            AssetResponse(
                id=a.id,
                original_filename=a.original_filename,
                file_type=a.file_type,
                file_size=a.file_size,
                uploaded_at=a.uploaded_at,
            )
            for a in (chapter.assets or [])
        ]
        
        return ChapterDetailResponse(
            id=chapter.id,
            product_dossier_id=chapter.product_dossier_id,
            chapter_code=chapter.chapter_code,
            chapter_title=chapter.chapter_title,
            level=chapter.level,
            has_content=chapter.has_content,
            has_assets=chapter.has_assets,
            source_file=chapter.source_file,
            working_file=chapter.working_file,
            assets=assets,
        )

    # ====== Asset Management ======

    async def _suggest_category(self, chapter_code: str, filename: str) -> Optional[UUID]:
        """根据文件名自动猜测素材分类"""
        from .field_models import AssetCategory
        
        stmt = select(AssetCategory).where(
            and_(
                AssetCategory.chapter_code == chapter_code,
                AssetCategory.is_deleted == False,
            )
        )
        result = await self.db.execute(stmt)
        categories = list(result.scalars().all())
        
        if not categories:
            return None
        
        fname_lower = filename.lower()
        
        for cat in categories:
            name_lower = cat.category_name.lower()
            # 分类名称出现在文件名中
            if name_lower in fname_lower:
                return cat.id
            # 描述中的关键词（按空格/标点分词）
            if cat.description:
                import re
                desc_words = [w for w in re.split(r'[\s，。、；：""''（）(),;]+', cat.description) if len(w) > 2]
                if any(w.lower() in fname_lower for w in desc_words):
                    return cat.id
        
        return None

    async def upload_chapter_asset(
        self, chapter_id: UUID, filename: str, file_content: bytes
    ) -> ChapterAsset:
        """上传章节素材"""
        chapter = await self.repo.get_chapter(chapter_id)
        if not chapter:
            raise ValueError(f"章节不存在: {chapter_id}")

        dossier = await self.repo.get_product_dossier(chapter.product_dossier_id)
        if not dossier:
            raise ValueError("品种资料不存在")

        # 保存文件
        assets_dir = Path(dossier.assets_path) / str(chapter_id)
        assets_dir.mkdir(parents=True, exist_ok=True)
        file_path = assets_dir / filename
        file_path.write_bytes(file_content)

        # 获取文件类型
        file_type = Path(filename).suffix.lower().lstrip('.')

        # 自动猜测分类
        suggested_category_id = await self._suggest_category(chapter.chapter_code, filename)
        
        # 创建记录
        asset = ChapterAsset(
            chapter_id=chapter_id,
            original_filename=filename,
            file_path=str(file_path),
            file_type=file_type,
            file_size=len(file_content),
            category_id=suggested_category_id,
        )
        asset = await self.repo.create_asset(asset)

        # 更新章节状态
        await self.repo.update_chapter(chapter_id, has_assets=True)
        await self.db.commit()
        
        return asset

    async def list_chapter_assets(self, chapter_id: UUID) -> List[ChapterAsset]:
        """获取章节素材列表"""
        return await self.repo.list_assets(chapter_id)

    async def delete_asset(self, asset_id: UUID) -> bool:
        """删除素材"""
        asset = await self.repo.get_asset(asset_id)
        if not asset:
            return False

        # 删除物理文件
        file_path = Path(asset.file_path)
        if file_path.exists():
            file_path.unlink()

        # 删除记录
        await self.repo.delete_asset(asset_id)
        
        # 更新章节状态
        remaining = await self.repo.count_assets(asset.chapter_id)
        if remaining == 0:
            await self.repo.update_chapter(asset.chapter_id, has_assets=False)
        
        await self.db.commit()
        return True

    # ====== Export ======

    async def export_dossier(
        self, dossier_id: UUID, chapter_ids: Optional[List[UUID]] = None
    ) -> Dict[str, Any]:
        """导出品种资料"""
        dossier = await self.repo.get_product_dossier(dossier_id)
        if not dossier:
            return {"success": False, "message": "品种资料不存在"}

        working_dir = Path(dossier.working_path)
        outputs_dir = Path(dossier.outputs_path)
        outputs_dir.mkdir(parents=True, exist_ok=True)

        export_path = None
        export_filename = None

        # 获取章节信息用于命名
        all_chapters = await self.repo.get_chapter_tree(dossier_id)
        chapter_map = {ch.id: ch for ch in all_chapters}

        if chapter_ids and len(chapter_ids) == 1:
            # 单章节导出：用 章节编号_章节标题.docx
            target_chapter = chapter_map.get(chapter_ids[0])
            if target_chapter and target_chapter.working_file:
                safe_title = re.sub(r'[\/:*?"<>|]', '_', target_chapter.chapter_title)
                code_part = target_chapter.chapter_code or "unknown"
                export_filename = f"{code_part}_{safe_title}.docx"
                export_path = outputs_dir / export_filename
                source_file = working_dir / target_chapter.working_file
                if source_file.exists():
                    shutil.copy2(source_file, export_path)
                else:
                    export_path = None
        elif chapter_ids:
            # 多章节导出：用 品种名_章节编号范围_申报资料.docx
            matched_chapters = [
                chapter_map[cid] for cid in chapter_ids
                if cid in chapter_map and chapter_map[cid].working_file
            ]
            if matched_chapters:
                codes = sorted(set(ch.chapter_code for ch in matched_chapters if ch.chapter_code))
                code_range = f"{codes[0]}-{codes[-1]}" if len(codes) > 1 else (codes[0] if codes else "export")
                export_filename = f"{dossier.product_name}_{code_range}_申报资料.docx"
                export_path = outputs_dir / export_filename
                source_file = working_dir / matched_chapters[0].working_file
                if source_file.exists():
                    shutil.copy2(source_file, export_path)
                else:
                    export_path = None

        # fallback：品种名_申报资料_时间戳
        if not export_path or not export_path.exists():
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            export_filename = f"{dossier.product_name}_申报资料_{timestamp}.docx"
            export_path = outputs_dir / export_filename
            working_files = list(working_dir.glob("*.docx"))
            if working_files:
                shutil.copy2(working_files[0], export_path)

        if export_path and export_path.exists():
            return {
                "success": True,
                "message": "导出成功",
                "file_path": str(export_path),
                "filename": export_filename,
            }
        else:
            return {"success": False, "message": "导出失败：无可用文件"}

    # ====== Storage Helpers ======

    def _create_storage_dirs(self, dossier_id: str) -> Dict[str, str]:
        """创建存储目录"""
        base = self.storage_root / "products" / dossier_id
        dirs = {
            "source_templates": base / "source_templates",
            "working": base / "working",
            "assets": base / "assets",
            "outputs": base / "outputs",
        }
        for d in dirs.values():
            d.mkdir(parents=True, exist_ok=True)
        return {k: str(v) for k, v in dirs.items()}

    def _delete_storage_dirs(self, dossier_id: str) -> None:
        """删除存储目录"""
        base = self.storage_root / "products" / dossier_id
        if base.exists():
            shutil.rmtree(base)

    async def _create_m3_chapters(self, dossier_id: UUID) -> None:
        """创建 M3 标准目录结构，按 M3_CHAPTERS 定义顺序赋值 sort_order"""
        code_to_id: Dict[str, UUID] = {}
        
        # 两遍遍历：先创建所有节点，再按 M3 标准顺序赋值 sort_order
        # 第一遍：创建所有章节记录（按原始列表顺序，保证父节点先于子节点）
        for idx, ch in enumerate(M3_CHAPTERS):
            parent_id = code_to_id.get(ch["parent_code"]) if ch["parent_code"] else None
            
            chapter = DossierChapter(
                product_dossier_id=dossier_id,
                parent_id=parent_id,
                chapter_code=ch["code"],
                chapter_title=ch["title"],
                level=ch["level"],
                sort_order=idx,
                has_content=False,
                has_assets=False,
            )
            chapter = await self.repo.create_chapter(chapter)
            code_to_id[ch["code"]] = chapter.id

    async def _create_working_copy_for_chapter(
        self, dossier: ProductDossier, template: DossierTemplate, chapter: DossierChapter
    ) -> None:
        """为章节创建 working copy"""
        source_path = Path(template.file_path)
        working_dir = Path(dossier.working_path)
        working_dir.mkdir(parents=True, exist_ok=True)

        # 生成 working copy 文件名
        working_filename = f"{chapter.chapter_code.replace('.', '_')}_{source_path.name}"
        working_path = working_dir / working_filename

        # 复制文件
        shutil.copy2(source_path, working_path)

        # 执行基础信息替换
        self._replace_basic_info(working_path, dossier)

        # 更新章节信息
        await self.repo.update_chapter(
            chapter.id,
            source_file=template.original_filename,
            working_file=working_filename,
            has_content=True,
        )


    async def get_chapter_preview(self, chapter_id: UUID) -> Dict[str, Any]:
        """获取章节预览内容（从 working copy 提取）"""
        chapter = await self.repo.get_chapter(chapter_id)
        if not chapter:
            return {"success": False, "message": "章节不存在"}
        
        if not chapter.working_file:
            return {"success": False, "message": "章节无工作副本"}
        
        dossier = await self.repo.get_product_dossier(chapter.product_dossier_id)
        if not dossier:
            return {"success": False, "message": "品种不存在"}
        
        working_path = Path(dossier.working_path) / chapter.working_file
        if not working_path.exists():
            return {"success": False, "message": "工作副本文件不存在"}
        
        try:
            doc = Document(str(working_path))
            
            # 提取文本内容
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal",
                    })
            
            # 提取表格
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                tables.append(rows)
            
            return {
                "success": True,
                "chapter_code": chapter.chapter_code,
                "chapter_title": chapter.chapter_title,
                "paragraphs": paragraphs,
                "tables": tables,
            }
        except Exception as e:
            return {"success": False, "message": f"预览失败: {str(e)}"}

    async def match_assets_to_chapters(self, dossier_id: UUID) -> Dict[str, Any]:
        """智能匹配素材到章节"""
        dossier = await self.repo.get_product_dossier(dossier_id)
        if not dossier:
            return {"success": False, "message": "品种不存在", "matched_count": 0, "unmatched_files": []}
        
        # 获取所有模板文件
        templates = await self.repo.list_templates(dossier_id)
        if not templates:
            return {"success": False, "message": "无模板文件", "matched_count": 0, "unmatched_files": []}
        
        # 获取所有章节（扁平列表）
        chapters = await self.repo.get_chapter_tree(dossier_id)
        
        matched_count = 0
        unmatched_files = []
        
        for template in templates:
            filename = template.original_filename
            
            # 尝试匹配章节
            matched_code = match_file_to_chapter(filename)
            
            if matched_code:
                # 找到对应章节
                matched_chapter = None
                for ch in chapters:
                    if ch.chapter_code == matched_code:
                        matched_chapter = ch
                        break
                
                if matched_chapter:
                    # 创建 working copy
                    await self._create_working_copy_for_chapter(dossier, template, matched_chapter)
                    matched_count += 1
            else:
                unmatched_files.append(filename)
        
        await self.db.commit()
        
        return {
            "success": True,
            "message": f"匹配完成：{matched_count} 个文件已匹配，{len(unmatched_files)} 个未匹配",
            "matched_count": matched_count,
            "unmatched_files": unmatched_files,
        }

    async def init_chapter_ai_config(self, chapter_code: str) -> Dict[str, Any]:
        """初始化章节的 AI 配置（FieldMapping + AssetCategory）
        
        从 scripts/seed_s6_ai_config.py 中的种子数据初始化，不再使用硬编码配置。
        如果该章节已有配置，跳过不重复创建。
        """
        from .field_models import FieldMapping, AssetCategory
        
        # 加载种子数据
        from scripts.seed_s6_ai_config import S6_FIELD_MAPPINGS, S6_ASSET_CATEGORIES

        
        # 只处理指定章节的配置
        field_configs = [c for c in S6_FIELD_MAPPINGS if c.get("chapter_code") == chapter_code]
        category_configs = [c for c in S6_ASSET_CATEGORIES if c.get("chapter_code") == chapter_code]
        
        if not field_configs and not category_configs:
            return {
                "success": False,
                "message": f"章节 {chapter_code} 没有可用的种子配置"
            }
        
        created_mappings = 0
        skipped_mappings = 0
        
        # 创建 FieldMapping
        for config in field_configs:
            stmt = select(FieldMapping).where(
                FieldMapping.chapter_code == config["chapter_code"],
                FieldMapping.field_name == config["field_name"],
                FieldMapping.is_deleted == False
            )
            result = await self.db.execute(stmt)
            existing = result.scalar_one_or_none()
            
            if existing:
                skipped_mappings += 1
                continue
            
            mapping = FieldMapping(**config)
            self.db.add(mapping)
            created_mappings += 1
        
        created_categories = 0
        skipped_categories = 0
        
        # 创建 AssetCategory
        for config in category_configs:
            stmt = select(AssetCategory).where(
                AssetCategory.chapter_code == config["chapter_code"],
                AssetCategory.category_name == config["category_name"],
                AssetCategory.is_deleted == False
            )
            result = await self.db.execute(stmt)
            existing = result.scalar_one_or_none()
            
            if existing:
                skipped_categories += 1
                continue
            
            category = AssetCategory(**config)
            self.db.add(category)
            created_categories += 1
        
        await self.db.commit()
        
        return {
            "success": True,
            "message": f"{chapter_code} 配置初始化完成: 字段映射 {created_mappings} 新建/{skipped_mappings} 跳过, 素材分类 {created_categories} 新建/{skipped_categories} 跳过",
            "chapter_code": chapter_code,
            "field_mappings_created": created_mappings,
            "field_mappings_skipped": skipped_mappings,
            "asset_categories_created": created_categories,
            "asset_categories_skipped": skipped_categories,
        }
