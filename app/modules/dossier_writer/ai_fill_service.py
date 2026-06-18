"""AI 填充服务 - 编排素材提取、AI 解析、模板填充的完整流程"""
import json
import subprocess
import tempfile
import uuid
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any

from docx import Document
_logger = logging.getLogger(__name__)
from sqlalchemy import select, and_
from sqlalchemy.ext.asyncio import AsyncSession

from .field_models import FieldMapping, FieldFillResult, AssetCategory, AssetPageSplit
from .models import ChapterAsset, DossierChapter, ProductDossier
from app.core.llm import llm_client
from .ai_prompts import (
    build_extract_fields_prompt,
    build_split_pages_prompt,
    build_fill_location_prompt,
)
from .asset_text_extractor import AssetTextExtractor


class AIFillService:
    """AI 填充服务"""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.llm = llm_client
        self.extractor = AssetTextExtractor()

    async def get_asset_categories(self, chapter_code: str) -> List[Dict]:
        """获取章节的素材分类列表"""
        stmt = select(AssetCategory).where(
            and_(
                AssetCategory.chapter_code == chapter_code,
                AssetCategory.is_deleted == False,
            )
        ).order_by(AssetCategory.sort_order)
        result = await self.db.execute(stmt)
        categories = list(result.scalars().all())
        return [
            {
                "id": str(c.id),
                "category_name": c.category_name,
                "category_type": c.category_type,
                "appendix_slot": c.appendix_slot,
                "description": c.description,
                "sort_order": c.sort_order,
            }
            for c in categories
        ]

    async def get_field_mappings(self, chapter_code: str) -> List[FieldMapping]:
        """获取章节的字段映射配置"""
        stmt = select(FieldMapping).where(
            and_(
                FieldMapping.chapter_code == chapter_code,
                FieldMapping.is_deleted == False,
            )
        ).order_by(FieldMapping.sort_order)
        result = await self.db.execute(stmt)
        return list(result.scalars().all())

    async def get_chapter_assets(self, chapter_id: uuid.UUID) -> List[ChapterAsset]:
        """获取章节的素材列表（附带分类名称）"""
        from .field_models import AssetCategory
        
        stmt = (
            select(ChapterAsset, AssetCategory.category_name)
            .outerjoin(AssetCategory, ChapterAsset.category_id == AssetCategory.id)
            .where(
                and_(
                    ChapterAsset.chapter_id == chapter_id,
                    ChapterAsset.is_deleted == False,
                )
            )
        )
        result = await self.db.execute(stmt)
        assets = []
        for row in result.all():
            asset = row[0]
            category_name = row[1]
            asset._category_name = category_name
            assets.append(asset)
        return assets

    async def preview_extraction(
        self,
        dossier: ProductDossier,
        chapter: DossierChapter,
    ) -> Dict[str, Any]:
        """预览 AI 提取结果（不写入文档，只返回提取值供用户确认）"""
        if False:  # Config check handled by core.llm
            return {"success": False, "message": "LLM 服务未配置"}

        mappings = await self.get_field_mappings(chapter.chapter_code)
        if not mappings:
            return {"success": False, "message": f"章节 {chapter.chapter_code} 没有配置字段映射"}

        assets = await self.get_chapter_assets(chapter.id)
        if not assets:
            return {"success": False, "message": "请先上传素材"}

        # 按 source_category 分组字段和素材
        extract_fields = [m for m in mappings if m.source_type == "asset_extract"]
        fixed_fields = [m for m in mappings if m.source_type == "fixed"]
        image_fields = [m for m in mappings if m.source_type == "asset_image"]

        results = []

        # 处理固定值字段
        for m in fixed_fields:
            results.append({
                "field_name": m.field_name,
                "field_type": m.field_type,
                "value": m.fixed_value,
                "confidence": 1.0,
                "source": "固定值",
                "field_mapping_id": str(m.id),
            })

        # 按 source_category 分组提取
        category_groups: Dict[str, List[FieldMapping]] = {}
        for m in extract_fields:
            cat = m.source_category or "_default"
            category_groups.setdefault(cat, []).append(m)

        # 对每个分类组调用 AI 提取
        total_usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}

        for category_name, group_mappings in category_groups.items():
            # 找到该分类下的素材文件
            category_assets = self._filter_assets_by_category(assets, category_name)
            if not category_assets:
                # 该分类下没有素材，所有字段标记为未找到
                for m in group_mappings:
                    results.append({
                        "field_name": m.field_name,
                        "field_type": m.field_type,
                        "value": None,
                        "confidence": 0.0,
                        "source": f"分类 '{category_name}' 下无素材",
                        "field_mapping_id": str(m.id),
                        "source_category": m.source_category,
                    })
                continue

            # 处理 table 类型字段：直接从素材文档中提取表格数据
            table_fields = [m for m in group_mappings if m.field_type == "table"]
            non_table_fields = [m for m in group_mappings if m.field_type != "table"]
            
            for m in table_fields:
                table_data = None
                for asset in category_assets:
                    table_data = self._extract_table_from_asset(asset)
                    if table_data:
                        results.append({
                            "field_name": m.field_name,
                            "field_type": m.field_type,
                            "value": table_data,
                            "confidence": 1.0,
                            "source": f"直接提取自: {asset.original_filename}",
                            "field_mapping_id": str(m.id),
                            "source_category": m.source_category,
                        })
                        break
                
                if not table_data:
                    results.append({
                        "field_name": m.field_name,
                        "field_type": m.field_type,
                        "value": None,
                        "confidence": 0.0,
                        "source": "未能从素材中提取表格数据",
                        "field_mapping_id": str(m.id),
                        "source_category": m.source_category,
                    })
            
            if not non_table_fields:
                continue

            # 提取素材文本
            asset_texts = {}
            for asset in category_assets:
                file_path = Path(asset.file_path)
                extracted = self.extractor.extract(file_path)
                if extracted.get("text"):
                    asset_texts[asset.original_filename] = extracted["text"]

            if not asset_texts:
                for m in non_table_fields:
                    results.append({
                        "field_name": m.field_name,
                        "field_type": m.field_type,
                        "value": None,
                        "confidence": 0.0,
                        "source": "素材文本提取失败",
                        "field_mapping_id": str(m.id),
                        "source_category": m.source_category,
                    })
                continue

            # 调用 AI 提取非表格字段
            fields_for_prompt = [
                {
                    "field_name": m.field_name,
                    "field_type": m.field_type,
                    "extraction_prompt": m.extraction_prompt or f"提取 {m.field_name} 的值",
                }
                for m in non_table_fields
            ]

            messages = build_extract_fields_prompt(
                fields=fields_for_prompt,
                asset_texts=asset_texts,
                product_name=dossier.product_name,
            )

            llm_result = await self.llm.chat_json(messages)
            if not llm_result["success"]:
                for m in non_table_fields:
                    results.append({
                        "field_name": m.field_name,
                        "field_type": m.field_type,
                        "value": None,
                        "confidence": 0.0,
                        "source": f"AI 提取失败: {llm_result.get('error', '')}",
                        "field_mapping_id": str(m.id),
                        "source_category": m.source_category,
                    })
                continue

            # 记录 token 用量
            usage = llm_result.get("usage", {})
            for k in total_usage:
                total_usage[k] += usage.get(k, 0)

            # 解析 AI 返回的字段
            parsed_fields = {
                f["field_name"]: f
                for f in llm_result.get("parsed", {}).get("fields", [])
            }

            for m in non_table_fields:
                extracted = parsed_fields.get(m.field_name, {})
                results.append({
                    "field_name": m.field_name,
                    "field_type": m.field_type,
                    "value": extracted.get("value"),
                    "confidence": extracted.get("confidence", 0.5),
                    "source": extracted.get("source", "未知来源"),
                    "field_mapping_id": str(m.id),
                    "source_category": m.source_category,
                })

        # 处理图片类字段
        for m in image_fields:
            results.append({
                "field_name": m.field_name,
                "field_type": m.field_type,
                "value": m.appendix_slot or "待插入",
                "confidence": 1.0,
                "source": "图片类字段，确认后插入",
                "field_mapping_id": str(m.id),
                "source_category": m.source_category,
            })

        return {
            "success": True,
            "message": f"提取完成: {len(results)} 个字段",
            "fields": results,
            "token_usage": total_usage,
        }

    async def confirm_and_fill(
        self,
        dossier: ProductDossier,
        chapter: DossierChapter,
        user_confirmed_fields: List[Dict[str, Any]],
    ) -> Dict[str, Any]:
        """用户确认后，将字段值写入文档"""
        working_path = Path(dossier.working_path) / chapter.working_file
        if not working_path.exists():
            return {"success": False, "message": f"工作副本不存在: {chapter.working_file}"}

        doc = Document(str(working_path))

        # 获取模板结构（段落和表格）
        template_paragraphs = [
            {"index": i, "text": p.text.strip()}
            for i, p in enumerate(doc.paragraphs)
            if p.text.strip()
        ]
        template_tables = []
        for t_idx, table in enumerate(doc.tables):
            rows_data = [
                [cell.text.strip()[:30] for cell in row.cells]
                for row in table.rows[:5]
            ]
            template_tables.append({
                "index": t_idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "preview": rows_data,
            })

        # 调用 AI 确定每个字段的填充位置
        text_fields = [f for f in user_confirmed_fields if f.get("value") is not None and f.get("field_type") != "image_appendix"]

        fill_instructions = []
        if text_fields:
            _logger.info(f"[Fill] Template has {len(template_paragraphs)} paragraphs, {len(template_tables)} tables")
            for p in template_paragraphs[:30]:
                _logger.info(f"  P[{p['index']}]: {p['text'][:80]}")
            for t in template_tables:
                _logger.info(f"  T[{t['index']}]: {t['rows']}x{t['cols']} preview={t['preview'][:2]}")
            _logger.info(f"[Fill] Fields to fill: {[f['field_name'] for f in text_fields]}")
            
            messages = build_fill_location_prompt(
                template_paragraphs=template_paragraphs,
                template_tables=template_tables,
                extracted_fields=text_fields,
            )
            llm_result = await self.llm.chat_json(messages)
            if llm_result["success"]:
                fill_instructions = llm_result.get("parsed", {}).get("fills", [])
                _logger.info(f"[Fill] AI returned {len(fill_instructions)} fill instructions:")
                for inst in fill_instructions:
                    _logger.info(f"  {inst.get('field_name')}: action={inst.get('fill_action')} target={inst.get('target')}")
            else:
                _logger.warning(f"[Fill] AI location prompt failed: {llm_result.get('message')}")

        # 加载素材（带分类名称），用于图片插入
        chapter_assets = await self.get_chapter_assets(chapter.id)

        # 执行填充
        fill_results = []
        for field_data in user_confirmed_fields:
            field_name = field_data["field_name"]
            value = field_data.get("value")
            field_type = field_data.get("field_type", "text")

            if field_type == "image_appendix":
                # 图片类字段：检查是否已手动插入
                if value and str(value).startswith("已插入"):
                    fill_results.append({
                        "field_name": field_name,
                        "status": "filled",
                        "message": "已通过选择页手动插入",
                    })
                else:
                    # 未手动插入，尝试自动插入
                    success = await self._auto_insert_image(doc, field_name, field_data, chapter, chapter_assets)
                    fill_results.append({
                        "field_name": field_name,
                        "status": "filled" if success else "skipped",
                        "message": "图片已自动插入" if success else "未插入（请通过选择页手动插入）",
                    })
                continue

            if value is None:
                fill_results.append({
                    "field_name": field_name,
                    "status": "skipped",
                    "message": "值为空，跳过",
                })
                continue

            # 找到对应的填充指令
            instruction = None
            for inst in fill_instructions:
                if inst.get("field_name") == field_name:
                    instruction = inst
                    break

            if instruction and instruction.get("fill_action") != "skip":
                success = self._execute_fill(doc, instruction, value)
                _logger.info(f"[Fill] {field_name}: action={instruction.get('fill_action')} result={'OK' if success else 'FAIL'}")
                fill_results.append({
                    "field_name": field_name,
                    "status": "filled" if success else "failed",
                    "message": "填充成功" if success else "填充执行失败",
                })
            else:
                # 没有填充指令，尝试 fallback 策略
                _logger.info(f"[Fill] {field_name}: no instruction, trying fallback")
                success = self._fallback_fill(doc, field_name, value, field_type)
                _logger.info(f"[Fill] {field_name}: fallback result={'OK' if success else 'FAIL'}")
                fill_results.append({
                    "field_name": field_name,
                    "status": "filled" if success else "no_match",
                    "message": "填充成功(fallback)" if success else "模板中未找到匹配位置",
                })

        # 保存文档
        doc.save(str(working_path))

        # 保存填充结果到数据库
        mappings = await self.get_field_mappings(chapter.chapter_code)
        mapping_map = {str(m.id): m for m in mappings}

        for field_data in user_confirmed_fields:
            mapping_id = field_data.get("field_mapping_id")
            if not mapping_id or mapping_id not in mapping_map:
                continue

            result = FieldFillResult(
                dossier_id=dossier.id,
                chapter_id=chapter.id,
                field_mapping_id=uuid.UUID(mapping_id),
                field_name=field_data["field_name"],
                filled_value=str(field_data.get("value", "")) if field_data.get("value") is not None else None,
                fill_method="ai",
                confidence=field_data.get("confidence"),
                status="filled" if field_data.get("value") is not None else "pending",
            )
            self.db.add(result)

        await self.db.commit()

        filled_count = sum(1 for r in fill_results if r["status"] == "filled")
        return {
            "success": True,
            "message": f"填充完成: {filled_count}/{len(user_confirmed_fields)} 个字段",
            "results": fill_results,
        }

    async def preview_page_splits(
        self,
        asset: ChapterAsset,
        available_appendix_slots: List[str],
    ) -> Dict[str, Any]:
        """预览多页 PDF 的拆分结果"""
        if False:  # Config check handled by core.llm
            return {"success": False, "message": "LLM 服务未配置"}

        file_path = Path(asset.file_path)
        if file_path.suffix.lower() != ".pdf":
            return {"success": False, "message": "仅支持 PDF 文件的页拆分"}

        extracted = self.extractor.extract(file_path)
        if not extracted.get("page_texts"):
            return {"success": False, "message": f"PDF 提取失败: {extracted.get('error', '')}"}

        messages = build_split_pages_prompt(
            page_texts=extracted["page_texts"],
            available_appendix_slots=available_appendix_slots,
        )

        llm_result = await self.llm.chat_json(messages)
        if not llm_result["success"]:
            return {"success": False, "message": f"AI 拆分失败: {llm_result.get('error', '')}"}

        pages = llm_result.get("parsed", {}).get("pages", [])

        # 保存拆分结果到数据库
        for page_info in pages:
            split = AssetPageSplit(
                asset_id=asset.id,
                page_number=page_info.get("page_number", 1),
                page_type=page_info.get("page_type", "未知"),
                content_summary=page_info.get("content_summary", ""),
                ocr_text=extracted["page_texts"][page_info.get("page_number", 1) - 1]["text"] if page_info.get("page_number") else "",
                appendix_slot=page_info.get("appendix_slot"),
                status="pending",
            )
            self.db.add(split)

        await self.db.commit()

        return {
            "success": True,
            "message": f"拆分完成: {len(pages)} 页",
            "pages": pages,
            "page_count": extracted.get("page_count", 0),
        }

    async def confirm_page_splits_and_insert(
        self,
        dossier: ProductDossier,
        chapter: DossierChapter,
        splits: List[Dict[str, Any]],
    ) -> Dict[str, Any]:
        """用户确认页拆分后，将各页转为图片插入模板"""
        working_path = Path(dossier.working_path) / chapter.working_file
        if not working_path.exists():
            return {"success": False, "message": "工作副本不存在"}

        doc = Document(str(working_path))
        inserted = 0

        for split_data in splits:
            split_id = split_data.get("split_id")
            appendix_slot = split_data.get("appendix_slot")
            asset_id = split_data.get("asset_id")
            page_number = split_data.get("page_number")

            if not appendix_slot or not asset_id:
                continue

            # 获取素材文件路径
            asset = await self.db.get(ChapterAsset, uuid.UUID(asset_id))
            if not asset:
                continue

            file_path = Path(asset.file_path)

            # 转换指定页为图片
            img_path = self.extractor.pdf_page_to_image(file_path, page_number)
            if not img_path:
                continue

            # 在模板中找到附录位置并插入图片
            success = self._insert_image_at_appendix(doc, appendix_slot, img_path)
            if success:
                inserted += 1

            # 更新拆分记录（仅在 split_id 是有效 UUID 时）
            if split_id:
                try:
                    split_uuid = uuid.UUID(split_id)
                    split_record = await self.db.get(AssetPageSplit, split_uuid)
                    if split_record:
                        split_record.appendix_slot = appendix_slot
                        split_record.image_path = str(img_path)
                        split_record.status = "inserted" if success else "failed"
                except ValueError:
                    pass  # split_id 不是有效 UUID，跳过记录更新

        doc.save(str(working_path))
        await self.db.commit()

        return {
            "success": True,
            "message": f"图片插入完成: {inserted}/{len(splits)}",
            "inserted_count": inserted,
        }

    def _filter_assets_by_category(
        self,
        assets: List[ChapterAsset],
        category_name: str,
    ) -> List[ChapterAsset]:
        """按素材分类 ID 精确过滤（基于用户确认的 category_id）"""
        if category_name == "_default":
            return assets

        # 精确匹配 category_id 对应的分类名称
        matched = []
        uncategorized = []
        for asset in assets:
            cat_name = getattr(asset, '_category_name', None)
            if cat_name == category_name:
                matched.append(asset)
            elif cat_name is None:
                uncategorized.append(asset)

        # 如果有精确匹配，返回精确匹配
        if matched:
            return matched

        # 如果没有任何素材被分类，返回空列表（让用户先分类）
        if all(getattr(a, '_category_name', None) is None for a in assets):
            return assets  # 全部未分类时回退到全部（兼容旧数据）

        return matched

    def _resolve_docx_path(self, file_path: Path) -> Optional[Path]:
        """将 .doc 文件解析为可用的 .docx 路径（优先使用已转换版本，否则调用 libreoffice）"""
        if file_path.suffix.lower() == ".docx":
            return file_path if file_path.exists() else None

        if file_path.suffix.lower() == ".doc":
            # 优先查找同目录已转换的 .docx
            docx_sibling = file_path.with_suffix(".docx")
            if docx_sibling.exists():
                return docx_sibling

            # 调用 libreoffice 转换
            try:
                with tempfile.TemporaryDirectory() as tmpdir:
                    result = subprocess.run(
                        ["libreoffice", "--headless", "--convert-to", "docx",
                         "--outdir", tmpdir, str(file_path)],
                        capture_output=True, text=True, timeout=30,
                    )
                    if result.returncode == 0:
                        converted = Path(tmpdir) / (file_path.stem + ".docx")
                        if converted.exists():
                            return converted
            except Exception:
                pass

        return None

    def _extract_table_from_asset(self, asset: ChapterAsset) -> Optional[List[List[str]]]:
        """直接从素材文档中提取检验项目表格数据（跳过 AI）"""
        file_path = Path(asset.file_path)
        docx_path = self._resolve_docx_path(file_path)
        if docx_path is None:
            return None

        try:
            doc = Document(str(docx_path))
            
            # Find the table with "检验项目" or "企业内控标准" header
            for table in doc.tables:
                if len(table.rows) < 3:
                    continue
                
                # Check all rows for header keywords
                header_row_idx = None
                for idx, row in enumerate(table.rows):
                    row_text = " ".join(cell.text.strip() for cell in row.cells)
                    if "检验项目" in row_text and "企业内控标准" in row_text:
                        header_row_idx = idx
                        break
                
                if header_row_idx is None:
                    continue
                
                # Found the table - extract data rows after header
                data_rows = []
                for row in table.rows[header_row_idx + 1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    # Deduplicate merged cell values
                    deduped = []
                    for j, cell_val in enumerate(cells):
                        if j == 0 or cell_val != cells[j-1]:
                            deduped.append(cell_val)
                    # Skip empty rows and "备注" rows
                    if not deduped or not any(v for v in deduped):
                        continue
                    if deduped[0].startswith("备注"):
                        continue
                    data_rows.append(deduped)
                
                return data_rows if data_rows else None
        except Exception:
            pass
        
        return None

    def _execute_fill(self, doc: Document, instruction: Dict, value: Any) -> bool:
        """执行单个字段的文档填充"""
        action = instruction.get("fill_action")
        target = instruction.get("target", {})

        try:
            if action == "replace_after_colon":
                return self._fill_paragraph_replace(doc, target, value)
            elif action == "fill_table_cell":
                return self._fill_table_cell(doc, target, value)
            elif action == "replace_table_rows":
                return self._fill_table_rows(doc, target, value)
            else:
                return False
        except Exception:
            return False

    def _fill_paragraph_replace(self, doc: Document, target: Dict, value: str) -> bool:
        """替换段落中冒号后的内容"""
        para_idx = target.get("paragraph_index")
        keyword = target.get("keyword", "")

        # 优先按 index 定位
        if para_idx is not None and para_idx < len(doc.paragraphs):
            para = doc.paragraphs[para_idx]
            text = para.text
            colon_pos = text.find("：")
            if colon_pos == -1:
                colon_pos = text.find(":")
            if colon_pos != -1:
                new_text = text[: colon_pos + 1] + str(value)
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                return True

        # fallback: 按关键词查找
        for para in doc.paragraphs:
            if keyword and keyword in para.text:
                text = para.text
                colon_pos = text.find("：")
                if colon_pos == -1:
                    colon_pos = text.find(":")
                if colon_pos != -1:
                    new_text = text[: colon_pos + 1] + str(value)
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_text
                    return True
        return False

    def _fill_table_cell(self, doc: Document, target: Dict, value: str) -> bool:
        """填充表格中关键词对应的单元格"""
        table_idx = target.get("table_index", 0)
        keyword = target.get("keyword", "")

        tables = doc.tables
        if table_idx >= len(tables):
            # fallback: 遍历所有表格
            for table in tables:
                if self._fill_in_table(table, keyword, str(value)):
                    return True
            return False

        table = tables[table_idx]
        return self._fill_in_table(table, keyword, str(value))

    def _fill_in_table(self, table, keyword: str, value: str) -> bool:
        """在表格中查找关键词并填充下一个单元格"""
        for row in table.rows:
            cells = list(row.cells)
            for i, cell in enumerate(cells):
                if keyword and keyword in cell.text and i + 1 < len(cells):
                    target_cell = cells[i + 1]
                    for para in target_cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                    if target_cell.paragraphs:
                        if target_cell.paragraphs[0].runs:
                            target_cell.paragraphs[0].runs[0].text = value
                        else:
                            target_cell.paragraphs[0].text = value
                    return True
        return False

    def _fill_table_rows(self, doc: Document, target: Dict, value: Any) -> bool:
        """替换表格数据行（用于完整表格字段）"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn
        from copy import deepcopy
        
        table_idx = target.get("table_index", 1)
        if table_idx >= len(doc.tables):
            return False

        if not isinstance(value, list) or len(value) == 0:
            return False

        table = doc.tables[table_idx]
        header_rows = target.get("header_rows", 2)
        
        if len(table.rows) <= header_rows:
            return False

        # Get the table XML element
        tbl_elem = table._tbl
        
        # Find the first data row to use as template (row after header)
        first_data_row = table.rows[header_rows]
        template_row = first_data_row._tr
        
        # Find footer row (last row, if it has merged cells or "备注")
        footer_row = None
        last_row = table.rows[-1]
        last_row_text = last_row.cells[0].text.strip()
        if "备注" in last_row_text or any(cell._tc.find(qn('w:tcPr')) is not None and 
                                           cell._tc.find(qn('w:tcPr')).find(qn('w:gridSpan')) is not None
                                           for cell in last_row.cells):
            footer_row = last_row._tr
        
        # Remove all data rows (between header and footer)
        rows_to_remove = table.rows[header_rows:-1] if footer_row is not None else table.rows[header_rows:]
        for row in rows_to_remove:
            tbl_elem.remove(row._tr)
        
        # Add new rows based on extracted data
        for idx, row_data in enumerate(value):
            if not isinstance(row_data, list):
                continue
            
            # Clone the template row
            new_tr = deepcopy(template_row)
            
            # Remove vMerge from all cells (each row is independent)
            cells = new_tr.findall(qn('w:tc'))
            for j, cell_elem in enumerate(cells):
                tcPr = cell_elem.find(qn('w:tcPr'))
                if tcPr is not None:
                    vmerge = tcPr.find(qn('w:vMerge'))
                    if vmerge is not None:
                        tcPr.remove(vmerge)
            
            # Fill in the cell values
            for j, cell_value in enumerate(row_data):
                if j < len(cells):
                    cell_elem = cells[j]
                    # Clear existing paragraphs
                    for p in cell_elem.findall(qn('w:p')):
                        cell_elem.remove(p)
                    
                    # Add new paragraph with text
                    new_p = parse_xml(
                        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        f'<w:r><w:t>{cell_value if cell_value else ""}</w:t></w:r></w:p>'
                    )
                    cell_elem.append(new_p)
            
            # Insert before footer or append to end
            if footer_row is not None:
                footer_row.addprevious(new_tr)
            else:
                tbl_elem.append(new_tr)
        
        return True

    def _fallback_fill(self, doc: Document, field_name: str, value: Any, field_type: str) -> bool:
        """Fallback 填充策略：当 AI 未返回填充指令时，按关键词在文档中查找"""
        str_value = str(value) if not isinstance(value, (list, dict)) else value

        if field_type == "table" and isinstance(value, list):
            # 尝试在所有表格中查找
            for table in doc.tables:
                for row in table.rows:
                    cells = list(row.cells)
                    for i, cell in enumerate(cells):
                        if field_name in cell.text and i + 1 < len(cells):
                            return self._fill_in_table(table, field_name, str_value)

        # 文本字段：在段落中查找
        for para in doc.paragraphs:
            if field_name in para.text:
                text = para.text
                colon_pos = text.find("：")
                if colon_pos == -1:
                    colon_pos = text.find(":")
                if colon_pos != -1:
                    new_text = text[: colon_pos + 1] + str_value
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_text
                    return True

        # 表格 fallback
        for table in doc.tables:
            if self._fill_in_table(table, field_name, str_value):
                return True

        return False

    async def _auto_insert_image(
        self,
        doc: Document,
        field_name: str,
        field_data: Dict[str, Any],
        chapter: DossierChapter,
        chapter_assets: List[ChapterAsset],
    ) -> bool:
        """自动从素材中提取图片并插入到文档的附录位置
        
        使用 FieldMapping.source_category 精确定位素材，不再依赖关键词硬编码匹配。
        """
        import logging
        _logger = logging.getLogger(__name__)
        from .field_models import FieldMapping
        
        # 1. 通过 field_mapping_id 查找 FieldMapping，获取 source_category
        mapping_id = field_data.get("field_mapping_id")
        source_category = None
        
        if mapping_id:
            mapping = await self.db.get(FieldMapping, uuid.UUID(mapping_id))
            if mapping:
                source_category = mapping.source_category
                _logger.info(f"[ImageInsert] {field_name}: field_mapping source_category={source_category}")
        
        # 2. 使用 _filter_assets_by_category 精确匹配素材
        target_asset = None
        if source_category:
            matched_assets = self._filter_assets_by_category(chapter_assets, source_category)
            if matched_assets:
                target_asset = matched_assets[0]
                _logger.info(f"[ImageInsert] {field_name}: matched asset via source_category: {target_asset.original_filename}")
        
        # 3. 如果 FieldMapping 匹配失败，回退到 AssetCategory.category_type 匹配
        if not target_asset:
            from .field_models import AssetCategory
            cat_stmt = select(AssetCategory).where(
                and_(
                    AssetCategory.chapter_code == chapter.chapter_code,
                    AssetCategory.category_type == "image_appendix",
                    AssetCategory.is_deleted == False,
                )
            )
            cat_result = await self.db.execute(cat_stmt)
            image_categories = list(cat_result.scalars().all())
            
            for cat in image_categories:
                matched = self._filter_assets_by_category(chapter_assets, cat.category_name)
                if matched:
                    # 检查附录编号是否匹配
                    appendix_slot = field_data.get("value", "")
                    if cat.appendix_slot and appendix_slot and cat.appendix_slot in appendix_slot:
                        target_asset = matched[0]
                        _logger.info(f"[ImageInsert] {field_name}: matched via appendix_slot: {target_asset.original_filename}")
                        break
            
            # 如果附录编号不匹配，取第一个图片类素材
            if not target_asset:
                for cat in image_categories:
                    matched = self._filter_assets_by_category(chapter_assets, cat.category_name)
                    if matched:
                        target_asset = matched[0]
                        _logger.info(f"[ImageInsert] {field_name}: fallback to image category asset: {target_asset.original_filename}")
                        break
        
        if not target_asset:
            _logger.warning(f"[ImageInsert] {field_name}: no matching asset found (source_category={source_category})")
            return False
        
        # 4. 转换素材的第一页为图片
        file_path = Path(target_asset.file_path)
        if not file_path.exists():
            _logger.warning(f"[ImageInsert] {field_name}: file not found {file_path}")
            return False
        
        img_path = self.extractor.pdf_page_to_image(file_path, 1)
        if not img_path:
            _logger.warning(f"[ImageInsert] {field_name}: failed to convert to image (only PDF supported)")
            return False
        
        # 5. 在文档中查找附录位置并插入图片
        appendix_slot = field_data.get("value", "")
        if not appendix_slot or appendix_slot == "待插入":
            appendix_slot = field_name.replace("图片", "")
        
        success = self._insert_image_at_appendix(doc, appendix_slot, img_path)
        if not success:
            success = self._insert_image_at_appendix(doc, field_name, img_path)
        
        _logger.info(f"[ImageInsert] {field_name}: insert result = {success}")
        return success

    def _insert_image_at_appendix(self, doc: Document, appendix_slot: str, img_path: Path) -> bool:
        """在模板的附录位置插入图片"""
        from docx.shared import Cm
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # Skip TOC entries (contain tabs and page numbers)
            if appendix_slot not in text or "\t" in para.text:
                continue
            # Skip if this is just the appendix slot name without full title
            if text == appendix_slot:
                continue
            
            # Found the appendix title in content area
            # Find the first empty paragraph after this title
            insert_idx = i + 1
            while insert_idx < len(doc.paragraphs):
                next_para = doc.paragraphs[insert_idx]
                if not next_para.text.strip():
                    # Found empty paragraph, insert image here
                    run = next_para.add_run()
                    run.add_picture(str(img_path), width=Cm(15))
                    return True
                # If we hit another non-empty paragraph (next appendix or section), stop
                if next_para.text.strip() and not next_para.text.strip().startswith(appendix_slot):
                    # Insert before this paragraph
                    new_para = doc.add_paragraph()
                    run = new_para.add_run()
                    run.add_picture(str(img_path), width=Cm(15))
                    # Move this new paragraph before the next appendix
                    next_para._element.addprevious(new_para._element)
                    return True
                insert_idx += 1
            
            # If no empty paragraph found, append at end
            new_para = doc.add_paragraph()
            run = new_para.add_run()
            run.add_picture(str(img_path), width=Cm(15))
            return True
        return False
