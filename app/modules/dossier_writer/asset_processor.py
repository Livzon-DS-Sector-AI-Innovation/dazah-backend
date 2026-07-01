"""素材处理器 - 从各种格式的素材中提取内容"""
import re
from pathlib import Path
from typing import Any

from docx import Document
from pdf2image import convert_from_path
from PIL import Image

from app.shared.ocr_service import get_ocr_service


class AssetExtractor:
    """素材内容提取器"""

    @staticmethod
    def extract_text_from_docx(file_path: Path) -> dict[str, Any]:
        """从 docx 文件提取文本和表格"""
        doc = Document(str(file_path))

        # 提取段落文本
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "index": i,
                    "text": text,
                    "style": para.style.name if para.style else "Normal"
                })

        # 提取表格
        tables = []
        for t_idx, table in enumerate(doc.tables):
            table_data = []
            for r_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                table_data.append(cells)
            tables.append({
                "index": t_idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": table_data
            })

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "full_text": "\n".join([p["text"] for p in paragraphs])
        }

    @staticmethod
    def extract_text_from_pdf_ocr(file_path: Path) -> dict[str, Any]:
        """使用 OCR 从 PDF 扫描件提取文本"""
        try:
            ocr_service = get_ocr_service()

            # 使用混合 API，PDF 自动使用 PP-StructureV3 保持结构
            full_text = ocr_service.extract(file_path, output_format="text")

            # 转换 PDF 为图片以获取页数
            images = convert_from_path(str(file_path), dpi=200)

            # 为每页提取文本（使用 PP-StructureV3）
            all_text = []
            for page_num, image in enumerate(images):
                text = ocr_service.extract(image, output_format="text")
                all_text.append({
                    "page": page_num + 1,
                    "text": text.strip()
                })

            return {
                "pages": all_text,
                "full_text": full_text,
                "page_count": len(images)
            }
        except Exception as e:
            return {
                "pages": [],
                "full_text": "",
                "error": str(e)
            }

    @staticmethod
    def extract_from_image(file_path: Path) -> dict[str, Any]:
        """从图片文件提取文本"""
        try:
            ocr_service = get_ocr_service()
            image = Image.open(str(file_path))
            # 使用 PP-OCR 进行快速文本提取
            text = ocr_service.extract_text(image)
            return {
                "text": text.strip(),
                "format": image.format,
                "size": image.size
            }
        except Exception as e:
            return {
                "text": "",
                "error": str(e)
            }

    @staticmethod
    def extract_field_value(content: str, field_name: str) -> str | None:
        """从内容中提取特定字段的值"""
        # 常见模式匹配
        patterns = [
            # "字段名：值" 或 "字段名: 值"
            rf'{field_name}[：:]\s*([^\n]+)',
            # "字段名  值" (多个空格)
            rf'{field_name}\s+([^\n]+)',
        ]

        for pattern in patterns:
            match = re.search(pattern, content)
            if match:
                return match.group(1).strip()

        return None

    @staticmethod
    def extract_table_value(tables: list[dict], field_name: str) -> str | None:
        """从表格中提取字段值"""
        for table in tables:
            for row in table["data"]:
                # 查找包含字段名的行
                for i, cell in enumerate(row):
                    if field_name in cell and i + 1 < len(row):
                        value = row[i + 1].strip()
                        if value:
                            return value
        return None


class ContentFiller:
    """内容填充器 - 将提取的值填充到模板中"""

    @staticmethod
    def fill_paragraph_text(doc: Document, keyword: str, value: str) -> bool:
        """在段落中查找关键词后的冒号位置并填充值"""
        for para in doc.paragraphs:
            if keyword in para.text:
                # 查找冒号位置
                text = para.text
                colon_pos = text.find('：')
                if colon_pos == -1:
                    colon_pos = text.find(':')

                if colon_pos != -1:
                    # 清空冒号后的内容并填入新值
                    new_text = text[:colon_pos + 1] + value
                    # 清空所有 run 并重建
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_text
                    return True
        return False

    @staticmethod
    def fill_table_cell(doc: Document, field_name: str, value: str) -> bool:
        """填充表格中指定字段的值"""
        for table in doc.tables:
            for row in table.rows:
                cells = list(row.cells)
                for i, cell in enumerate(cells):
                    if field_name in cell.text and i + 1 < len(cells):
                        # 填充下一列
                        target_cell = cells[i + 1]
                        # 清空现有内容
                        for para in target_cell.paragraphs:
                            for run in para.runs:
                                run.text = ""
                        # 填入新值
                        if target_cell.paragraphs:
                            if target_cell.paragraphs[0].runs:
                                target_cell.paragraphs[0].runs[0].text = value
                            else:
                                target_cell.paragraphs[0].text = value
                        return True
        return False

    @staticmethod
    def insert_image_at_placeholder(doc: Document, placeholder: str, image_path: Path) -> bool:
        """在占位符位置插入图片"""
        for para in doc.paragraphs:
            if placeholder in para.text:
                # 清空段落
                for run in para.runs:
                    run.text = ""
                # 插入图片
                if para.runs:
                    run = para.runs[0]
                    run.add_picture(str(image_path), width=500000)  # 宽度约 5cm
                    return True
        return False
