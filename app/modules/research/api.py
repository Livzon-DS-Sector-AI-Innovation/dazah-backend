"""研发项目 API 路由."""



import uuid

from uuid import UUID

from typing import Optional

from fastapi import APIRouter, Body, Depends, File, HTTPException, Query, UploadFile

from fastapi.responses import JSONResponse, StreamingResponse

from sqlalchemy.ext.asyncio import AsyncSession



from app.core.database import get_db

from app.core.deps import CurrentUser

from app.platform.identity.permissions import require_permission, require_role, RequireManager, RequireMember

from app.core.response import paginated_response, success_response, error_response

from app.modules.research import service

from app.modules.research.schemas import (

    ResearchProjectCreate,

    ResearchProjectResponse,

    ResearchProjectUpdate,

    RdProjectCreate, RdProjectUpdate, RdProjectResponse,

    RdMilestoneCreate, RdMilestoneUpdate, RdMilestoneResponse,

    RdStageRecordCreate, RdStageRecordUpdate, RdStageRecordResponse,

    RdResearchTrackCreate, RdResearchTrackUpdate, RdResearchTrackResponse,

    RdResearchFindingCreate, RdResearchFindingUpdate, RdResearchFindingResponse,

    RdPilotStudyCreate, RdPilotStudyUpdate, RdPilotStudyResponse,

    RdProcessValidationCreate, RdProcessValidationUpdate, RdProcessValidationResponse,

    RdRegistrationFilingCreate, RdRegistrationFilingUpdate, RdRegistrationFilingResponse,

    RdStageDeliverableCreate, RdStageDeliverableUpdate, RdStageDeliverableResponse,

    RdExperimentLogCreate, RdExperimentLogUpdate, RdExperimentLogResponse,

    RdReportCreate, RdReportUpdate, RdReportResponse,

    RdInitiationCreate, RdInitiationUpdate, RdInitiationResponse,

    RdDeliverableTemplateCreate, RdDeliverableTemplateUpdate, RdDeliverableTemplateResponse,

    RdReportGenerateRequest, RdReportGenerateResponse,

    RdTrackConclusionVersionCreate, RdTrackConclusionVersionResponse,

)

from app.shared.module_api import create_module_router

from app.shared.module_registry import MODULES_BY_CODE



router = create_module_router(MODULES_BY_CODE["research"])





@router.post("/projects", summary="创建研发项目")

async def create_project(

    current_user: CurrentUser,
    data: ResearchProjectCreate,

    db: AsyncSession = Depends(get_db),


) -> JSONResponse:

    project = await service.create_project(db, data)

    return success_response(data=ResearchProjectResponse.model_validate(project))





@router.get("/projects", summary="获取研发项目列表")

async def get_projects(

    current_user: CurrentUser,
    stage: str | None = Query(None, description="项目阶段"),

    status: str | None = Query(None, description="项目状态"),

    keyword: str | None = Query(None, description="搜索项目编号或名称"),

    project_type: str | None = Query(None, description="项目类型过滤"),

    page: int = Query(1, ge=1),

    page_size: int = Query(20, ge=1, le=100),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    projects, total = await service.get_projects(

        db, stage=stage, status=status, keyword=keyword, project_type=project_type,

        page=page, page_size=page_size

    )

    return paginated_response(

        data=[ResearchProjectResponse.model_validate(p) for p in projects],

        page=page,

        page_size=page_size,

        total=total,

    )





@router.get("/projects/{project_id}", summary="获取研发项目详情")

async def get_project(

    current_user: CurrentUser,
    project_id: uuid.UUID,

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    project = await service.get_project(db, project_id)

    return success_response(data=ResearchProjectResponse.model_validate(project))





@router.put("/projects/{project_id}", summary="更新研发项目")

async def update_project(

    project_id: uuid.UUID,

    data: ResearchProjectUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    project = await service.update_project(db, project_id, data)

    return success_response(data=ResearchProjectResponse.model_validate(project))





@router.delete("/projects/{project_id}", summary="删除研发项目")

async def delete_project(

    project_id: uuid.UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    await service.delete_project(db, project_id)

    return success_response(data={"message": "项目已删除"})





@router.post("/ich/q3c/analyze", summary="ICH Q3C 溶剂残留分析")

async def analyze_ich_q3c(

    file: UploadFile = File(...),

    route: str = Query("oral", description="给药途径"),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    file_content = await file.read()

    filename = file.filename or "unknown"

    result = await service.analyze_ich_q3c(db, file_content, filename, route)

    return success_response(data=result)







@router.post("/ich/analyze", summary="ICH Q3C/Q3D 联合分析")

async def analyze_ich_combined(

    file: UploadFile = File(...),

    route: str = Query("oral", description="给药途径"),

    use_llm: bool = Query(False, description="是否使用 LLM 增强"),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    file_content = await file.read()

    filename = file.filename or "unknown"

    result = await service.analyze_ich_combined(db, file_content, filename, route, use_llm)

    return success_response(data=result)





@router.get("/ich/records", summary="获取 ICH Q3C/Q3D 杂质识别记录列表")

async def get_ich_records(

    page: int = Query(1, ge=1),

    page_size: int = Query(20, ge=1, le=100),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    records, total = await service.get_ich_records(db, page=page, page_size=page_size)

    return paginated_response(

        data=[

            {

                "id": str(r.id),

                "filename": r.filename,

                "route": r.route,

                "q3c_result": r.q3c_result,

                "q3d_result": r.q3d_result,

                "llm_used": r.llm_used,

                "notes": r.notes,

                "created_at": r.created_at.isoformat(),

            }

            for r in records

        ],

        page=page,

        page_size=page_size,

        total=total,

    )





@router.get("/ich/records/{record_id}", summary="获取 ICH Q3C/Q3D 杂质识别记录详情")

async def get_ich_record(

    record_id: uuid.UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    record = await service.get_ich_record(db, record_id)

    return success_response(

        data={

            "id": str(record.id),

            "filename": record.filename,

            "route": record.route,

            "q3c_result": record.q3c_result,

            "q3d_result": record.q3d_result,

            "llm_used": record.llm_used,

            "notes": record.notes,

            "created_at": record.created_at.isoformat(),

        }

    )





@router.delete("/ich/records/{record_id}", summary="删除 ICH Q3C/Q3D 杂质识别记录")

async def delete_ich_record(

    record_id: uuid.UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    await service.delete_ich_record(db, record_id)

    return success_response(data={"message": "记录已删除"})





@router.post("/edbo/optimize", summary="EDBO+ 贝叶斯优化")

async def edbo_optimize(

    file: UploadFile = File(..., description="反应范围 CSV 文件"),

    objectives: str = Body(..., description="目标列名，逗号分隔"),

    objective_modes: str = Body("max", description="目标方向，逗号分隔（max/min）"),

    batch_size: int = Body(5, ge=1, le=100, description="建议实验数量"),

    save_prediction: bool = Body(False, description="是否保存预测文件"),

    current_user: CurrentUser = None,

) -> JSONResponse:

    """

    使用 EDBO+ 进行贝叶斯反应优化。



    上传 CSV 文件（反应范围），指定目标列和优化方向，返回建议的实验列表。

    如果 save_prediction=True，还会返回预测文件（包含所有实验的模型预测结果）。

    """

    from app.modules.research.edbo_runner import run_edbo_optimization

    from app.modules.research.schemas import EDBOOptimizeResponse



    # Parse comma-separated strings

    obj_list = [o.strip() for o in objectives.split(",") if o.strip()]

    mode_list = [m.strip() for m in objective_modes.split(",") if m.strip()]



    if len(obj_list) != len(mode_list):

        from fastapi import HTTPException

        raise HTTPException(

            status_code=400,

            detail=f"objectives ({len(obj_list)}) and objective_modes ({len(mode_list)}) must have the same length"

        )



    # Read CSV content

    try:

        csv_content = (await file.read()).decode("utf-8")

    except UnicodeDecodeError:

        from app.core.response import error_response

        return error_response(message="文件编码错误：请使用 UTF-8 编码的 CSV 文件")



    # Run EDBO+ optimization

    try:

        result = await run_edbo_optimization(

            csv_content=csv_content,

            objectives=obj_list,

            objective_modes=mode_list,

            batch_size=batch_size,

            save_prediction=save_prediction,

        )

    except RuntimeError as e:

        from app.core.response import error_response

        return error_response(message=str(e))



    return success_response(data=EDBOOptimizeResponse(**result))





@router.post("/edbo/generate-scope", summary="生成反应范围")

async def edbo_generate_scope(

    current_user: CurrentUser,
    components: dict = Body(..., description="组件定义，支持两种格式：1) 直接值列表 {name: [v1,v2,...]} 2) 范围定义 {name: {type: 'numeric', lower: x, upper: y, data_points: n}} 或 {name: {type: 'categorical', values: [...]}}"),

    objectives: list[str] = Body(default=[], description="优化目标名称列表，会作为新列添加到结果CSV中，初始值为PENDING"),

    batch_size: int = Body(default=5, ge=1, le=100, description="通量大小，表示同时进行的实验数量")

):

    """

    生成反应范围 CSV（所有组合的笛卡尔积）

    

    支持两种组件格式：

    

    1. 简单值列表（直接枚举）:

    ```json

    {

        "solvent": ["THF", "DMSO"],

        "catalyst": ["Pd", "Ni"]

    }

    ```

    

    2. 范围定义（数值型自动生成等间隔值）:

    ```json

    {

        "temperature": {"type": "numeric", "lower": 30, "upper": 90, "data_points": 4},

        "solvent": {"type": "categorical", "values": ["THF", "DMSO", "MeOH"]}

    }

    ```

    

    数值型组件会自动处理有效数字，避免浮点精度问题。

    """

    import itertools

    import math

    import pandas as pd

    from fastapi.responses import StreamingResponse

    import io

    import logging

    

    logger = logging.getLogger(__name__)

    logger.info(f"generate-scope called with objectives: {objectives}")

    logger.info(f"generate-scope components: {components}")

    logger.info(f"generate-scope batch_size: {batch_size}")

    

    def count_significant_digits(num):

        """Count the number of significant digits in a number."""

        if num == 0:

            return 1

        num_str = str(num).lower()

        if 'e' in num_str:

            mantissa = num_str.split('e')[0].replace('.', '').lstrip('0')

            return len(mantissa)

        if '.' in num_str:

            integer_part, decimal_part = num_str.split('.')

            integer_part = integer_part.lstrip('0')

            if integer_part:

                return len(integer_part) + len(decimal_part)

            else:

                decimal_part = decimal_part.lstrip('0')

                return len(decimal_part)

        else:

            return len(num_str.lstrip('0'))

    

    def round_to_significant_digits(num, sig_digits):

        """Round a number to the specified number of significant digits."""

        if num == 0:

            return 0

        if abs(num) == float('inf') or math.isnan(num):

            return num

        magnitude = 10 ** (sig_digits - 1 - int(math.floor(math.log10(abs(num)))))

        rounded = round(num * magnitude) / magnitude

        return rounded

    

    def generate_numeric_values(lower, upper, data_points):

        """Generate numeric values with proper significant digit handling."""

        if lower > upper:

            raise ValueError(f"下限 ({lower}) 不能大于上限 ({upper})")

        if data_points <= 1:

            raise ValueError(f"数据点数必须大于 1")

        

        values = []

        if data_points == 2:

            return [lower, upper]

        

        interval = (upper - lower) / (data_points - 1)

        

        # Check for extreme case: too many data points for small intervals

        if abs(lower) > 1e-15:

            relative_interval = interval / abs(lower)

            if relative_interval < 1e-6:

                raise ValueError(f"数据点过多，间隔过小")

        

        # Determine significant digits

        lower_sig_digits = count_significant_digits(lower)

        upper_sig_digits = count_significant_digits(upper)

        target_sig_digits = min(lower_sig_digits, upper_sig_digits)

        

        # Check if first in-between value rounds to lower limit

        first_in_between = lower + interval

        rounded_first = round_to_significant_digits(first_in_between, target_sig_digits)

        if abs(rounded_first - lower) < 1e-10:

            # Increase significant digits until different

            current_sig_digits = target_sig_digits + 1

            while current_sig_digits <= 15:

                rounded_first = round_to_significant_digits(first_in_between, current_sig_digits)

                if abs(rounded_first - lower) > 1e-10:

                    target_sig_digits = current_sig_digits

                    break

                current_sig_digits += 1

        

        # Generate all values

        for i in range(data_points):

            exact_value = lower + i * interval

            if abs(exact_value) == float('inf') or math.isnan(exact_value):

                values.append(exact_value)

            else:

                rounded_value = round_to_significant_digits(exact_value, target_sig_digits)

                values.append(rounded_value)

        

        return values

    

    # Validate input

    if not components:

        raise HTTPException(status_code=400, detail="组件定义不能为空")

    

    # Process each component

    processed_components = {}

    for key, value in components.items():

        if isinstance(value, list):

            # Simple list format

            processed_components[key] = value

        elif isinstance(value, dict):

            # Range definition format

            comp_type = value.get('type', 'categorical')

            if comp_type == 'numeric':

                lower = value.get('lower')

                upper = value.get('upper')

                data_points = value.get('data_points')

                if lower is None or upper is None or data_points is None:

                    raise HTTPException(

                        status_code=400,

                        detail=f"组件 '{key}' 缺少必要参数 (lower, upper, data_points)"

                    )

                try:

                    processed_components[key] = generate_numeric_values(

                        float(lower), float(upper), int(data_points)

                    )

                except ValueError as e:

                    raise HTTPException(status_code=400, detail=f"组件 '{key}': {str(e)}")

            elif comp_type == 'categorical':

                cat_values = value.get('values', [])

                if not cat_values:

                    raise HTTPException(

                        status_code=400,

                        detail=f"组件 '{key}' 的值列表不能为空"

                    )

                processed_components[key] = cat_values

            else:

                raise HTTPException(

                    status_code=400,

                    detail=f"组件 '{key}' 的类型无效: {comp_type}"

                )

        else:

            raise HTTPException(

                status_code=400,

                detail=f"组件 '{key}' 的格式无效"

            )

    

    # Calculate total combinations

    n_combinations = 1

    for key, values in processed_components.items():

        n_combinations *= len(values)

    

    if n_combinations > 10000:

        raise HTTPException(

            status_code=400,

            detail=f"组合数量过大 ({n_combinations})，请减少组件选项"

        )

    

    # Generate Cartesian product

    keys = list(processed_components.keys())

    values = [processed_components[key] for key in keys]

    

    scope = [dict(zip(keys, combination)) for combination in itertools.product(*values)]

    df_scope = pd.DataFrame(scope)

    

    # Convert to CSV (without objective columns - EDBO+ will add them)

    csv_buffer = io.StringIO()

    df_scope.to_csv(csv_buffer, index=False)

    csv_buffer.seek(0)

    scope_csv = csv_buffer.getvalue()

    

    # If objectives are defined, run EDBO+ to get initial sampling recommendations

    # IMPORTANT: Do NOT add objective columns with PENDING values before calling EDBO+.

    # EDBO+ detects missing objective columns and performs initial sampling automatically.

    # If we add PENDING columns first, EDBO+ thinks the scope already exists and refuses to run.

    logger.info(f"Checking objectives: {objectives}, length: {len(objectives)}")

    if objectives:

        logger.info(f"Objectives found, calling EDBO+ for initial sampling with: {objectives}")

        try:

            from app.modules.research.edbo_runner import run_edbo_optimization

            objective_modes = ['max'] * len(objectives)

            

            result = await run_edbo_optimization(

                csv_content=scope_csv,

                objectives=objectives,

                objective_modes=objective_modes,

                batch_size=batch_size,

                save_prediction=False,

            )

            

            result_csv = result.get("csv_data", scope_csv)

            return {

                "csv_data": result_csv,

                "row_count": len(scope),

                "columns": keys + objectives + ["priority"],

                "recommended_experiments": result_csv,

                "optimization_completed": True

            }

        except Exception as e:

            logger.error(f"EDBO+ optimization failed: {e}")

            # Add PENDING objective columns for the fallback response

            for obj in objectives:

                df_scope[obj] = "PENDING"

            csv_buffer = io.StringIO()

            df_scope.to_csv(csv_buffer, index=False)

            csv_buffer.seek(0)

            scope_csv_with_pending = csv_buffer.getvalue()

            

            return {

                "csv_data": scope_csv_with_pending,

                "row_count": len(scope),

                "columns": keys + objectives,

                "optimization_completed": False,

                "optimization_error": str(e)

            }

    

    return {

        "csv_data": scope_csv,

        "row_count": len(scope),

        "columns": keys,

        "optimization_completed": False

    }





# ===== Pilot Workflow API =====



import os

import uuid as uuid_module



from app.modules.research import repository as pilot_repo

from app.modules.research.pilot_workflow.engine import (

    start_workflow as start_workflow_engine,

    approve_step as approve_step_engine,

)

from app.modules.research.schemas import (

    PilotWorkflowCreate,

    PilotWorkflowListResponse,

    PilotWorkflowResponse,

    PilotWorkflowStepResponse,

)





@router.post("/pilot/workflow", summary="创建中试研究")

async def create_pilot_workflow(

    data: PilotWorkflowCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    workflow_data = data.model_dump()

    workflow = await pilot_repo.create_workflow(db, workflow_data)

    await db.flush()



    # 构建响应（不含步骤）

    response_data = PilotWorkflowResponse.model_validate(workflow)

    response_data.steps = []

    return success_response(data=response_data)





@router.get("/pilot/workflow", summary="获取中试研究列表")

async def get_pilot_workflows(

    current_user: CurrentUser,
    status: str | None = Query(None, description="状态筛选"),

    keyword: str | None = Query(None, description="搜索产品名称"),

    page: int = Query(1, ge=1),

    page_size: int = Query(20, ge=1, le=100),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    workflows, total = await pilot_repo.get_workflows(

        db, status=status, keyword=keyword, page=page, page_size=page_size

    )



    # 获取每个工作流的步骤数

    items = []

    for wf in workflows:

        steps = await pilot_repo.get_workflow_steps(db, wf.id)

        completed = sum(1 for s in steps if s.status == "completed")

        item = PilotWorkflowListResponse.model_validate(wf)

        item.step_count = len(steps)

        item.completed_step_count = completed

        items.append(item)



    return paginated_response(

        data=items,

        page=page,

        page_size=page_size,

        total=total,

    )





@router.get("/pilot/workflow/{workflow_id}", summary="获取工作流详情")

async def get_pilot_workflow_detail(

    current_user: CurrentUser,
    workflow_id: uuid_module.UUID,

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    workflow = await pilot_repo.get_workflow_by_id(db, workflow_id)

    if not workflow:

        raise HTTPException(status_code=404, detail="工作流不存在")



    steps = await pilot_repo.get_workflow_steps(db, workflow_id)

    response_data = PilotWorkflowResponse.model_validate(workflow)

    response_data.steps = [

        PilotWorkflowStepResponse.model_validate(s) for s in steps

    ]

    return success_response(data=response_data)





@router.post("/pilot/workflow/{workflow_id}/start", summary="启动工作流执行")

async def start_pilot_workflow(

    workflow_id: uuid_module.UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    workflow = await pilot_repo.get_workflow_by_id(db, workflow_id)

    if not workflow:

        raise HTTPException(status_code=404, detail="工作流不存在")



    if workflow.status != "pending":

        raise HTTPException(status_code=400, detail=f"工作流状态为 {workflow.status}，无法启动")



    # 异步启动工作流

    await start_workflow_engine(workflow_id)



    return success_response(data={"message": "工作流已启动", "workflow_id": str(workflow_id)})









@router.post("/pilot/workflow/{workflow_id}/approve", summary="确认当前步骤并执行下一步")

async def approve_pilot_workflow_step(

    workflow_id: uuid_module.UUID,

    current_user: CurrentUser,

) -> JSONResponse:

    result = await approve_step_engine(workflow_id)

    if "error" in result:

        raise HTTPException(status_code=400, detail=result["error"])

    return success_response(data=result)



@router.get("/pilot/workflow/{workflow_id}/steps/{step_id}", summary="获取步骤详情")

async def get_pilot_workflow_step(

    current_user: CurrentUser,
    workflow_id: uuid_module.UUID,

    step_id: uuid_module.UUID,

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    step = await pilot_repo.get_workflow_step_by_id(db, step_id)

    if not step or step.workflow_id != workflow_id:

        raise HTTPException(status_code=404, detail="步骤不存在")



    return success_response(

        data=PilotWorkflowStepResponse.model_validate(step)

    )





@router.post("/pilot/workflow/{workflow_id}/upload", summary="上传工艺文档")

async def upload_pilot_document(

    workflow_id: uuid_module.UUID,

    file: UploadFile = File(...),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    workflow = await pilot_repo.get_workflow_by_id(db, workflow_id)

    if not workflow:

        raise HTTPException(status_code=404, detail="工作流不存在")



    if workflow.status != "pending":

        raise HTTPException(status_code=400, detail="只有 pending 状态的工作流可以上传文档")



    # 保存文件

    upload_dir = "storage/pilot_workflow"

    os.makedirs(upload_dir, exist_ok=True)

    file_path = os.path.join(

        upload_dir, f"{workflow_id}_{file.filename}"

    )

    content = await file.read()

    with open(file_path, "wb") as f:

        f.write(content)



    # 更新工作流

    workflow.input_document_path = file_path

    await db.flush()



    return success_response(

        data={

            "file_path": file_path,

            "filename": file.filename,

            "file_size": len(content),

        }

    )





@router.delete("/pilot/workflow/{workflow_id}", summary="删除工作流")

async def delete_pilot_workflow(

    workflow_id: uuid_module.UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    workflow = await pilot_repo.get_workflow_by_id(db, workflow_id)

    if not workflow:

        raise HTTPException(status_code=404, detail="工作流不存在")



    # Soft delete

    workflow.is_deleted = True

    workflow.updated_by = current_user.id if current_user else None

    await db.flush()



    return success_response(data={"id": str(workflow_id)})



@router.post("/literature/analyze", summary="AI 文献解析（流式）")

async def analyze_literature(

    file: UploadFile = File(...),

    current_user: CurrentUser = None,

):

    """解析上传的文献文件（PDF/TXT），提取合成路线 - SSE 流式响应"""

    import json

    import asyncio

    from app.modules.research.literature_service import analyze_literature_with_ai

    

    # 读取文件内容

    content = await file.read()

    

    # 根据文件类型提取文本

    text = ""

    filename = file.filename.lower() if file.filename else ""

    

    async def event_stream():

        nonlocal text

        

        # 进度 10%: 开始解析文件

        yield f"data: {json.dumps({'progress': 10, 'status': 'reading', 'message': '正在读取文件...'})}\n\n"

        await asyncio.sleep(0.1)

        

        try:

            if filename.endswith('.txt'):

                text = content.decode('utf-8')

            elif filename.endswith('.md') or filename.endswith('.markdown'):

                text = content.decode('utf-8')

            elif filename.endswith('.pdf'):

                # 进度 20%: 解析 PDF

                yield f"data: {json.dumps({'progress': 20, 'status': 'extracting', 'message': '正在提取 PDF 文本...'})}\n\n"

                from pypdf import PdfReader

                import io

                reader = PdfReader(io.BytesIO(content))

                for page in reader.pages:

                    text += page.extract_text() + "\n"

            elif filename.endswith('.docx'):

                # 进度 20%: 解析 Word

                yield f"data: {json.dumps({'progress': 20, 'status': 'extracting', 'message': '正在提取 Word 文档...'})}\n\n"

                import docx

                import io

                doc = docx.Document(io.BytesIO(content))

                for para in doc.paragraphs:

                    text += para.text + "\n"

                for table in doc.tables:

                    for row in table.rows:

                        for cell in row.cells:

                            text += cell.text + " "

                        text += "\n"

            elif filename.endswith('.doc'):

                yield f"data: {json.dumps({'progress': 100, 'status': 'error', 'error': '不支持旧版 .doc 格式，请转换为 .docx 或 PDF 后上传'})}\n\n"

                return

            elif filename.endswith('.rtf'):

                text = content.decode('utf-8', errors='ignore')

                import re

                text = re.sub(r'\\[a-z]+[\d]*\s?', '', text)

                text = re.sub(r'[{}]', '', text)

            else:

                yield f"data: {json.dumps({'progress': 100, 'status': 'error', 'error': '不支持的文件格式，请上传 PDF、Word、TXT、Markdown 或 RTF 文件'})}\n\n"

                return

        except Exception as e:

            yield f"data: {json.dumps({'progress': 100, 'status': 'error', 'error': f'文件解析失败: {str(e)}'})}\n\n"

            return

        

        # 进度 30%: 文件提取完成

        yield f"data: {json.dumps({'progress': 30, 'status': 'extracted', 'message': f'文件提取完成，共 {len(text)} 字符'})}\n\n"

        await asyncio.sleep(0.1)

        

        # 进度 40%: 开始 AI 分析

        yield f"data: {json.dumps({'progress': 40, 'status': 'analyzing', 'message': 'AI 正在分析文献，提取合成路线...'})}\n\n"

        

        try:

            # 调用 AI 解析

            result = await analyze_literature_with_ai(text)

            

            # 进度 90%: AI 分析完成

            yield f"data: {json.dumps({'progress': 90, 'status': 'analyzed', 'message': 'AI 分析完成，正在整理结果...'})}\n\n"

            await asyncio.sleep(0.1)

            

            # 进度 100%: 完成

            yield f"data: {json.dumps({'progress': 100, 'status': 'complete', 'data': result})}\n\n"

        except Exception as e:

            yield f"data: {json.dumps({'progress': 100, 'status': 'error', 'error': f'AI 解析失败: {str(e)}'})}\n\n"

    

    return StreamingResponse(event_stream(), media_type="text/event-stream")





# ==================== 打通路线 CRUD ====================







@router.get("/routes", summary="获取打通路线列表")

async def get_routes(

    current_user: CurrentUser,
    project_id: str | None = Query(None),

    status: str | None = Query(None),

    keyword: str | None = Query(None),

    page: int = Query(1, ge=1),

    page_size: int = Query(20, ge=1, le=100),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    from app.modules.research.models import RouteDevelopment, RouteExperiment

    from sqlalchemy import select, func, or_

    

    query = select(RouteDevelopment).where(RouteDevelopment.is_deleted == False)

    count_query = select(func.count()).select_from(RouteDevelopment).where(RouteDevelopment.is_deleted == False)

    

    if project_id:

        query = query.where(RouteDevelopment.project_id == project_id)

        count_query = count_query.where(RouteDevelopment.project_id == project_id)

    if status:

        query = query.where(RouteDevelopment.status == status)

        count_query = count_query.where(RouteDevelopment.status == status)

    if keyword:

        pattern = f"%{keyword}%"

        like_filter = or_(

            RouteDevelopment.name.ilike(pattern),

            RouteDevelopment.route_no.ilike(pattern),

        )

        query = query.where(like_filter)

        count_query = count_query.where(like_filter)

    

    total = (await db.execute(count_query)).scalar_one()

    query = query.order_by(RouteDevelopment.updated_at.desc()).offset((page - 1) * page_size).limit(page_size)

    result = await db.execute(query)

    items = result.scalars().all()

    

    data = []

    for item in items:

        data.append({

            "id": str(item.id),

            "project_id": item.project_id,

            "route_no": item.route_no,

            "name": item.name,

            "source": item.source,

            "source_reference": item.source_reference,

            "description": item.description,

            "status": item.status,

            "current_module": item.current_module,

            "literature_sources": item.literature_sources or [],

            "candidate_routes": item.candidate_routes or [],

            "selected_route_ids": item.selected_route_ids or [],

            "experiment_plans": item.experiment_plans or [],

            "experiments": [],  # Will be loaded separately

            "assessment": item.assessment,

            "deliverables": item.deliverables or [],

            "start_date": item.start_date.isoformat() if item.start_date else None,

            "end_date": item.end_date.isoformat() if item.end_date else None,

            "created_at": item.created_at.isoformat() if item.created_at else None,

            "updated_at": item.updated_at.isoformat() if item.updated_at else None,

            "created_by": str(item.created_by) if item.created_by else None,

        })

    

    return paginated_response(data=data, page=page, page_size=page_size, total=total)





@router.get("/routes/{route_id}", summary="获取路线详情")

async def get_route(

    current_user: CurrentUser,
    route_id: str,

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    from app.modules.research.models import RouteDevelopment, RouteExperiment

    from sqlalchemy import select

    

    result = await db.execute(

        select(RouteDevelopment).where(

            RouteDevelopment.id == route_id,

            RouteDevelopment.is_deleted == False,

        )

    )

    route = result.scalar_one_or_none()

    if not route:

        raise HTTPException(status_code=404, detail="路线不存在")

    

    # Load experiments

    exp_result = await db.execute(

        select(RouteExperiment).where(

            RouteExperiment.route_id == str(route.id),

            RouteExperiment.is_deleted == False,

        ).order_by(RouteExperiment.experiment_date.desc())

    )

    experiments = exp_result.scalars().all()

    

    data = {

        "id": str(route.id),

        "project_id": route.project_id,

        "route_no": route.route_no,

        "name": route.name,

        "source": route.source,

        "source_reference": route.source_reference,

        "description": route.description,

        "status": route.status,

        "current_module": route.current_module,

        "literature_sources": route.literature_sources or [],

        "candidate_routes": route.candidate_routes or [],

        "selected_route_ids": route.selected_route_ids or [],

        "experiment_plans": route.experiment_plans or [],

        "experiments": [

            {

                "id": str(exp.id),

                "route_id": exp.route_id,

                "experiment_no": exp.experiment_no,

                "title": exp.title,

                "description": exp.description,

                "date": exp.experiment_date.isoformat() if exp.experiment_date else None,

                "operator": exp.operator,

                "status": exp.status,

                "reaction_temp": exp.reaction_temp,

                "reaction_time": exp.reaction_time,

                "yield": exp.yield_pct,

                "purity": exp.purity,

                "impurities": exp.impurities,

                "result_summary": exp.result_summary,

                "created_at": exp.created_at.isoformat() if exp.created_at else None,

                "updated_at": exp.updated_at.isoformat() if exp.updated_at else None,

            }

            for exp in experiments

        ],

        "assessment": route.assessment,

        "deliverables": route.deliverables or [],

        "start_date": route.start_date.isoformat() if route.start_date else None,

        "end_date": route.end_date.isoformat() if route.end_date else None,

        "created_at": route.created_at.isoformat() if route.created_at else None,

        "updated_at": route.updated_at.isoformat() if route.updated_at else None,

        "created_by": str(route.created_by) if route.created_by else None,

    }

    

    return success_response(data=data)





@router.post("/routes", summary="创建路线")

async def create_route(

    data: dict = Body(...),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    from app.modules.research.models import RouteDevelopment, RouteExperiment

    import uuid as uuid_mod

    from datetime import date as date_mod

    

    route = RouteDevelopment(

        id=str(uuid_mod.uuid4()),

        project_id=data.get("project_id", "default"),

        route_no=data.get("route_no", f"RD-{date_mod.today().year}-{str(uuid_mod.uuid4())[:6].upper()}"),

        name=data.get("name", "新路线"),

        source=data.get("source", "manual"),

        source_reference=data.get("source_reference"),

        description=data.get("description"),

        status="planning",

        current_module="research",

        created_by=current_user.id if current_user else None,

    )

    db.add(route)

    await db.flush()

    

    return success_response(data={"id": str(route.id), "route_no": route.route_no})





@router.put("/routes/{route_id}", summary="更新路线（保存工作流状态）")

async def update_route(

    route_id: str,

    data: dict = Body(...),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    from app.modules.research.models import RouteDevelopment, RouteExperiment

    from sqlalchemy import select

    

    result = await db.execute(

        select(RouteDevelopment).where(

            RouteDevelopment.id == route_id,

            RouteDevelopment.is_deleted == False,

        )

    )

    route = result.scalar_one_or_none()

    if not route:

        raise HTTPException(status_code=404, detail="路线不存在")

    

    # Update fields

    updatable = [

        "name", "source", "source_reference", "description", "status",

        "current_module", "literature_sources", "candidate_routes",

        "selected_route_ids", "experiment_plans", "assessment", "deliverables",

        "end_date",

    ]

    for field in updatable:

        if field in data:

            setattr(route, field, data[field])

    

    if current_user:

        route.updated_by = current_user.id

    

    await db.flush()

    return success_response(data={"id": str(route.id), "message": "已保存"})





@router.delete("/routes/{route_id}", summary="删除路线")

async def delete_route(

    route_id: str,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    from app.modules.research.models import RouteDevelopment, RouteExperiment

    from sqlalchemy import select

    

    # 先查询路线是否存在（包括已删除的）

    result = await db.execute(

        select(RouteDevelopment).where(

            RouteDevelopment.id == route_id,

        )

    )

    route = result.scalar_one_or_none()

    if not route:

        raise HTTPException(status_code=404, detail="路线不存在")

    

    # 如果已经删除，直接返回成功（幂等性）

    if route.is_deleted:

        return success_response(data={"message": "已删除"})

    

    route.is_deleted = True

    await db.flush()

    return success_response(data={"message": "已删除"})





# ==================== 实验记录 CRUD ====================



@router.post("/routes/{route_id}/experiments", summary="添加实验记录")

async def create_experiment(

    route_id: str,

    data: dict = Body(...),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    from app.modules.research.models import RouteDevelopment, RouteExperiment

    import uuid as uuid_mod

    from datetime import date as date_mod

    

    # Count existing experiments for this route

    from sqlalchemy import select, func

    count_result = await db.execute(

        select(func.count()).select_from(RouteExperiment).where(

            RouteExperiment.route_id == route_id,

            RouteExperiment.is_deleted == False,

        )

    )

    count = count_result.scalar_one()

    

    exp = RouteExperiment(

        id=str(uuid_mod.uuid4()),

        route_id=route_id,

        experiment_no=data.get("experiment_no", f"EXP-{count + 1:03d}"),

        title=data.get("title", "实验记录"),

        description=data.get("description"),

        experiment_date=date_mod.fromisoformat(data["date"]) if data.get("date") else date_mod.today(),

        operator=data.get("operator"),

        status=data.get("status", "planned"),

        reaction_temp=data.get("reaction_temp"),

        reaction_time=data.get("reaction_time"),

        yield_pct=data.get("yield"),

        purity=data.get("purity"),

        impurities=data.get("impurities"),

        result_summary=data.get("result_summary"),

        created_by=current_user.id if current_user else None,

    )

    db.add(exp)

    await db.flush()

    

    return success_response(data={"id": str(exp.id), "experiment_no": exp.experiment_no})





@router.put("/experiments/{exp_id}", summary="更新实验记录")

async def update_experiment(

    exp_id: str,

    data: dict = Body(...),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    from app.modules.research.models import RouteDevelopment, RouteExperiment

    from sqlalchemy import select

    from datetime import date as date_mod

    

    result = await db.execute(

        select(RouteExperiment).where(

            RouteExperiment.id == exp_id,

            RouteExperiment.is_deleted == False,

        )

    )

    exp = result.scalar_one_or_none()

    if not exp:

        raise HTTPException(status_code=404, detail="实验记录不存在")

    

    updatable = [

        "title", "description", "operator", "status",

        "reaction_temp", "reaction_time", "impurities", "result_summary",

    ]

    for field in updatable:

        if field in data:

            setattr(exp, field, data[field])

    

    if "date" in data and data["date"]:

        exp.experiment_date = date_mod.fromisoformat(data["date"])

    if "yield" in data:

        exp.yield_pct = data["yield"]

    if "purity" in data:

        exp.purity = data["purity"]

    if current_user:

        exp.updated_by = current_user.id

    

    await db.flush()

    return success_response(data={"id": str(exp.id), "message": "已保存"})





@router.delete("/experiments/{exp_id}", summary="删除实验记录")

async def delete_experiment(

    exp_id: str,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    from app.modules.research.models import RouteDevelopment, RouteExperiment

    from sqlalchemy import select

    

    result = await db.execute(

        select(RouteExperiment).where(

            RouteExperiment.id == exp_id,

            RouteExperiment.is_deleted == False,

        )

    )

    exp = result.scalar_one_or_none()

    if not exp:

        raise HTTPException(status_code=404, detail="实验记录不存在")

    

    exp.is_deleted = True

    await db.flush()

    return success_response(data={"message": "已删除"})





# ============ 工艺优化 API ============



@router.get("/optimizations", summary="获取工艺优化列表")

async def get_optimizations(

    current_user: CurrentUser,
    project_id: str | None = Query(None),

    status: str | None = Query(None),

    keyword: str | None = Query(None),

    page: int = Query(1, ge=1),

    page_size: int = Query(20, ge=1, le=100),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    from sqlalchemy import select, func

    from app.modules.research.models import ProcessOptimization



    query = select(ProcessOptimization)

    count_query = select(func.count()).select_from(ProcessOptimization)



    if project_id:

        query = query.where(ProcessOptimization.project_id == project_id)

        count_query = count_query.where(ProcessOptimization.project_id == project_id)

    if status:

        query = query.where(ProcessOptimization.status == status)

        count_query = count_query.where(ProcessOptimization.status == status)

    if keyword:

        query = query.where(

            ProcessOptimization.name.ilike(f"%{keyword}%")

        )

        count_query = count_query.where(

            ProcessOptimization.name.ilike(f"%{keyword}%")

        )



    query = query.order_by(ProcessOptimization.created_at.desc())

    query = query.offset((page - 1) * page_size).limit(page_size)



    result = await db.execute(query)

    optimizations = result.scalars().all()

    total_result = await db.execute(count_query)

    total = total_result.scalar()



    return paginated_response(

        data=[{

            "id": str(opt.id),

            "project_id": opt.project_id,

            "source_route_id": opt.source_route_id,

            "name": opt.name,

            "description": opt.description,

            "status": opt.status,

            "current_module": opt.current_module,

            "doe_experiment": opt.doe_experiment,

            

            "impurity_study": opt.impurity_study,

            "crystal_form_study": opt.crystal_form_study,

            "quality_standard_set": opt.quality_standard_set,

            "scale_up_study": opt.scale_up_study,

            

            "start_date": str(opt.start_date) if opt.start_date else None,

            "end_date": str(opt.end_date) if opt.end_date else None,

            "created_at": opt.created_at.isoformat() if opt.created_at else None,

            "updated_at": opt.updated_at.isoformat() if opt.updated_at else None,

        } for opt in optimizations],

        page=page,

        page_size=page_size,

        total=total,

    )





@router.get("/optimizations/{optimization_id}", summary="获取工艺优化详情")

async def get_optimization(

    current_user: CurrentUser,
    optimization_id: str,

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    from sqlalchemy import select

    from app.modules.research.models import ProcessOptimization



    result = await db.execute(

        select(ProcessOptimization).where(ProcessOptimization.id == optimization_id)

    )

    opt = result.scalar_one_or_none()

    if not opt:

        raise HTTPException(status_code=404, detail="工艺优化记录不存在")



    return success_response(data={

        "id": str(opt.id),

        "project_id": opt.project_id,

        "source_route_id": opt.source_route_id,

        "name": opt.name,

        "description": opt.description,

        "status": opt.status,

        "current_module": opt.current_module,

        "doe_experiment": opt.doe_experiment,

        

        "impurity_study": opt.impurity_study,

        "crystal_form_study": opt.crystal_form_study,

        "quality_standard_set": opt.quality_standard_set,

        "scale_up_study": opt.scale_up_study,

        

        "start_date": str(opt.start_date) if opt.start_date else None,

        "end_date": str(opt.end_date) if opt.end_date else None,

        "created_at": opt.created_at.isoformat() if opt.created_at else None,

        "updated_at": opt.updated_at.isoformat() if opt.updated_at else None,

    })





@router.post("/optimizations", summary="创建工艺优化")

async def create_optimization(

    data: dict = Body(...),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    import uuid

    from datetime import date

    from app.modules.research.models import ProcessOptimization



    opt = ProcessOptimization(

        id=str(uuid.uuid4()),

        optimization_no=data.get("optimization_no") or f"OPT-{uuid.uuid4().hex[:8].upper()}",

        project_id=data.get("project_id"),

        source_route_id=data.get("source_route_id"),

        source_route_name=data.get("source_route_name"),

        name=data.get("name", "新工艺优化"),

        description=data.get("description", ""),

        status="in_progress",

        current_module="doe",

        start_date=date.today(),

    )

    db.add(opt)

    await db.commit()

    await db.refresh(opt)



    return success_response(data={

        "id": str(opt.id),

        "source_route_id": opt.source_route_id,

        "name": opt.name,

        "status": opt.status,

        "current_module": opt.current_module,

    })





@router.put("/optimizations/{optimization_id}", summary="更新工艺优化")

async def update_optimization(

    optimization_id: str,

    data: dict = Body(...),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    from sqlalchemy import select

    from app.modules.research.models import ProcessOptimization



    result = await db.execute(

        select(ProcessOptimization).where(ProcessOptimization.id == optimization_id)

    )

    opt = result.scalar_one_or_none()

    if not opt:

        raise HTTPException(status_code=404, detail="工艺优化记录不存在")



    # 更新字段

    for key, value in data.items():

        if hasattr(opt, key) and key not in ("id", "created_at"):

            setattr(opt, key, value)



    await db.commit()

    await db.refresh(opt)



    return success_response(data={

        "id": str(opt.id),

        "source_route_id": opt.source_route_id,

        "name": opt.name,

        "status": opt.status,

        "current_module": opt.current_module,

    })





@router.delete("/optimizations/{optimization_id}", summary="删除工艺优化")

async def delete_optimization(

    optimization_id: str,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    from sqlalchemy import select

    from app.modules.research.models import ProcessOptimization



    result = await db.execute(

        select(ProcessOptimization).where(ProcessOptimization.id == optimization_id)

    )

    opt = result.scalar_one_or_none()

    if not opt:

        # Idempotent delete: return success even if record doesn't exist

        return success_response(data={"message": "工艺优化记录已删除"})



    await db.delete(opt)

    await db.commit()



    return success_response(data={"message": "工艺优化记录已删除"})





# ===== Pilot Workflow Endpoints =====



import uuid as _uuid

from datetime import datetime, timezone



from app.modules.research.models import PilotWorkflow, PilotWorkflowStep

from app.modules.research.schemas import (

    PilotWorkflowCreate,

    PilotWorkflowResponse,

    PilotWorkflowStepResponse,

    PilotWorkflowListItem,

)



# Rd Project schemas

from app.modules.research.schemas import (

    RdProjectCreate, RdProjectUpdate, RdProjectResponse,

    RdMilestoneCreate, RdMilestoneUpdate, RdMilestoneResponse,

    RdStageRecordCreate, RdStageRecordUpdate, RdStageRecordResponse,

    RdResearchTrackCreate, RdResearchTrackUpdate, RdResearchTrackResponse,

    RdResearchFindingCreate, RdResearchFindingUpdate, RdResearchFindingResponse

)





# ===== Milestone Endpoints =====



@router.post("/projects/{project_id}/milestones", response_model=RdMilestoneResponse, summary="创建里程碑")

async def create_milestone(

    project_id: UUID,

    data: RdMilestoneCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None

):

    user_id = current_user.id

    return await service.create_milestone(db, project_id, data, user_id)





@router.get("/projects/{project_id}/milestones", response_model=list[RdMilestoneResponse], summary="获取里程碑列表")

async def get_milestones(

    current_user: CurrentUser,
    project_id: UUID,

    db: AsyncSession = Depends(get_db)

):

    return await service.get_milestones(db, project_id)





@router.put("/projects/milestones/{milestone_id}", response_model=RdMilestoneResponse, summary="更新里程碑")

async def update_milestone(

    milestone_id: UUID,

    data: RdMilestoneUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None

):

    user_id = current_user.id

    return await service.update_milestone(db, milestone_id, data, user_id)





# ===== Stage Record Endpoints =====



@router.post("/projects/{project_id}/stages", response_model=RdStageRecordResponse, summary="创建阶段记录")

async def create_stage_record(

    project_id: UUID,

    data: RdStageRecordCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None

):

    user_id = current_user.id

    return await service.create_stage_record(db, project_id, data, user_id)





@router.get("/projects/{project_id}/stages", response_model=list[RdStageRecordResponse], summary="获取阶段记录列表")

async def get_stage_records(

    current_user: CurrentUser,
    project_id: UUID,

    db: AsyncSession = Depends(get_db)

):

    return await service.get_stage_records(db, project_id)





@router.put("/projects/stages/{record_id}", response_model=RdStageRecordResponse, summary="更新阶段记录")

async def update_stage_record(

    record_id: UUID,

    data: RdStageRecordUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None

):

    user_id = current_user.id

    return await service.update_stage_record(db, record_id, data, user_id)





# ===== Research Track Endpoints =====



@router.post("/projects/{project_id}/tracks", response_model=RdResearchTrackResponse, summary="创建研究项")

async def create_research_track(

    project_id: UUID,

    data: RdResearchTrackCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None

):

    user_id = current_user.id

    return await service.create_research_track(db, project_id, data, user_id)





@router.get("/tracks", summary="获取所有研究项列表")

async def get_all_research_tracks(

    current_user: CurrentUser,
    project_id: str | None = Query(None, description="按项目ID筛选"),

    track_type: str | None = Query(None, description="按研究项类型筛选"),

    db: AsyncSession = Depends(get_db),

):

    """获取所有研究项，支持按项目和类型筛选"""

    from sqlalchemy import select

    from app.modules.research.models import RdResearchTrack, RdProject



    stmt = (

        select(RdResearchTrack, RdProject.name.label("project_name"))

        .outerjoin(RdProject, RdResearchTrack.project_id == RdProject.id)

        .where(RdResearchTrack.is_deleted == False)

        .where((RdProject.is_deleted == False) | (RdProject.id.is_(None)))

    )

    if project_id and project_id.strip():

        stmt = stmt.where(RdResearchTrack.project_id == project_id.strip())

    if track_type and track_type.strip():

        stmt = stmt.where(RdResearchTrack.type == track_type.strip())

    stmt = stmt.order_by(RdResearchTrack.created_at.desc())



    result = await db.execute(stmt)

    tracks = []

    for row in result.all():

        track = row[0]

        project_name = row[1]

        tracks.append({

            "id": str(track.id),

            "project_id": str(track.project_id),

            "project_name": project_name or "",

            "type": track.type,

            "name": track.name,

            "description": track.description,

            "status": track.status,

            "priority": track.priority,

            "current_conclusion": track.current_conclusion,

            "conclusion_version": track.conclusion_version,

            "conclusion_confidence": track.conclusion_confidence,

            "active_stages": track.active_stages,

            "owner_id": str(track.owner_id) if track.owner_id else None,

            "created_at": track.created_at.isoformat() if track.created_at else None,

            "updated_at": track.updated_at.isoformat() if track.updated_at else None,

        })

    return success_response(data=tracks)





@router.get("/projects/{project_id}/tracks", response_model=list[RdResearchTrackResponse], summary="获取研究项列表")

async def get_research_tracks(

    current_user: CurrentUser,
    project_id: UUID,

    db: AsyncSession = Depends(get_db)

):

    return await service.get_research_tracks(db, project_id)





@router.put("/projects/tracks/{track_id}", response_model=RdResearchTrackResponse, summary="更新研究项")

async def update_research_track(

    track_id: UUID,

    data: RdResearchTrackUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None

):

    user_id = current_user.id

    return await service.update_research_track(db, track_id, data, user_id)





# ===== Research Finding Endpoints =====



@router.post("/tracks/{track_id}/findings", response_model=RdResearchFindingResponse, summary="创建研究发现")

async def create_research_finding(

    track_id: UUID,

    data: RdResearchFindingCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None

):

    user_id = current_user.id

    return await service.create_research_finding(db, track_id, data, user_id)





@router.get("/tracks/{track_id}/findings", response_model=list[RdResearchFindingResponse], summary="获取研究发现列表")

async def get_research_findings(

    current_user: CurrentUser,
    track_id: UUID,

    db: AsyncSession = Depends(get_db)

):

    return await service.get_research_findings(db, track_id)





@router.put("/findings/{finding_id}", response_model=RdResearchFindingResponse, summary="更新研究发现")

async def update_research_finding(

    finding_id: UUID,

    data: RdResearchFindingUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None

):

    user_id = current_user.id

    return await service.update_research_finding(db, finding_id, data, user_id)





# ===== Conclusion Version Endpoints =====



@router.post("/tracks/{track_id}/conclusions", summary="发布新结论版本")

async def publish_conclusion_version(

    track_id: UUID,

    data: dict,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None

):

    """发布新的结论版本，更新研究项的当前结论"""

    user_id = current_user.id

    conclusion = data.get("conclusion", "")

    confidence = data.get("confidence", "preliminary")

    return await service.publish_conclusion_version(db, track_id, conclusion, confidence, user_id)





@router.get("/tracks/{track_id}/conclusions", summary="获取结论历史")

async def get_conclusion_history(

    current_user: CurrentUser,
    track_id: UUID,

    db: AsyncSession = Depends(get_db)

):

    """获取研究项的结论版本历史"""

    return await service.get_conclusion_history(db, track_id)





# ===== RdProject API Endpoints =====



@router.get("/rd-projects", summary="获取研发项目列表(RdProject)")

async def get_rd_projects(

    current_user: CurrentUser,
    stage: str | None = Query(None, description="当前阶段"),

    status: str | None = Query(None, description="状态"),

    keyword: str | None = Query(None, description="搜索关键词"),

    project_type: str | None = Query(None, description="项目类型"),

    page: int = Query(1, ge=1),

    page_size: int = Query(20, ge=1, le=100),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    projects, total = await service.get_rd_projects(

        db, stage=stage, status=status, keyword=keyword,

        project_type=project_type, page=page, page_size=page_size

    )

    return paginated_response(

        data=[RdProjectResponse.model_validate(p) for p in projects],

        page=page,

        page_size=page_size,

        total=total,

    )





@router.get("/rd-projects/{project_id}", summary="获取研发项目详情(RdProject)")

async def get_rd_project(

    current_user: CurrentUser,
    project_id: UUID,

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    project = await service.get_rd_project(db, project_id)

    return success_response(data=RdProjectResponse.model_validate(project))





@router.post("/rd-projects", summary="创建研发项目(RdProject)")

async def create_rd_project(

    data: RdProjectCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    project = await service.create_rd_project(db, data, user_id)

    return success_response(

        data=RdProjectResponse.model_validate(project),

        message="创建成功",

    )





@router.put("/rd-projects/{project_id}", summary="更新研发项目(RdProject)")

async def update_rd_project(

    project_id: UUID,

    data: RdProjectUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    project = await service.update_rd_project(db, project_id, data, user_id)

    return success_response(

        data=RdProjectResponse.model_validate(project),

        message="更新成功",

    )





@router.delete("/rd-projects/{project_id}", summary="删除研发项目(RdProject)")

async def delete_rd_project(

    project_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    await service.delete_rd_project(db, project_id, user_id)

    return success_response(message="删除成功")





# ===== Stage Transition API Endpoints =====



@router.post("/rd-projects/{project_id}/transition", summary="阶段流转")

async def transition_stage(

    project_id: UUID,

    data: dict = Body(...),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    target_stage = data.get("target_stage")

    review_notes = data.get("review_notes")

    

    if not target_stage:

        return error_response(message="缺少 target_stage 参数")

    

    user_id = current_user.id if current_user else None

    result = await service.transition_stage(db, project_id, target_stage, review_notes, user_id)

    

    if result.get("success"):

        return success_response(data=result, message="阶段流转成功")

    else:

        return error_response(message=result.get("message", "阶段流转失败"))





@router.get("/rd-projects/{project_id}/transition-check", summary="检查阶段流转条件")

async def check_stage_transition(

    current_user: CurrentUser,
    project_id: UUID,

    target_stage: str = Query(..., description="目标阶段"),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    result = await service.check_stage_transition(db, project_id, target_stage)

    return success_response(data=result)





# ===== 中试研究 API Endpoints =====



@router.get("/pilot-studies", summary="获取中试研究记录列表")

async def get_pilot_studies(

    current_user: CurrentUser,
    project_id: UUID = Query(..., description="项目ID"),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    studies = await service.get_pilot_studies(db, project_id)

    return success_response(

        data=[RdPilotStudyResponse.model_validate(s) for s in studies]

    )





@router.post("/pilot-studies", summary="创建中试研究记录")

async def create_pilot_study(

    data: RdPilotStudyCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    study = await service.create_pilot_study(db, data, user_id)

    return success_response(

        data=RdPilotStudyResponse.model_validate(study),

        message="创建成功",

    )





@router.put("/pilot-studies/{study_id}", summary="更新中试研究记录")

async def update_pilot_study(

    study_id: UUID,

    data: RdPilotStudyUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    study = await service.update_pilot_study(db, study_id, data, user_id)

    return success_response(

        data=RdPilotStudyResponse.model_validate(study),

        message="更新成功",

    )





# ===== 工艺验证 API Endpoints =====



@router.get("/process-validations", summary="获取工艺验证记录列表")

async def get_validations(

    current_user: CurrentUser,
    project_id: UUID = Query(..., description="项目ID"),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    validations = await service.get_validations(db, project_id)

    return success_response(

        data=[RdProcessValidationResponse.model_validate(v) for v in validations]

    )





@router.post("/process-validations", summary="创建工艺验证记录")

async def create_validation(

    data: RdProcessValidationCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    validation = await service.create_validation(db, data, user_id)

    return success_response(

        data=RdProcessValidationResponse.model_validate(validation),

        message="创建成功",

    )





@router.put("/process-validations/{validation_id}", summary="更新工艺验证记录")

async def update_validation(

    validation_id: UUID,

    data: RdProcessValidationUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    validation = await service.update_validation(db, validation_id, data, user_id)

    return success_response(

        data=RdProcessValidationResponse.model_validate(validation),

        message="更新成功",

    )





# ===== 申报资料 API Endpoints =====



@router.get("/registration-filings", summary="获取申报资料记录列表")

async def get_filings(

    current_user: CurrentUser,
    project_id: UUID = Query(..., description="项目ID"),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    filings = await service.get_filings(db, project_id)

    return success_response(

        data=[RdRegistrationFilingResponse.model_validate(f) for f in filings]

    )





@router.post("/registration-filings", summary="创建申报资料记录")

async def create_filing(

    data: RdRegistrationFilingCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    filing = await service.create_filing(db, data, user_id)

    return success_response(

        data=RdRegistrationFilingResponse.model_validate(filing),

        message="创建成功",

    )





@router.put("/registration-filings/{filing_id}", summary="更新申报资料记录")

async def update_filing(

    filing_id: UUID,

    data: RdRegistrationFilingUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    filing = await service.update_filing(db, filing_id, data, user_id)

    return success_response(

        data=RdRegistrationFilingResponse.model_validate(filing),

        message="更新成功",

    )





# ===== RdStageDeliverable API Endpoints =====



@router.post(

    "/rd-stage-deliverables",

    response_model=RdStageDeliverableResponse,

    summary="创建阶段交付物",

)

async def create_rd_stage_deliverable_api(

    data: RdStageDeliverableCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> RdStageDeliverableResponse:

    user_id = current_user.id if current_user else None

    deliverable = await service.create_rd_stage_deliverable(db, data, user_id)

    return RdStageDeliverableResponse.model_validate(deliverable.__dict__)





@router.get(

    "/rd-stage-deliverables/{deliverable_id}",

    response_model=RdStageDeliverableResponse,

    summary="获取阶段交付物",

)

async def get_rd_stage_deliverable_api(

    current_user: CurrentUser,
    deliverable_id: UUID,

    db: AsyncSession = Depends(get_db),

) -> RdStageDeliverableResponse:

    deliverable = await service.get_rd_stage_deliverable(db, deliverable_id)

    return RdStageDeliverableResponse.model_validate(deliverable.__dict__)





@router.get(

    "/rd-stage-deliverables",

    summary="获取阶段交付物列表",

)

async def list_rd_stage_deliverables_api(

    current_user: CurrentUser,
    project_id: Optional[UUID] = None,

    stage: Optional[str] = None,

    deliverable_type: Optional[str] = None,

    status: Optional[str] = None,

    page: int = 1,

    page_size: int = 20,

    db: AsyncSession = Depends(get_db),

):

    items, total = await service.list_rd_stage_deliverables(

        db, project_id, stage, deliverable_type, status, page, page_size

    )

    return {

        "items": [RdStageDeliverableResponse.model_validate(i.__dict__).model_dump() for i in items],

        "total": total,

        "page": page,

        "page_size": page_size,

    }





@router.put(

    "/rd-stage-deliverables/{deliverable_id}",

    response_model=RdStageDeliverableResponse,

    summary="更新阶段交付物",

)

async def update_rd_stage_deliverable_api(

    deliverable_id: UUID,

    data: RdStageDeliverableUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> RdStageDeliverableResponse:

    user_id = current_user.id if current_user else None

    deliverable = await service.update_rd_stage_deliverable(db, deliverable_id, data, user_id)

    return RdStageDeliverableResponse.model_validate(deliverable.__dict__)





@router.delete(

    "/rd-stage-deliverables/{deliverable_id}",

    summary="删除阶段交付物",

)

async def delete_rd_stage_deliverable_api(

    deliverable_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> list:

    user_id = current_user.id if current_user else None

    await service.delete_rd_stage_deliverable(db, deliverable_id, user_id)

    return {"message": "删除成功"}



@router.post(

    "/rd-stage-deliverables/{deliverable_id}/upload",

    summary="上传交付物附件",

)

async def upload_deliverable_file(

    deliverable_id: UUID,

    file: UploadFile = File(...),

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> list:

    """上传交付物附件文件（本地存储）"""

    import os

    UPLOAD_DIR = "/tmp/dazah_uploads/research/deliverables"

    os.makedirs(UPLOAD_DIR, exist_ok=True)

    

    content = await file.read()

    file_size = len(content)

    filename = file.filename or "unknown"

    

    # 保存文件到本地

    file_path = os.path.join(UPLOAD_DIR, f"{deliverable_id}_{filename}")

    with open(file_path, "wb") as f:

        f.write(content)

    

    # 更新交付物记录

    file_url = f"/api/v1/research/rd-stage-deliverables/{deliverable_id}/download"

    await service.update_rd_stage_deliverable(

        db, 

        deliverable_id, 

        RdStageDeliverableUpdate(file_url=file_url, file_name=filename, file_size=file_size),

        None

    )

    

    return success_response(data={

        "file_url": file_url,

        "file_name": filename,

        "file_size": file_size,

    })





@router.get(

    "/rd-stage-deliverables/{deliverable_id}/download",

    summary="下载交付物附件",

)

async def download_deliverable_file(

    current_user: CurrentUser,
    deliverable_id: UUID,

    db: AsyncSession = Depends(get_db),

):

    """下载交付物附件文件（本地存储）"""

    import os

    from fastapi.responses import FileResponse

    

    deliverable = await service.get_rd_stage_deliverable(db, deliverable_id)

    

    if not deliverable.file_name:

        raise HTTPException(status_code=404, detail="未找到附件文件")

    

    UPLOAD_DIR = "/tmp/dazah_uploads/research/deliverables"

    file_path = os.path.join(UPLOAD_DIR, f"{deliverable_id}_{deliverable.file_name}")

    

    if not os.path.exists(file_path):

        raise HTTPException(status_code=404, detail="文件不存在")

    

    return FileResponse(

        file_path,

        filename=deliverable.file_name,

        media_type='application/octet-stream'

    )





# ===== 中试研究 API =====



@router.get("/rd-pilot-studies", summary="获取中试研究列表")

async def get_pilot_studies(

    current_user: CurrentUser,
    project_id: UUID,

    db: AsyncSession = Depends(get_db),

) -> list:

    items = await service.get_pilot_studies(db, project_id)

    return [RdPilotStudyResponse.model_validate(i.__dict__).model_dump() for i in items]





@router.post("/rd-pilot-studies", summary="创建中试研究")

async def create_pilot_study_api(

    data: RdPilotStudyCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    item = await service.create_pilot_study(db, data, user_id)

    return RdPilotStudyResponse.model_validate(item.__dict__).model_dump()





@router.put("/rd-pilot-studies/{study_id}", summary="更新中试研究")

async def update_pilot_study_api(

    study_id: UUID,

    data: RdPilotStudyUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    item = await service.update_pilot_study(db, study_id, data, user_id)

    return RdPilotStudyResponse.model_validate(item.__dict__).model_dump()





@router.delete("/rd-pilot-studies/{study_id}", summary="删除中试研究")

async def delete_pilot_study_api(

    study_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    await service.delete_pilot_study(db, study_id, user_id)

    return {"message": "删除成功"}





# ===== 工艺验证 API =====



@router.get("/rd-process-validations", summary="获取工艺验证列表")

async def get_validations(

    current_user: CurrentUser,
    project_id: UUID,

    db: AsyncSession = Depends(get_db),

) -> list:

    items = await service.get_validations(db, project_id)

    return [RdProcessValidationResponse.model_validate(i.__dict__).model_dump() for i in items]





@router.post("/rd-process-validations", summary="创建工艺验证")

async def create_validation_api(

    data: RdProcessValidationCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    item = await service.create_validation(db, data, user_id)

    return RdProcessValidationResponse.model_validate(item.__dict__).model_dump()





@router.put("/rd-process-validations/{validation_id}", summary="更新工艺验证")

async def update_validation_api(

    validation_id: UUID,

    data: RdProcessValidationUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    item = await service.update_validation(db, validation_id, data, user_id)

    return RdProcessValidationResponse.model_validate(item.__dict__).model_dump()





@router.delete("/rd-process-validations/{validation_id}", summary="删除工艺验证")

async def delete_validation_api(

    validation_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    await service.delete_validation(db, validation_id, user_id)

    return {"message": "删除成功"}





# ===== 申报资料 API =====



@router.get("/rd-registration-filings", summary="获取申报资料列表")

async def get_filings(

    current_user: CurrentUser,
    project_id: UUID,

    db: AsyncSession = Depends(get_db),

) -> list:

    items = await service.get_filings(db, project_id)

    return [RdRegistrationFilingResponse.model_validate(i.__dict__).model_dump() for i in items]





@router.post("/rd-registration-filings", summary="创建申报资料")

async def create_filing_api(

    data: RdRegistrationFilingCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    item = await service.create_filing(db, data, user_id)

    return RdRegistrationFilingResponse.model_validate(item.__dict__).model_dump()





@router.put("/rd-registration-filings/{filing_id}", summary="更新申报资料")

async def update_filing_api(

    filing_id: UUID,

    data: RdRegistrationFilingUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    item = await service.update_filing(db, filing_id, data, user_id)

    return RdRegistrationFilingResponse.model_validate(item.__dict__).model_dump()





@router.delete("/rd-registration-filings/{filing_id}", summary="删除申报资料")

async def delete_filing_api(

    filing_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    await service.delete_filing(db, filing_id, user_id)

    return {"message": "删除成功"}





# ===== 实验记录 API Endpoints =====



@router.get("/experiment-logs", summary="获取实验记录列表")

async def get_experiment_logs(

    current_user: CurrentUser,
    project_id: UUID = Query(..., description="项目ID"),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    logs = await service.get_experiment_logs(db, project_id)

    return success_response(

        data=[RdExperimentLogResponse.model_validate(log) for log in logs]

    )





@router.post("/experiment-logs", summary="创建实验记录")

async def create_experiment_log(

    data: RdExperimentLogCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    log = await service.create_experiment_log(db, data, user_id)

    await db.commit()

    return success_response(

        data=RdExperimentLogResponse.model_validate(log),

        message="创建成功",

    )





@router.put("/experiment-logs/{log_id}", summary="更新实验记录")

async def update_experiment_log(

    log_id: UUID,

    data: RdExperimentLogUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    log = await service.update_experiment_log(db, log_id, data, user_id)

    await db.commit()

    return success_response(

        data=RdExperimentLogResponse.model_validate(log),

        message="更新成功",

    )





@router.delete("/experiment-logs/{log_id}", summary="删除实验记录")

async def delete_experiment_log(

    log_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    await service.delete_experiment_log(db, log_id, user_id)

    return {"message": "删除成功"}





# ===== 研发报告 API Endpoints =====



@router.get("/reports", summary="获取研发报告列表")

async def get_reports(

    current_user: CurrentUser,
    project_id: UUID = Query(..., description="项目ID"),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    reports = await service.get_reports(db, project_id)

    return success_response(

        data=[RdReportResponse.model_validate(r) for r in reports]

    )





@router.post("/reports", summary="创建研发报告")

async def create_report(

    data: RdReportCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    report = await service.create_report(db, data, user_id)

    await db.commit()

    return success_response(

        data=RdReportResponse.model_validate(report),

        message="创建成功",

    )





@router.put("/reports/{report_id}", summary="更新研发报告")

async def update_report(

    report_id: UUID,

    data: RdReportUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    report = await service.update_report(db, report_id, data, user_id)

    await db.commit()

    return success_response(

        data=RdReportResponse.model_validate(report),

        message="更新成功",

    )





@router.delete("/reports/{report_id}", summary="删除研发报告")

async def delete_report(

    report_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    await service.delete_report(db, report_id, user_id)

    return {"message": "删除成功"}





# ===== 立项申请 API Endpoints =====



@router.get("/initiations", summary="获取立项申请列表")

async def get_initiations(

    current_user: CurrentUser,
    project_id: UUID = Query(..., description="项目ID"),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    items = await service.get_initiations(db, project_id)

    return success_response(

        data=[RdInitiationResponse.model_validate(i.__dict__) for i in items]

    )





@router.post("/initiations", summary="创建立项申请")

async def create_initiation(

    data: RdInitiationCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    item = await service.create_initiation(db, data, user_id)

    await db.commit()

    return success_response(

        data=RdInitiationResponse.model_validate(item.__dict__),

        message="创建成功",

    )





@router.put("/initiations/{initiation_id}", summary="更新立项申请")

async def update_initiation(

    initiation_id: UUID,

    data: RdInitiationUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    item = await service.update_initiation(db, initiation_id, data, user_id)

    await db.commit()

    return success_response(

        data=RdInitiationResponse.model_validate(item.__dict__),

        message="更新成功",

    )





@router.delete("/initiations/{initiation_id}", summary="删除立项申请")

async def delete_initiation(

    initiation_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    await service.delete_initiation(db, initiation_id, user_id)

    return {"message": "删除成功"}







# ===== 研究项/发现删除 API =====



@router.delete("/projects/tracks/{track_id}", summary="删除研究项")

async def delete_research_track(

    track_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    """软删除研究项"""

    from app.modules.research.models import RdResearchTrack

    result = await db.execute(

        select(RdResearchTrack).where(

            RdResearchTrack.id == track_id,

            RdResearchTrack.is_deleted == False,

        )

    )

    track = result.scalar_one_or_none()

    if not track:

        raise HTTPException(status_code=404, detail="研究项不存在")

    track.is_deleted = True

    track.updated_by = current_user.id

    await db.commit()

    return {"message": "删除成功"}





@router.delete("/findings/{finding_id}", summary="删除研究发现")

async def delete_research_finding(

    finding_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    """软删除研究发现"""

    from app.modules.research.models import RdResearchFinding

    result = await db.execute(

        select(RdResearchFinding).where(

            RdResearchFinding.id == finding_id,

            RdResearchFinding.is_deleted == False,

        )

    )

    finding = result.scalar_one_or_none()

    if not finding:

        raise HTTPException(status_code=404, detail="研究发现不存在")

    finding.is_deleted = True

    finding.updated_by = current_user.id

    await db.commit()

    return {"message": "删除成功"}



# ===== 研究项详情 API =====



@router.get("/tracks/{track_id}", summary="获取研究项详情")

async def get_track_detail(

    current_user: CurrentUser,
    track_id: UUID,

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    """获取研究项详情，包含研究发现和结论历史"""

    from sqlalchemy import select

    from app.modules.research.models import RdResearchTrack

    result = await db.execute(

        select(RdResearchTrack).where(

            RdResearchTrack.id == track_id,

            RdResearchTrack.is_deleted == False,

        )

    )

    track = result.scalar_one_or_none()

    if not track:

        raise HTTPException(status_code=404, detail="研究项不存在")

    

    # Get findings

    findings = await service.get_research_findings(db, track_id)

    

    # Get conclusion history

    conclusion_history = await service.get_conclusion_history(db, track_id)

    

    track_data = RdResearchTrackResponse.model_validate(track.__dict__).model_dump()

    track_data["findings"] = [RdResearchFindingResponse.model_validate(f.__dict__).model_dump() for f in findings]

    track_data["conclusion_history"] = conclusion_history

    

    return success_response(data=track_data)





# ===== 结论版本管理 API =====



@router.post("/tracks/{track_id}/conclusion-versions", summary="发布新结论版本")

async def create_conclusion_version_api(

    track_id: UUID,

    data: RdTrackConclusionVersionCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    """发布新的结论版本"""

    version_data = data.model_dump()

    version_data["track_id"] = track_id

    version_data["author_id"] = current_user.id

    

    # Also update the track's current conclusion

    result = await service.publish_conclusion_version(

        db, track_id, 

        data.conclusion or "", 

        data.confidence,

        current_user.id,

        data.change_summary,

        data.evidence_refs,

    )

    await db.commit()

    return success_response(data=result, message="结论版本发布成功")





@router.get("/tracks/{track_id}/conclusion-versions", summary="获取结论版本历史")

async def get_conclusion_versions_api(

    current_user: CurrentUser,
    track_id: UUID,

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    """获取研究项的结论版本历史"""

    versions = await service.get_conclusion_history(db, track_id)

    return success_response(data=versions)





# ===== 数据导出 API =====



import csv

import io

from fastapi.responses import StreamingResponse



@router.get("/export/projects", summary="导出项目列表 CSV")

async def export_projects_csv(

    current_user: CurrentUser,
    db: AsyncSession = Depends(get_db),

):

    """导出所有研发项目为 CSV 文件"""

    from sqlalchemy import select

    from app.modules.research.models import RdProject

    result = await db.execute(

        select(RdProject).where(RdProject.is_deleted == False)

    )

    projects = result.scalars().all()

    

    # Create CSV

    output = io.StringIO()

    writer = csv.writer(output)

    

    # Header

    writer.writerow([

        'ID', '品种名称', 'API全称', 'CAS号', '分子式', '分子量',

        '适应症', '项目类型', '优先级', '当前阶段', '状态', '总体进度',

        '开始日期', '目标申报日期', '实际申报日期', '备注'

    ])

    

    # Data rows

    for p in projects:

        writer.writerow([

            str(p.id), p.name, p.api_name, p.cas_number or '',

            p.molecular_formula or '', p.molecular_weight or '',

            p.indication or '', p.project_type or '', p.priority or '',

            p.current_stage or '', p.status or '', p.overall_progress or '',

            p.start_date.isoformat() if p.start_date else '', 

            p.target_filing_date.isoformat() if p.target_filing_date else '',

            p.actual_filing_date.isoformat() if p.actual_filing_date else '', 

            p.notes or ''

        ])

    

    # Return as streaming response

    output.seek(0)

    return StreamingResponse(

        iter([output.getvalue()]),

        media_type="text/csv",

        headers={"Content-Disposition": "attachment; filename=rd_projects.csv"}

    )





@router.get("/export/tracks", summary="导出研究项 CSV")

async def export_tracks_csv(

    current_user: CurrentUser,
    project_id: UUID = Query(..., description="项目ID"),

    db: AsyncSession = Depends(get_db),

):

    """导出指定项目的研究项为 CSV 文件"""

    from sqlalchemy import select

    from app.modules.research.models import RdResearchTrack

    result = await db.execute(

        select(RdResearchTrack).where(

            RdResearchTrack.project_id == project_id,

            RdResearchTrack.is_deleted == False,

        )

    )

    tracks = result.scalars().all()

    

    # Create CSV

    output = io.StringIO()

    writer = csv.writer(output)

    

    # Header

    writer.writerow([

        '研究项名称', '类型', '状态', '优先级', '负责人',

        '当前结论', '结论版本', '置信度', '创建时间'

    ])

    

    # Data rows

    for t in tracks:

        writer.writerow([

            t.name, t.type, t.status, t.priority, t.owner_id or '',

            t.current_conclusion or '', t.conclusion_version or 0,

            t.conclusion_confidence or '',

            t.created_at.strftime('%Y-%m-%d %H:%M') if t.created_at else ''

        ])

    

    # Return as streaming response

    output.seek(0)

    return StreamingResponse(

        iter([output.getvalue()]),

        media_type="text/csv",

        headers={"Content-Disposition": "attachment; filename=rd_tracks.csv"}

    )





@router.get("/export/experiment-logs", summary="导出实验记录 CSV")

async def export_experiment_logs_csv(

    current_user: CurrentUser,
    project_id: UUID = Query(..., description="项目ID"),

    db: AsyncSession = Depends(get_db),

):

    """导出指定项目的实验记录为 CSV 文件"""

    from sqlalchemy import select

    from app.modules.research.models import RdExperimentLog

    result = await db.execute(

        select(RdExperimentLog).where(

            RdExperimentLog.project_id == project_id,

            RdExperimentLog.is_deleted == False,

        )

    )

    logs = result.scalars().all()

    

    # Create CSV

    output = io.StringIO()

    writer = csv.writer(output)

    

    # Header

    writer.writerow([

        '实验标题', '实验类型', '实验日期', '操作人', '状态',

        '实验目的', '实验步骤', '实验现象', '实验结论', '备注'

    ])

    

    # Data rows

    for log in logs:

        writer.writerow([

            log.title, log.experiment_type,

            log.experiment_date.strftime('%Y-%m-%d') if log.experiment_date else '',

            log.operator or '', log.status,

            log.objective or '', log.procedure or '',

            log.observations or '', log.conclusion or '', log.notes or ''

        ])

    

    # Return as streaming response

    output.seek(0)

    return StreamingResponse(

        iter([output.getvalue()]),

        media_type="text/csv",

        headers={"Content-Disposition": "attachment; filename=rd_experiment_logs.csv"}

    )





# ===== 统计报表 API =====



@router.get("/stats/overview", summary="研发概览统计")

async def get_stats_overview(

    current_user: CurrentUser,
    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    """获取研发模块概览统计数据"""

    from sqlalchemy import select, func

    from app.modules.research.models import RdProject, RdResearchTrack, RdExperimentLog, RdStageDeliverable

    

    # 项目统计

    total_projects = await db.execute(

        select(func.count(RdProject.id)).where(RdProject.is_deleted == False)

    )

    total_projects = total_projects.scalar() or 0

    

    # 按阶段统计

    stage_stats = await db.execute(

        select(RdProject.current_stage, func.count(RdProject.id))

        .where(RdProject.is_deleted == False)

        .group_by(RdProject.current_stage)

    )

    stage_distribution = {stage or 'unknown': count for stage, count in stage_stats.all()}

    

    # 按状态统计

    status_stats = await db.execute(

        select(RdProject.status, func.count(RdProject.id))

        .where(RdProject.is_deleted == False)

        .group_by(RdProject.status)

    )

    status_distribution = {status or 'unknown': count for status, count in status_stats.all()}

    

    # 研究项统计 (只统计未删除项目中的研究项)

    from app.modules.research.models import RdProject

    total_tracks = await db.execute(

        select(func.count(RdResearchTrack.id))

        .where(RdResearchTrack.is_deleted == False)

        .where(RdResearchTrack.project_id.in_(

            select(RdProject.id).where(RdProject.is_deleted == False)

        ))

    )

    total_tracks = total_tracks.scalar() or 0

    

    # 按类型统计研究项 (只统计未删除项目中的研究项)

    track_type_stats = await db.execute(

        select(RdResearchTrack.type, func.count(RdResearchTrack.id))

        .where(RdResearchTrack.is_deleted == False)

        .where(RdResearchTrack.project_id.in_(

            select(RdProject.id).where(RdProject.is_deleted == False)

        ))

        .group_by(RdResearchTrack.type)

    )

    track_type_distribution = {t or 'unknown': c for t, c in track_type_stats.all()}

    

    # 实验记录统计

    total_experiments = await db.execute(

        select(func.count(RdExperimentLog.id)).where(RdExperimentLog.is_deleted == False)

    )

    total_experiments = total_experiments.scalar() or 0

    

    # 交付物统计

    total_deliverables = await db.execute(

        select(func.count(RdStageDeliverable.id)).where(RdStageDeliverable.is_deleted == False)

    )

    total_deliverables = total_deliverables.scalar() or 0

    

    # 按状态统计交付物

    deliverable_status_stats = await db.execute(

        select(RdStageDeliverable.status, func.count(RdStageDeliverable.id))

        .where(RdStageDeliverable.is_deleted == False)

        .group_by(RdStageDeliverable.status)

    )

    deliverable_status_distribution = {s or 'unknown': c for s, c in deliverable_status_stats.all()}

    

    return success_response(data={

        "projects": {

            "total": total_projects,

            "by_stage": stage_distribution,

            "by_status": status_distribution,

        },

        "tracks": {

            "total": total_tracks,

            "by_type": track_type_distribution,

        },

        "experiments": {

            "total": total_experiments,

        },

        "deliverables": {

            "total": total_deliverables,

            "by_status": deliverable_status_distribution,

        },

    })





@router.get("/stats/project-progress", summary="项目进度统计")

async def get_project_progress(

    current_user: CurrentUser,
    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    """获取各项目进度统计"""

    from sqlalchemy import select

    from app.modules.research.models import RdProject

    

    result = await db.execute(

        select(

            RdProject.id,

            RdProject.name,

            RdProject.current_stage,

            RdProject.status,

            RdProject.overall_progress,

            RdProject.start_date,

            RdProject.target_filing_date,

        ).where(RdProject.is_deleted == False)

    )

    

    # Stage-to-progress mapping: each stage maps to a base progress percentage

    stage_progress_map = {

        'initiation': 8,

        'route_dev': 20,

        'optimization': 40,

        'pilot': 60,

        'validation': 80,

        'filing': 95,

    }



    projects = []

    for row in result.all():

        # Calculate progress from current_stage if overall_progress is not set

        progress = row.overall_progress

        if not progress and row.current_stage:

            progress = stage_progress_map.get(row.current_stage, 0)

        elif not progress:

            progress = 0



        projects.append({

            "id": str(row.id),

            "name": row.name,

            "current_stage": row.current_stage,

            "status": row.status,

            "progress": progress,

            "start_date": row.start_date.isoformat() if row.start_date else None,

            "target_filing_date": row.target_filing_date.isoformat() if row.target_filing_date else None,

        })

    

    return success_response(data=projects)





# ===== 交付物模板 API =====



@router.get("/deliverable-templates", summary="获取交付物模板列表")

async def get_deliverable_templates(

    current_user: CurrentUser,
    stage: str | None = Query(None, description="阶段"),

    deliverable_type: str | None = Query(None, description="交付物类型"),

    is_active: bool | None = Query(None, description="是否启用"),

    db: AsyncSession = Depends(get_db),

) -> JSONResponse:

    templates = await service.get_deliverable_templates(db, stage, deliverable_type, is_active)

    return success_response(

        data=[RdDeliverableTemplateResponse.model_validate(t.__dict__) for t in templates]

    )





@router.post("/deliverable-templates", summary="创建交付物模板")

async def create_deliverable_template(

    data: RdDeliverableTemplateCreate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    template = await service.create_deliverable_template(db, data, user_id)

    await db.commit()

    return success_response(

        data=RdDeliverableTemplateResponse.model_validate(template.__dict__),

        message="创建成功",

    )





@router.put("/deliverable-templates/{template_id}", summary="更新交付物模板")

async def update_deliverable_template(

    template_id: UUID,

    data: RdDeliverableTemplateUpdate,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    user_id = current_user.id if current_user else None

    template = await service.update_deliverable_template(db, template_id, data, user_id)

    await db.commit()

    return success_response(

        data=RdDeliverableTemplateResponse.model_validate(template.__dict__),

        message="更新成功",

    )





@router.delete("/deliverable-templates/{template_id}", summary="删除交付物模板")

async def delete_deliverable_template(

    template_id: UUID,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> dict:

    user_id = current_user.id if current_user else None

    await service.delete_deliverable_template(db, template_id, user_id)

    return {"message": "删除成功"}





# ===== AI 报告生成 API =====



@router.post("/generate-report", summary="使用 AI 生成报告")

async def generate_report(

    data: RdReportGenerateRequest,

    db: AsyncSession = Depends(get_db),

    current_user: CurrentUser = None,

) -> JSONResponse:

    result = await service.generate_report_with_ai(

        db,

        data.project_id,

        data.deliverable_type,

        data.template_id,

        data.additional_context,

    )

    return success_response(data=result, message="生成成功")
