from uuid import UUID

from fastapi import Depends, File, Query, UploadFile
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.core.response import paginated_response, success_response
from app.modules.administration.schemas import (
    GiftInventoryCreate,
    GiftInventoryResponse,
    GiftInventoryUpdate,
    GiftRequisitionCreate,
    GiftRequisitionResponse,
    GiftRequisitionUpdate,
    ITServiceTicketCreate,
    ITServiceTicketResponse,
    ITServiceTicketUpdate,
    RegulationCreate,
    RegulationResponse,
    RegulationUpdate,
    VehicleCreate,
    VehicleRequestCreate,
    VehicleRequestResponse,
    VehicleRequestUpdate,
    VehicleResponse,
    VehicleUpdate,
)
from app.modules.administration.service import (
    GiftInventoryService,
    GiftRequisitionService,
    ITServiceTicketService,
    RegulationService,
    VehicleRequestService,
    VehicleService,
)


def _extract_text_from_file(file_data: str | None, file_type: str | None, file_name: str | None) -> str:
    """尝试从上传的文件中提取文本内容."""
    if not file_data:
        return ""

    import base64
    from io import BytesIO

    try:
        file_bytes = base64.b64decode(file_data)
    except Exception:
        return ""

    ext = (file_name or "").lower().split(".")[-1] if file_name else ""
    mime = (file_type or "").lower()

    # PDF
    if ext == "pdf" or mime == "application/pdf":
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            texts = []
            for page in doc:
                texts.append(page.get_text())
            doc.close()
            return "\n".join(texts).strip()
        except Exception:
            return ""

    # Word (.docx)
    if ext == "docx" or mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        try:
            from docx import Document
            doc = Document(BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(row_texts))
                tables.append("\n".join(table_rows))
            all_text = "\n".join(paragraphs)
            if tables:
                all_text += "\n\n[表格内容]\n" + "\n\n".join(tables)
            return all_text.strip()
        except Exception:
            return ""

    # Plain text / markdown
    if ext in ("txt", "md") or mime.startswith("text/"):
        try:
            return file_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            try:
                return file_bytes.decode("gbk").strip()
            except UnicodeDecodeError:
                return ""

    return ""


async def _extract_text_with_ai(file_data: str | None, file_type: str | None, file_name: str | None) -> str:
    """当常规提取失败时，使用AI vision识别图片/PDF内容."""
    if not file_data:
        return ""

    import base64
    from io import BytesIO
    import openai
    from app.core.config import get_settings

    settings = get_settings()
    if not settings.MOONSHOT_API_KEY:
        return ""

    client = openai.AsyncOpenAI(
        api_key=settings.MOONSHOT_API_KEY,
        base_url="https://api.moonshot.cn/v1",
    )

    try:
        file_bytes = base64.b64decode(file_data)
    except Exception:
        return ""

    ext = (file_name or "").lower().split(".")[-1] if file_name else ""
    mime = (file_type or "").lower()
    images: list[str] = []

    # PDF: 转前3页为图片
    if ext == "pdf" or mime == "application/pdf":
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page_idx in range(min(3, len(doc))):
                page = doc.load_page(page_idx)
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                img_b64 = base64.b64encode(img_bytes).decode()
                images.append(img_b64)
            doc.close()
        except Exception:
            return ""

    # 图片文件
    elif mime.startswith("image/"):
        images = [file_data]

    if not images:
        return ""

    content_parts: list[dict] = [
        {
            "type": "text",
            "text": (
                "请识别以下图片中的全部文字内容，保留原有的段落结构、表格格式和数字。"
                "只输出识别到的原始文字，不要添加任何解释、总结或额外说明。"
            ),
        }
    ]
    for img_b64 in images:
        content_parts.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{img_b64}"},
            }
        )

    messages = [
        {
            "role": "system",
            "content": "你是一个文档OCR识别助手，擅长从图片中提取文字并保留原始格式。",
        },
        {"role": "user", "content": content_parts},
    ]

    try:
        response = await client.chat.completions.create(
            model=settings.AI_MODEL or "moonshot-v1-128k",
            messages=messages,  # type: ignore
            temperature=0.1,
            max_tokens=4096,
        )
        return response.choices[0].message.content or ""
    except Exception:
        return ""


from app.shared.module_api import create_module_router
from app.shared.module_registry import MODULES_BY_CODE
from app.shared.schemas import PageParams

router = create_module_router(MODULES_BY_CODE["administration"])


def get_vehicle_service(session: AsyncSession = Depends(get_db)) -> VehicleService:
    return VehicleService(session)


def get_vehicle_request_service(
    session: AsyncSession = Depends(get_db),
) -> VehicleRequestService:
    return VehicleRequestService(session)


def get_it_ticket_service(
    session: AsyncSession = Depends(get_db),
) -> ITServiceTicketService:
    return ITServiceTicketService(session)


def get_regulation_service(
    session: AsyncSession = Depends(get_db),
) -> RegulationService:
    return RegulationService(session)


def get_gift_inventory_service(
    session: AsyncSession = Depends(get_db),
) -> GiftInventoryService:
    return GiftInventoryService(session)


def get_gift_requisition_service(
    session: AsyncSession = Depends(get_db),
) -> GiftRequisitionService:
    return GiftRequisitionService(session)


# ─── Regulation Routes ───

@router.get("/regulations", summary="规章制度列表")
async def list_regulations(
    keyword: str | None = Query(None, description="制度名称或内容关键词"),
    page_params: PageParams = Depends(),
    service: RegulationService = Depends(get_regulation_service),
):
    regulations, total = await service.list_regulations(
        keyword=keyword,
        page=page_params.page,
        page_size=page_params.page_size,
    )
    data = [RegulationResponse.model_validate(r).model_dump(mode="json") for r in regulations]
    return paginated_response(
        data=data,
        page=page_params.page,
        page_size=page_params.page_size,
        total=total,
    )


@router.post("/regulations/extract", summary="提取制度文件文字内容")
async def extract_regulation_text(
    payload: RegulationCreate,
):
    """接收文件数据，先用pymupdf/python-docx提取文本；
    如果提取为空（扫描件），调用AI vision模型进行OCR识别。"""
    if not payload.file_data:
        return success_response(data={"text": "", "source": "none"}, message="未提供文件数据")

    # 1. 常规提取
    text = _extract_text_from_file(payload.file_data, payload.file_type, payload.file_name)
    if text:
        return success_response(
            data={"text": text, "source": "native"},
            message="已提取文本内容",
        )

    # 2. AI vision 识别（扫描件 fallback）
    ai_text = await _extract_text_with_ai(payload.file_data, payload.file_type, payload.file_name)
    if ai_text:
        return success_response(
            data={"text": ai_text, "source": "ai"},
            message="AI识别完成",
        )

    return success_response(
        data={"text": "", "source": "none"},
        message="未能提取到文本内容，请手动输入",
    )


@router.post("/regulations", summary="创建规章制度")
async def create_regulation(
    payload: RegulationCreate,
    service: RegulationService = Depends(get_regulation_service),
):
    regulation = await service.create_regulation(payload)
    return success_response(
        data=RegulationResponse.model_validate(regulation).model_dump(mode="json"),
        message="规章制度创建成功",
        status_code=201,
    )


@router.get("/regulations/{regulation_id}", summary="规章制度详情")
async def get_regulation(
    regulation_id: UUID,
    service: RegulationService = Depends(get_regulation_service),
):
    regulation = await service.get_regulation(regulation_id)
    return success_response(
        data=RegulationResponse.model_validate(regulation).model_dump(mode="json"),
    )


@router.put("/regulations/{regulation_id}", summary="更新规章制度")
async def update_regulation(
    regulation_id: UUID,
    payload: RegulationUpdate,
    service: RegulationService = Depends(get_regulation_service),
):
    regulation = await service.update_regulation(regulation_id, payload)
    return success_response(
        data=RegulationResponse.model_validate(regulation).model_dump(mode="json"),
        message="规章制度更新成功",
    )


@router.delete("/regulations/{regulation_id}", summary="删除规章制度")
async def delete_regulation(
    regulation_id: UUID,
    service: RegulationService = Depends(get_regulation_service),
):
    await service.delete_regulation(regulation_id)
    return success_response(message="规章制度删除成功")


# ─── Vehicle Routes ───

@router.get("/vehicles", summary="车辆列表")
async def list_vehicles(
    keyword: str | None = Query(None, description="车牌号或品牌关键词"),
    status: str | None = Query(None, description="状态筛选"),
    page_params: PageParams = Depends(),
    service: VehicleService = Depends(get_vehicle_service),
):
    vehicles, total = await service.list_vehicles(
        keyword=keyword,
        status=status,
        page=page_params.page,
        page_size=page_params.page_size,
    )
    data = [VehicleResponse.model_validate(v).model_dump(mode="json") for v in vehicles]
    return paginated_response(
        data=data,
        page=page_params.page,
        page_size=page_params.page_size,
        total=total,
    )


@router.post("/vehicles", summary="创建车辆")
async def create_vehicle(
    payload: VehicleCreate,
    service: VehicleService = Depends(get_vehicle_service),
):
    vehicle = await service.create_vehicle(payload)
    return success_response(
        data=VehicleResponse.model_validate(vehicle).model_dump(mode="json"),
        message="车辆创建成功",
        status_code=201,
    )


@router.get("/vehicles/{vehicle_id}", summary="车辆详情")
async def get_vehicle(
    vehicle_id: UUID,
    service: VehicleService = Depends(get_vehicle_service),
):
    vehicle = await service.get_vehicle(vehicle_id)
    return success_response(
        data=VehicleResponse.model_validate(vehicle).model_dump(mode="json"),
    )


@router.put("/vehicles/{vehicle_id}", summary="更新车辆")
async def update_vehicle(
    vehicle_id: UUID,
    payload: VehicleUpdate,
    service: VehicleService = Depends(get_vehicle_service),
):
    vehicle = await service.update_vehicle(vehicle_id, payload)
    return success_response(
        data=VehicleResponse.model_validate(vehicle).model_dump(mode="json"),
        message="车辆更新成功",
    )


@router.delete("/vehicles/{vehicle_id}", summary="删除车辆")
async def delete_vehicle(
    vehicle_id: UUID,
    service: VehicleService = Depends(get_vehicle_service),
):
    await service.delete_vehicle(vehicle_id)
    return success_response(message="车辆删除成功")


@router.get("/vehicles/{vehicle_id}/photo", summary="查看车辆照片")
async def download_vehicle_photo(
    vehicle_id: UUID,
    service: VehicleService = Depends(get_vehicle_service),
):
    from fastapi.responses import StreamingResponse
    import base64
    from io import BytesIO

    vehicle = await service.get_vehicle(vehicle_id)
    if not vehicle.photo_data:
        raise ValueError("该车辆没有上传照片")

    file_bytes = base64.b64decode(vehicle.photo_data)
    mime_type = vehicle.photo_type or "image/jpeg"

    return StreamingResponse(
        BytesIO(file_bytes),
        media_type=mime_type,
    )


@router.post("/vehicles/batch-import", summary="批量导入车辆")
async def batch_import_vehicles(
    file: UploadFile = File(...),
    service: VehicleService = Depends(get_vehicle_service),
):
    contents = await file.read()
    result = await service.batch_import(contents, file.content_type or file.filename or "")
    return success_response(
        data=result,
        message=f"导入完成：新增 {result.get('created', 0)} 条，恢复 {result.get('restored', 0)} 条，失败 {result.get('failed', 0)} 条",
    )


# ─── Vehicle Request Routes ───

@router.get("/vehicle-requests", summary="用车申请列表")
async def list_vehicle_requests(
    keyword: str | None = Query(None, description="申请人或事由关键词"),
    status: str | None = Query(None, description="状态筛选"),
    page_params: PageParams = Depends(),
    service: VehicleRequestService = Depends(get_vehicle_request_service),
):
    requests, total = await service.list_requests(
        keyword=keyword,
        status=status,
        page=page_params.page,
        page_size=page_params.page_size,
    )
    data = [VehicleRequestResponse.model_validate(r).model_dump(mode="json") for r in requests]
    return paginated_response(
        data=data,
        page=page_params.page,
        page_size=page_params.page_size,
        total=total,
    )


@router.post("/vehicle-requests", summary="创建用车申请")
async def create_vehicle_request(
    payload: VehicleRequestCreate,
    service: VehicleRequestService = Depends(get_vehicle_request_service),
):
    request = await service.create_request(payload)
    return success_response(
        data=VehicleRequestResponse.model_validate(request).model_dump(mode="json"),
        message="用车申请创建成功",
        status_code=201,
    )


@router.get("/vehicle-requests/{request_id}", summary="用车申请详情")
async def get_vehicle_request(
    request_id: UUID,
    service: VehicleRequestService = Depends(get_vehicle_request_service),
):
    request = await service.get_request(request_id)
    return success_response(
        data=VehicleRequestResponse.model_validate(request).model_dump(mode="json"),
    )


@router.put("/vehicle-requests/{request_id}", summary="更新用车申请")
async def update_vehicle_request(
    request_id: UUID,
    payload: VehicleRequestUpdate,
    service: VehicleRequestService = Depends(get_vehicle_request_service),
):
    request = await service.update_request(request_id, payload)
    return success_response(
        data=VehicleRequestResponse.model_validate(request).model_dump(mode="json"),
        message="用车申请更新成功",
    )


@router.delete("/vehicle-requests/{request_id}", summary="删除用车申请")
async def delete_vehicle_request(
    request_id: UUID,
    service: VehicleRequestService = Depends(get_vehicle_request_service),
):
    await service.delete_request(request_id)
    return success_response(message="用车申请删除成功")


@router.post("/vehicle-requests/sync-from-feishu", summary="从飞书同步用车申请数据")
async def sync_vehicle_requests_from_feishu(
    service: VehicleRequestService = Depends(get_vehicle_request_service),
):
    """从飞书多维表格同步用车申请数据到本地数据库."""
    stats = await service.sync_from_feishu()
    return success_response(
        data=stats,
        message=f"同步完成：新增 {stats['created']} 条，更新 {stats['updated']} 条，失败 {stats['failed']} 条",
    )


# ─── IT Service Ticket Routes ───

@router.get("/it-service-tickets", summary="IT工单列表")
async def list_it_service_tickets(
    keyword: str | None = Query(None, description="标题或报障人关键词"),
    status: str | None = Query(None, description="状态筛选"),
    ticket_type: str | None = Query(None, description="工单类型筛选"),
    page_params: PageParams = Depends(),
    service: ITServiceTicketService = Depends(get_it_ticket_service),
):
    tickets, total = await service.list_tickets(
        keyword=keyword,
        status=status,
        ticket_type=ticket_type,
        page=page_params.page,
        page_size=page_params.page_size,
    )
    data = [ITServiceTicketResponse.model_validate(t).model_dump(mode="json") for t in tickets]
    return paginated_response(
        data=data,
        page=page_params.page,
        page_size=page_params.page_size,
        total=total,
    )


@router.post("/it-service-tickets", summary="创建IT工单")
async def create_it_service_ticket(
    payload: ITServiceTicketCreate,
    service: ITServiceTicketService = Depends(get_it_ticket_service),
):
    ticket = await service.create_ticket(payload)
    return success_response(
        data=ITServiceTicketResponse.model_validate(ticket).model_dump(mode="json"),
        message="IT工单创建成功",
        status_code=201,
    )


@router.get("/it-service-tickets/{ticket_id}", summary="IT工单详情")
async def get_it_service_ticket(
    ticket_id: UUID,
    service: ITServiceTicketService = Depends(get_it_ticket_service),
):
    ticket = await service.get_ticket(ticket_id)
    return success_response(
        data=ITServiceTicketResponse.model_validate(ticket).model_dump(mode="json"),
    )


@router.put("/it-service-tickets/{ticket_id}", summary="更新IT工单")
async def update_it_service_ticket(
    ticket_id: UUID,
    payload: ITServiceTicketUpdate,
    service: ITServiceTicketService = Depends(get_it_ticket_service),
):
    ticket = await service.update_ticket(ticket_id, payload)
    return success_response(
        data=ITServiceTicketResponse.model_validate(ticket).model_dump(mode="json"),
        message="IT工单更新成功",
    )


@router.delete("/it-service-tickets/{ticket_id}", summary="删除IT工单")
async def delete_it_service_ticket(
    ticket_id: UUID,
    service: ITServiceTicketService = Depends(get_it_ticket_service),
):
    await service.delete_ticket(ticket_id)
    return success_response(message="IT工单删除成功")


# ─── Regulation File Routes ───

@router.get("/regulations/{regulation_id}/file", summary="下载/预览规章制度原文件")
async def download_regulation_file(
    regulation_id: UUID,
    service: RegulationService = Depends(get_regulation_service),
):
    from fastapi.responses import StreamingResponse
    import base64
    from io import BytesIO

    regulation = await service.get_regulation(regulation_id)
    if not regulation.file_data:
        raise ValueError("该制度没有上传原始文件")

    file_bytes = base64.b64decode(regulation.file_data)
    mime_type = regulation.file_type or "application/octet-stream"
    file_name = regulation.file_name or f"{regulation.title}.pdf"

    # RFC 5987 encoding for non-ASCII filenames
    from urllib.parse import quote
    encoded_name = quote(file_name, safe='')

    return StreamingResponse(
        BytesIO(file_bytes),
        media_type=mime_type,
        headers={
            "Content-Disposition": f"inline; filename*=UTF-8''{encoded_name}",
        },
    )


# ─── Gift Requisition Routes ───

@router.get("/gift-requisitions", summary="领用记录列表")
async def list_gift_requisitions(
    department: str | None = Query(None, description="费用所属部门"),
    item_name: str | None = Query(None, description="领用物品名称"),
    recipient: str | None = Query(None, description="领用人"),
    page_params: PageParams = Depends(),
    service: GiftRequisitionService = Depends(get_gift_requisition_service),
):
    requisitions, total = await service.list_requisitions(
        department=department,
        item_name=item_name,
        recipient=recipient,
        page=page_params.page,
        page_size=page_params.page_size,
    )
    data = [GiftRequisitionResponse(**r).model_dump(mode="json") for r in requisitions]
    return paginated_response(
        data=data,
        page=page_params.page,
        page_size=page_params.page_size,
        total=total,
    )


@router.post("/gift-requisitions", summary="创建领用记录")
async def create_gift_requisition(
    payload: GiftRequisitionCreate,
    service: GiftRequisitionService = Depends(get_gift_requisition_service),
):
    req = await service.create_requisition(payload)
    return success_response(
        data=GiftRequisitionResponse(**req).model_dump(mode="json"),
        message="领用记录创建成功",
        status_code=201,
    )


@router.get("/gift-requisitions/{req_id}", summary="领用记录详情")
async def get_gift_requisition(
    req_id: UUID,
    service: GiftRequisitionService = Depends(get_gift_requisition_service),
):
    req = await service.get_requisition(req_id)
    return success_response(
        data=GiftRequisitionResponse(**req).model_dump(mode="json"),
    )


@router.put("/gift-requisitions/{req_id}", summary="更新领用记录")
async def update_gift_requisition(
    req_id: UUID,
    payload: GiftRequisitionUpdate,
    service: GiftRequisitionService = Depends(get_gift_requisition_service),
):
    req = await service.update_requisition(req_id, payload)
    return success_response(
        data=GiftRequisitionResponse(**req).model_dump(mode="json"),
        message="领用记录更新成功",
    )


@router.delete("/gift-requisitions/{req_id}", summary="删除领用记录")
async def delete_gift_requisition(
    req_id: UUID,
    service: GiftRequisitionService = Depends(get_gift_requisition_service),
):
    await service.delete_requisition(req_id)
    return success_response(message="领用记录删除成功")

@router.get("/gift-inventories", summary="礼品酒水库存列表")
async def list_gift_inventories(
    keyword: str | None = Query(None, description="物品名称关键词"),
    status: str | None = Query(None, description="状态筛选"),
    page_params: PageParams = Depends(),
    service: GiftInventoryService = Depends(get_gift_inventory_service),
):
    inventories, total = await service.list_inventories(
        keyword=keyword,
        status=status,
        page=page_params.page,
        page_size=page_params.page_size,
    )
    data = [GiftInventoryResponse(**i).model_dump(mode="json") for i in inventories]
    return paginated_response(
        data=data,
        page=page_params.page,
        page_size=page_params.page_size,
        total=total,
    )


@router.post("/gift-inventories", summary="创建库存记录")
async def create_gift_inventory(
    payload: GiftInventoryCreate,
    service: GiftInventoryService = Depends(get_gift_inventory_service),
):
    inventory = await service.create_inventory(payload)
    return success_response(
        data=GiftInventoryResponse(**inventory).model_dump(mode="json"),
        message="库存记录创建成功",
        status_code=201,
    )


@router.get("/gift-inventories/{inventory_id}", summary="库存记录详情")
async def get_gift_inventory(
    inventory_id: UUID,
    service: GiftInventoryService = Depends(get_gift_inventory_service),
):
    inventory = await service.get_inventory(inventory_id)
    return success_response(
        data=GiftInventoryResponse(**inventory).model_dump(mode="json"),
    )


@router.put("/gift-inventories/{inventory_id}", summary="更新库存记录")
async def update_gift_inventory(
    inventory_id: UUID,
    payload: GiftInventoryUpdate,
    service: GiftInventoryService = Depends(get_gift_inventory_service),
):
    inventory = await service.update_inventory(inventory_id, payload)
    return success_response(
        data=GiftInventoryResponse(**inventory).model_dump(mode="json"),
        message="库存记录更新成功",
    )


@router.delete("/gift-inventories/{inventory_id}", summary="删除库存记录")
async def delete_gift_inventory(
    inventory_id: UUID,
    service: GiftInventoryService = Depends(get_gift_inventory_service),
):
    await service.delete_inventory(inventory_id)
    return success_response(message="库存记录删除成功")
