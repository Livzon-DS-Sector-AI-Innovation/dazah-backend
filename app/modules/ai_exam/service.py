"""AI exam generation service."""

import json
import logging
import re
from io import BytesIO

import httpx
from docx import Document as DocxDocument

from app.core.config import get_settings

logger = logging.getLogger(__name__)


def _extract_docx_content(file_bytes: bytes) -> dict:
    """Parse docx file, return {full_text, bold_texts}."""
    doc = DocxDocument(BytesIO(file_bytes))
    all_paragraphs: list[str] = []
    bold_texts: list[str] = []

    for para in doc.paragraphs:
        para_parts: list[str] = []
        for run in para.runs:
            text = run.text.strip()
            if not text:
                continue
            para_parts.append(text)
            if run.bold:
                bold_texts.append(text)
        if para_parts:
            all_paragraphs.append(" ".join(para_parts))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text.strip()
                        if text and run.bold:
                            bold_texts.append(text)

    return {"full_text": "\n".join(all_paragraphs), "bold_texts": bold_texts}


def _extract_text_content(file_bytes: bytes) -> dict:
    text = file_bytes.decode("utf-8", errors="ignore")
    return {"full_text": text, "bold_texts": []}


def _parse_file(file_bytes: bytes, filename: str) -> dict:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("docx", "doc"):
        return _extract_docx_content(file_bytes)
    return _extract_text_content(file_bytes)


def _build_prompt(full_text: str, bold_texts: list[str], config: dict | None = None) -> str:
    """Build AI prompt with configurable question types and counts."""
    if config:
        choice_count = config.get("choice_count", 5)
        tf_count = config.get("true_false_count", 5)
        qa_count = config.get("qa_count", 0)
    else:
        choice_count, tf_count, qa_count = 5, 5, 0

    bold_hint = ""
    if bold_texts:
        bold_hint = (
            "\n\n【重点关注】以下内容在原文件中被加粗标记，请优先作为考点出题：\n"
            + "\n".join(f"  - {t}" for t in bold_texts[:20])
        )

    req_lines = []
    i = 1
    if choice_count > 0:
        req_lines.append(f"{i}. 出 {choice_count} 道单选题（每道 4 个选项 A/B/C/D，只有一个正确答案）")
        i += 1
    if tf_count > 0:
        req_lines.append(f"{i}. 出 {tf_count} 道判断题（正确/错误）")
        i += 1
    if qa_count > 0:
        req_lines.append(f"{i}. 出 {qa_count} 道简答题（需简要回答）")
        i += 1
    req_lines.append(f"{i}. 题目应覆盖材料的关键知识点，加粗内容优先出题")
    i += 1
    req_lines.append(f"{i}. 答案必须基于材料内容，不得编造")
    i += 1
    req_lines.append(f"{i}. 严格按 JSON 格式输出，不要输出其他内容")

    json_parts = []
    if choice_count > 0:
        json_parts.append('"choice_questions": [{"question": "...", "options": [{"label": "A", "text": "..."}, {"label": "B", "text": "..."}, {"label": "C", "text": "..."}, {"label": "D", "text": "..."}], "answer": "A"}]')
    if tf_count > 0:
        json_parts.append('"true_false_questions": [{"question": "...", "answer": "正确"}]')
    if qa_count > 0:
        json_parts.append('"qa_questions": [{"question": "...", "answer": "参考答案"}]')

    return f"""你是一名专业的培训考核出题老师。请根据以下培训材料，生成一套考试试卷。

要求：
{chr(10).join(req_lines)}

输出格式：
{{ {", ".join(json_parts)} }}

材料内容：
{full_text[:8000]}{bold_hint}"""


async def generate_exam(file_bytes: bytes, filename: str, config: dict | None = None) -> dict:
    """Generate exam questions from uploaded file."""
    content = _parse_file(file_bytes, filename)
    if not content["full_text"].strip():
        raise ValueError("文件中未检测到文本内容")

    bold_texts = content["bold_texts"]
    logger.info(
        "AI exam: %d chars, %d bold segments from %s, config=%s",
        len(content["full_text"]), len(bold_texts), filename, config,
    )

    settings = get_settings()
    api_key = settings.HR_AI_API_KEY
    if not api_key:
        raise ValueError("HR_AI_API_KEY 未配置，无法调用 AI")

    prompt = _build_prompt(content["full_text"], bold_texts, config)

    async with httpx.AsyncClient(timeout=120) as client:
        resp = await client.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": settings.HR_AI_MODEL,
                "messages": [
                    {"role": "system", "content": "你是一名严谨的考核出题老师，只输出 JSON。"},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.7,
                "max_tokens": 4096,
            },
        )
        resp.raise_for_status()
        data = resp.json()
        raw = data["choices"][0]["message"]["content"]

    json_match = re.search(r"\{[\s\S]*\}", raw)
    if not json_match:
        raise ValueError(f"AI 返回格式异常: {raw[:200]}")

    result = json.loads(json_match.group())

    for i, q in enumerate(result.get("choice_questions", [])):
        q["number"] = i + 1
    for i, q in enumerate(result.get("true_false_questions", [])):
        q["number"] = i + 1
    for i, q in enumerate(result.get("qa_questions", [])):
        q["number"] = i + 1

    return result


def export_exam(data: dict) -> BytesIO:
    """Export exam questions as a Word document."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    title = data.get("title", "培训考试试卷")
    doc.add_heading(title, level=1)
    doc.add_paragraph("")

    info = []
    if data.get("examiner"):
        info.append(f"出卷人：{data['examiner']}")
    if data.get("exam_date"):
        info.append(f"考试日期：{data['exam_date']}")
    if data.get("assessment_date"):
        info.append(f"评估日期：{data['assessment_date']}")
    if info:
        doc.add_paragraph("  |  ".join(info))
        doc.add_paragraph("")

    num = 1
    choice_qs = data.get("choice_questions", [])
    if choice_qs:
        doc.add_heading(f"{num}、单选题", level=2)
        num += 1
        for q in choice_qs:
            doc.add_paragraph(f"{q['number']}. {q['question']}")
            for opt in q.get("options", []):
                doc.add_paragraph(f"    {opt['label']}. {opt['text']}")
            doc.add_paragraph("")

    tf_qs = data.get("true_false_questions", [])
    if tf_qs:
        doc.add_heading(f"{num}、判断题", level=2)
        num += 1
        for q in tf_qs:
            doc.add_paragraph(f"{q['number']}. {q['question']} （  ）")
            doc.add_paragraph("")

    qa_qs = data.get("qa_questions", [])
    if qa_qs:
        doc.add_heading(f"{num}、简答题", level=2)
        for q in qa_qs:
            doc.add_paragraph(f"{q['number']}. {q['question']}")
            doc.add_paragraph("")

    doc.add_heading("参考答案", level=2)
    if choice_qs:
        doc.add_paragraph("单选题答案：")
        for q in choice_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")
    if tf_qs:
        doc.add_paragraph("判断题答案：")
        for q in tf_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")
    if qa_qs:
        doc.add_paragraph("简答题答案：")
        for q in qa_qs:
            doc.add_paragraph(f"  {q['number']}. {q.get('answer', '?')}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
