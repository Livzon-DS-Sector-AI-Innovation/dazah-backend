"""原料检验数据表 API"""

import json
import logging
from typing import Optional
from uuid import UUID

from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile
from pydantic import BaseModel

from app.core.database import get_db, AsyncSession

logger = logging.getLogger(__name__)


class ApiResponse(BaseModel):
    """统一响应格式"""
    code: int = 200
    message: str = "Success"
    data: Optional[dict | list] = None


class CreateTableRequest(BaseModel):
    """创建数据表请求"""
    table_name: str
    table_description: Optional[str] = None
    columns_config: list = []


class UpdateTableRequest(BaseModel):
    """更新数据表请求"""
    table_name: Optional[str] = None
    table_description: Optional[str] = None
    columns_config: Optional[list] = None
    is_active: Optional[bool] = None


class RowDataRequest(BaseModel):
    """行数据请求"""
    row_data: dict


class BatchRowsRequest(BaseModel):
    """批量行数据请求"""
    rows: list[dict]


router = APIRouter(prefix="/inspection-table", tags=["原料检验数据"])


# ============ 数据表管理 API ============

@router.get("/", summary="获取数据表列表")
async def list_tables(
    is_active: Optional[bool] = None,
    keyword: Optional[str] = None,
    page: int = 1,
    page_size: int = 20,
    session: AsyncSession = Depends(get_db),
):
    """获取数据表列表"""
    from app.modules.quality.inspection_table_service import InspectionTableService

    service = InspectionTableService(session)
    tables, total = await service.list_tables(
        is_active=is_active,
        keyword=keyword,
        page=page,
        page_size=page_size,
    )

    return ApiResponse(
        data={
            "items": tables,
            "total": total,
            "page": page,
            "page_size": page_size,
        }
    )


@router.post("/", summary="创建数据表")
async def create_table(
    request: Request,
    session: AsyncSession = Depends(get_db),
):
    """创建数据表"""
    from app.modules.quality.inspection_table_service import InspectionTableService

    # 读取原始请求体
    body = await request.body()
    logger.info(f"Received raw body: {body}")

    try:
        data = json.loads(body)
        logger.info(f"Parsed JSON: {data}")
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error: {e}")
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {e}")

    # 手动验证数据
    if 'table_name' not in data:
        raise HTTPException(status_code=400, detail="Missing required field: table_name")

    service = InspectionTableService(session)

    try:
        table = await service.create_table(
            data.get('table_name'),
            data.get('table_description'),
            data.get('columns_config', [])
        )
        return ApiResponse(data=table)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.get("/{table_id}", summary="获取数据表详情")
async def get_table(
    table_id: UUID,
    session: AsyncSession = Depends(get_db),
):
    """获取数据表详情"""
    from app.modules.quality.inspection_table_service import InspectionTableService

    service = InspectionTableService(session)
    table = await service.get_table(table_id)

    if not table:
        raise HTTPException(status_code=404, detail="数据表不存在")

    return ApiResponse(data=table)


@router.put("/{table_id}", summary="更新数据表")
async def update_table(
    table_id: UUID,
    request: UpdateTableRequest,
    session: AsyncSession = Depends(get_db),
):
    """更新数据表"""
    from app.modules.quality.inspection_table_service import InspectionTableService

    service = InspectionTableService(session)

    update_data = {}
    if request.table_name is not None:
        update_data["table_name"] = request.table_name
    if request.table_description is not None:
        update_data["table_description"] = request.table_description
    if request.columns_config is not None:
        update_data["columns_config"] = request.columns_config
    if request.is_active is not None:
        update_data["is_active"] = request.is_active

    table = await service.update_table(table_id, update_data)

    if not table:
        raise HTTPException(status_code=404, detail="数据表不存在")

    return ApiResponse(data=table)


@router.delete("/{table_id}", summary="删除数据表")
async def delete_table(
    table_id: UUID,
    session: AsyncSession = Depends(get_db),
):
    """删除数据表"""
    from app.modules.quality.inspection_table_service import InspectionTableService

    service = InspectionTableService(session)
    success = await service.delete_table(table_id)

    if not success:
        raise HTTPException(status_code=404, detail="数据表不存在")

    return ApiResponse(message="删除成功")


# ============ 数据行 API ============

@router.post("/{table_id}/rows", summary="添加数据行")
async def add_row(
    table_id: UUID,
    request: RowDataRequest,
    session: AsyncSession = Depends(get_db),
):
    """添加数据行"""
    from app.modules.quality.inspection_table_service import InspectionTableService

    service = InspectionTableService(session)

    try:
        row = await service.add_row(table_id, request.row_data)
        return ApiResponse(data=row)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.put("/{table_id}/rows/{row_id}", summary="更新数据行")
async def update_row(
    table_id: UUID,
    row_id: int,
    request: RowDataRequest,
    session: AsyncSession = Depends(get_db),
):
    """更新数据行"""
    from app.modules.quality.inspection_table_service import InspectionTableService

    service = InspectionTableService(session)
    row = await service.update_row(row_id, request.row_data)

    if not row:
        raise HTTPException(status_code=404, detail="数据行不存在")

    return ApiResponse(data=row)


@router.delete("/{table_id}/rows/{row_id}", summary="删除数据行")
async def delete_row(
    table_id: UUID,
    row_id: int,
    session: AsyncSession = Depends(get_db),
):
    """删除数据行"""
    from app.modules.quality.inspection_table_service import InspectionTableService

    service = InspectionTableService(session)
    success = await service.delete_row(row_id)

    if not success:
        raise HTTPException(status_code=404, detail="数据行不存在")

    return ApiResponse(message="删除成功")


@router.post("/{table_id}/rows/batch", summary="批量保存数据行")
async def batch_save_rows(
    table_id: UUID,
    request: BatchRowsRequest,
    session: AsyncSession = Depends(get_db),
):
    """批量保存数据行"""
    from app.modules.quality.inspection_table_service import InspectionTableService

    try:
        logger.info(f"Batch save rows for table {table_id}, rows count: {len(request.rows)}")

        service = InspectionTableService(session)
        saved_rows = await service.batch_save_rows(table_id, request.rows)

        logger.info(f"Batch save completed, saved {len(saved_rows)} rows")
        return ApiResponse(data={"rows": saved_rows})
    except Exception as e:
        logger.error(f"Batch save error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"保存失败: {str(e)}")


# ============ AI 识别 API ============

@router.post("/{table_id}/recognize", summary="AI识别扫描件")
async def recognize_image(
    table_id: UUID,
    request: Request,
    session: AsyncSession = Depends(get_db),
):
    """通过AI识别扫描件图片，提取数据"""
    from app.modules.quality.inspection_table_service import InspectionTableService

    # 获取数据表配置
    service = InspectionTableService(session)
    table = await service.get_table(table_id)

    if not table:
        raise HTTPException(status_code=404, detail="数据表不存在")

    columns_config = table.get('columns_config', [])

    # 读取原始请求体
    body = await request.body()

    # 解析 multipart form data
    import multipart
    from fastapi import UploadFile, File

    # 这里直接接收 form data 中的 image 字段
    # 由于 FastAPI 处理 multipart 比较复杂，我们改用简化的方式

    return ApiResponse(message="请使用表单上传图片")


@router.post("/{table_id}/recognize/upload", summary="上传图片并识别")
async def upload_and_recognize(
    table_id: UUID,
    file: UploadFile = File(...),
    session: AsyncSession = Depends(get_db),
):
    """上传扫描件图片并通过AI识别"""
    from app.modules.quality.inspection_table_service import InspectionTableService
    from app.platform.ai.minimax_util import get_vision_util
    import shutil
    from pathlib import Path

    # 获取数据表配置
    service = InspectionTableService(session)
    table = await service.get_table(table_id)

    if not table:
        raise HTTPException(status_code=404, detail="数据表不存在")

    columns_config = table.get('columns_config', [])

    if not columns_config:
        raise HTTPException(status_code=400, detail="该数据表没有配置列，请先编辑表头配置")

    # 保存上传的文件
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    uploads_dir = backend_dir / "uploads" / "inspection"
    uploads_dir.mkdir(parents=True, exist_ok=True)

    # 生成唯一文件名
    import uuid
    file_ext = Path(file.filename).suffix if file.filename else '.png'
    saved_filename = f"{uuid.uuid4().hex}{file_ext}"
    file_path = uploads_dir / saved_filename

    # 保存文件
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    logger.info(f"Uploaded file saved to: {file_path}")

    # 构建 AI 识别提示词
    # 生成列名列表用于提示 AI
    column_labels = [col.get('label', col.get('key', '')) for col in columns_config]
    column_keys = [col.get('key', '') for col in columns_config]

    prompt = f"""你是一个专业的表单数据提取AI。请仔细识别图片中的表单数据。

表单包含以下列：
{chr(10).join([f"- {label} ({key})" for label, key in zip(column_labels, column_keys)])}

请严格按照以下JSON格式返回识别结果（只返回JSON，不要其他内容）：
{{
    "rows": [
        {{{", ".join([f'"{key}": "识别的值"' for key in column_keys])}}}
    ]
}}

注意事项：
1. 每个字段的值应该是识别到的实际文本内容
2. 如果某列在图片中没有对应数据，值留空字符串 ""
3. 如果是日期字段，请格式化为 YYYY-MM-DD 格式
4. 如果是数字字段，只保留数字部分
5. 返回的JSON必须可以被Python/JS解析"""

    try:
        # 调用 AI 识别
        vision_util = get_vision_util()
        result = await vision_util.recognize_image(
            image_urls=[f"/uploads/inspection/{saved_filename}"],
            prompt=prompt,
            max_tokens=4096,
        )

        logger.info(f"AI recognition result: {result}")

        # 解析 AI 返回的 JSON
        try:
            # 尝试提取 JSON 部分
            import re
            json_match = re.search(r'\{[\s\S]*\}', result)
            if json_match:
                recognized_data = json.loads(json_match.group())
            else:
                recognized_data = json.loads(result)
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI response as JSON: {e}")
            logger.error(f"Raw response: {result}")
            raise HTTPException(status_code=500, detail=f"AI返回格式解析失败: {str(e)}")

        return ApiResponse(data={
            "image_url": f"/uploads/inspection/{saved_filename}",
            "recognized_rows": recognized_data.get('rows', []),
            "columns_config": columns_config,
        })

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/{table_id}/recognize/multiple", summary="多图片上传并识别")
async def upload_and_recognize_multiple(
    table_id: UUID,
    files: list[UploadFile] = File(...),
    session: AsyncSession = Depends(get_db),
):
    """上传多个扫描件图片并通过AI识别，汇总所有识别结果"""
    from app.modules.quality.inspection_table_service import InspectionTableService
    from app.platform.ai.minimax_util import get_vision_util
    import shutil
    import uuid
    from pathlib import Path

    # 获取数据表配置
    service = InspectionTableService(session)
    table = await service.get_table(table_id)

    if not table:
        raise HTTPException(status_code=404, detail="数据表不存在")

    columns_config = table.get('columns_config', [])

    if not columns_config:
        raise HTTPException(status_code=400, detail="该数据表没有配置列，请先编辑表头配置")

    if not files:
        raise HTTPException(status_code=400, detail="请上传至少一张图片")

    # 保存上传的文件
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    uploads_dir = backend_dir / "uploads" / "inspection"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    logger.info(f"[DEBUG] Uploads directory: {uploads_dir}")
    logger.info(f"[DEBUG] Files count: {len(files)}")

    # 保存所有文件
    saved_files = []
    for i, file in enumerate(files):
        file_ext = Path(file.filename).suffix if file.filename else '.png'
        saved_filename = f"{uuid.uuid4().hex}{file_ext}"
        file_path = uploads_dir / saved_filename

        logger.info(f"[DEBUG] Processing file {i}: {file.filename} -> {file_path}")

        with open(file_path, "wb") as f:
            shutil.copyfileobj(file.file, f)

        logger.info(f"[DEBUG] File saved: {file_path}, exists: {file_path.exists()}")

        saved_files.append({
            "original_name": file.filename,
            "saved_path": f"/uploads/inspection/{saved_filename}",
        })
        logger.info(f"Uploaded file saved to: {file_path}")

    # 构建 AI 识别提示词
    column_labels = [col.get('label', col.get('key', '')) for col in columns_config]
    column_keys = [col.get('key', '') for col in columns_config]

    # 为每张图片构建图片URL列表
    image_urls = [f["saved_path"] for f in saved_files]

    prompt = f"""你是一个专业的表单数据提取AI。请仔细识别以下多张图片中的表单数据，并将所有识别到的数据汇总。

表单包含以下列：
{chr(10).join([f"- {label} ({key})" for label, key in zip(column_labels, column_keys)])}

请严格按照以下JSON格式返回识别结果（只返回JSON，不要其他内容）：
{{
    "rows": [
        {{{", ".join([f'"{key}": "识别的值"' for key in column_keys])}}}
    ]
}}

注意事项：
1. 图片数量：{len(files)} 张
2. 请识别所有图片中的数据，并将结果统一汇总到 rows 数组中
3. 每行数据代表一条检验记录
4. 每个字段的值应该是识别到的实际文本内容
5. 如果某列在图片中没有对应数据，值留空字符串 ""
6. 如果是日期字段，请格式化为 YYYY-MM-DD 格式
7. 如果是数字字段，只保留数字部分
8. 返回的JSON必须可以被Python/JS解析"""

    try:
        # 调用 AI 识别（一次性传入所有图片）
        logger.info(f"[DEBUG] Calling AI with {len(image_urls)} images")
        vision_util = get_vision_util()
        result = await vision_util.recognize_image(
            image_urls=image_urls,
            prompt=prompt,
            max_tokens=8192,
        )

        logger.info(f"[DEBUG] AI recognition result length: {len(result) if result else 0}")

        # 解析 AI 返回的 JSON
        try:
            import re
            json_match = re.search(r'\{[\s\S]*\}', result)
            if json_match:
                recognized_data = json.loads(json_match.group())
            else:
                recognized_data = json.loads(result)
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI response as JSON: {e}")
            logger.error(f"Raw response: {result}")
            raise HTTPException(status_code=500, detail=f"AI返回格式解析失败: {str(e)}")

        return ApiResponse(data={
            "images": saved_files,
            "recognized_rows": recognized_data.get('rows', []),
            "columns_config": columns_config,
        })

    except ValueError as e:
        logger.error(f"ValueError in multiple recognition: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Unexpected error in multiple recognition: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"服务器内部错误: {str(e)}")


# ============ Word 模板管理 API ============

@router.post("/{table_id}/template", summary="上传Word模板")
async def upload_template(
    table_id: UUID,
    file: UploadFile = File(...),
    session: AsyncSession = Depends(get_db),
):
    """上传Word模板文件"""
    from app.modules.quality.inspection_table_service import InspectionTableService
    import shutil
    import uuid
    from pathlib import Path

    # 验证文件类型
    if not file.filename or not file.filename.lower().endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="仅支持 .docx 或 .doc 格式的Word文件")

    # 获取数据表
    service = InspectionTableService(session)
    table = await service.get_table(table_id)
    if not table:
        raise HTTPException(status_code=404, detail="数据表不存在")

    # 保存模板文件
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    templates_dir = backend_dir / "uploads" / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)

    # 使用 UUID 生成唯一文件名
    file_ext = Path(file.filename).suffix.lower()
    saved_filename = f"{uuid.uuid4().hex}{file_ext}"
    file_path = templates_dir / saved_filename

    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # 更新数据库
    await service.update_table(
        table_id,
        {
            "template_path": f"/uploads/templates/{saved_filename}",
            "template_name": file.filename,
        }
    )

    return ApiResponse(
        message="模板上传成功",
        data={
            "template_path": f"/uploads/templates/{saved_filename}",
            "template_name": file.filename,
        }
    )


@router.delete("/{table_id}/template", summary="删除Word模板")
async def delete_template(
    table_id: UUID,
    session: AsyncSession = Depends(get_db),
):
    """删除Word模板文件"""
    from app.modules.quality.inspection_table_service import InspectionTableService
    from pathlib import Path

    service = InspectionTableService(session)
    table = await service.get_table(table_id)
    if not table:
        raise HTTPException(status_code=404, detail="数据表不存在")

    # 删除文件
    if table.get("template_path"):
        backend_dir = Path(__file__).resolve().parent.parent.parent.parent
        file_path = backend_dir / table["template_path"].lstrip("/")
        if file_path.exists():
            file_path.unlink()

    # 更新数据库
    await service.update_table(
        table_id,
        {"template_path": None, "template_name": None}
    )

    return ApiResponse(message="模板删除成功")


# ============ Word 导出 API ============

@router.get("/{table_id}/rows/{row_id}/export", summary="导出单条数据为Word")
async def export_row_to_word(
    table_id: UUID,
    row_id: int,
    session: AsyncSession = Depends(get_db),
):
    """导出单条数据为Word文档"""
    from app.modules.quality.inspection_table_service import InspectionTableService
    from fastapi.responses import FileResponse
    from docx import Document
    import copy

    service = InspectionTableService(session)
    table = await service.get_table(table_id)
    if not table:
        raise HTTPException(status_code=404, detail="数据表不存在")

    # 获取行数据
    row = await service.get_row(row_id)
    if not row or str(row.get("table_id")) != str(table_id):
        raise HTTPException(status_code=404, detail="数据行不存在")

    # 检查是否有模板
    if not table.get("template_path"):
        raise HTTPException(status_code=400, detail="该数据表未设置Word模板，请先上传模板")

    # 读取模板
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    template_path = backend_dir / table["template_path"].lstrip("/")

    if not template_path.exists():
        raise HTTPException(status_code=500, detail="模板文件不存在，请重新上传")

    # 复制模板
    import tempfile
    import os
    temp_dir = tempfile.gettempdir()
    output_filename = f"export_{uuid.uuid4().hex}.docx"
    output_path = Path(temp_dir) / output_filename

    import shutil
    shutil.copy(template_path, output_path)

    # 填充数据
    doc = Document(output_path)
    row_data = row.get("row_data", {})
    columns_config = table.get("columns_config", [])

    for paragraph in doc.paragraphs:
        for col in columns_config:
            key = col.get("key", "")
            label = col.get("label", key)
            value = str(row_data.get(key, ""))
            # 替换 {{列名}} 或 {{key}} 格式的占位符
            paragraph.text = paragraph.text.replace(f"{{{label}}}", value)
            paragraph.text = paragraph.text.replace(f"{{{key}}}", value)

    # 处理表格中的占位符
    for table_obj in doc.tables:
        for row_obj in table_obj.rows:
            for cell in row_obj.cells:
                for col in columns_config:
                    key = col.get("key", "")
                    label = col.get("label", key)
                    value = str(row_data.get(key, ""))
                    cell.text = cell.text.replace(f"{{{label}}}", value)
                    cell.text = cell.text.replace(f"{{{key}}}", value)

    doc.save(output_path)

    # 生成下载文件名
    download_name = f"{table.get('table_name', '导出')}_{row_id}.docx"

    return FileResponse(
        path=str(output_path),
        filename=download_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.get("/{table_id}/export", summary="批量导出数据为Word")
async def export_rows_to_word(
    table_id: UUID,
    row_ids: str = "",  # 逗号分隔的ID列表，空表示全部
    session: AsyncSession = Depends(get_db),
):
    """批量导出数据为Word文档（每行数据生成一个文档）"""
    from app.modules.quality.inspection_table_service import InspectionTableService
    from fastapi.responses import FileResponse
    from docx import Document
    from pathlib import Path
    import tempfile
    import shutil
    import zipfile
    import os

    service = InspectionTableService(session)
    table = await service.get_table(table_id)
    if not table:
        raise HTTPException(status_code=404, detail="数据表不存在")

    # 检查是否有模板
    if not table.get("template_path"):
        raise HTTPException(status_code=400, detail="该数据表未设置Word模板，请先上传模板")

    # 获取要导出的行
    if row_ids:
        id_list = [int(x.strip()) for x in row_ids.split(",") if x.strip()]
        rows = []
        for rid in id_list:
            row = await service.get_row(rid)
            if row and str(row.get("table_id")) == str(table_id):
                rows.append(row)
    else:
        rows = await service.get_rows_by_table(table_id)

    if not rows:
        raise HTTPException(status_code=400, detail="没有可导出的数据")

    # 读取模板
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    template_path = backend_dir / table["template_path"].lstrip("/")

    if not template_path.exists():
        raise HTTPException(status_code=500, detail="模板文件不存在，请重新上传")

    # 创建临时目录存放导出文件
    temp_dir = Path(tempfile.mkdtemp())
    columns_config = table.get("columns_config", [])

    for idx, row in enumerate(rows):
        # 复制模板
        output_path = temp_dir / f"row_{row['id']}.docx"
        shutil.copy(template_path, output_path)

        # 填充数据
        doc = Document(output_path)
        row_data = row.get("row_data", {})

        for paragraph in doc.paragraphs:
            for col in columns_config:
                key = col.get("key", "")
                label = col.get("label", key)
                value = str(row_data.get(key, ""))
                paragraph.text = paragraph.text.replace(f"{{{label}}}", value)
                paragraph.text = paragraph.text.replace(f"{{{key}}}", value)

        for table_obj in doc.tables:
            for row_obj in table_obj.rows:
                for cell in row_obj.cells:
                    for col in columns_config:
                        key = col.get("key", "")
                        label = col.get("label", key)
                        value = str(row_data.get(key, ""))
                        cell.text = cell.text.replace(f"{{{label}}}", value)
                        cell.text = cell.text.replace(f"{{{key}}}", value)

        doc.save(output_path)

    # 创建 ZIP 文件
    zip_path = Path(tempfile.gettempdir()) / f"export_{uuid.uuid4().hex}.zip"
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file in temp_dir.glob("*.docx"):
            zipf.write(file, file.name)

    # 清理临时目录
    shutil.rmtree(temp_dir)

    download_name = f"{table.get('table_name', '导出')}_批量导出.zip"

    return FileResponse(
        path=str(zip_path),
        filename=download_name,
        media_type="application/zip"
    )
