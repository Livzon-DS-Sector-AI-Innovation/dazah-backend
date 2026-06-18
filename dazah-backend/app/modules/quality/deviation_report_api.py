"""偏差报告 API"""

import logging
import uuid
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from pydantic import BaseModel

from app.core.database import get_db, AsyncSession

logger = logging.getLogger(__name__)


class ApiResponse(BaseModel):
    """统一响应格式"""
    code: int = 200
    message: str = "Success"
    data: Optional[dict | list] = None


router = APIRouter(prefix="/deviation-report", tags=["偏差报告"])


# ============ Word 模板上传 API ============

@router.post("/upload", summary="上传并解析Word模板")
async def upload_word_template(
    file: UploadFile = File(...),
):
    """上传Word文档并解析为HTML格式"""
    import mammoth
    import tempfile
    import shutil

    # 验证文件类型
    if not file.filename or not file.filename.lower().endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="仅支持 .docx 或 .doc 格式的Word文件")

    # 保存上传的文件
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    uploads_dir = backend_dir / "uploads" / "deviation_reports"
    uploads_dir.mkdir(parents=True, exist_ok=True)

    # 使用 UUID 生成唯一文件名
    file_ext = Path(file.filename).suffix.lower()
    saved_filename = f"{uuid.uuid4().hex}{file_ext}"
    file_path = uploads_dir / saved_filename

    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        # 转换为 HTML
        with open(file_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value

        # 处理警告
        messages = [msg.message for msg in result.messages]

        return ApiResponse(
            message="文档解析成功",
            data={
                "file_id": saved_filename,
                "file_name": file.filename,
                "html_content": html_content,
                "file_path": f"/uploads/deviation_reports/{saved_filename}",
                "warnings": messages,
            }
        )
    except Exception as e:
        logger.error(f"文档解析失败: {e}")
        raise HTTPException(status_code=500, detail=f"文档解析失败: {str(e)}")


@router.post("/save", summary="保存HTML内容为Word")
async def save_html_to_word(
    file_id: str,
    html_content: str,
):
    """将HTML内容保存为Word文档"""
    from fastapi.responses import FileResponse
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import tempfile
    import re
    from bs4 import BeautifulSoup

    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    uploads_dir = backend_dir / "uploads" / "deviation_reports"
    original_path = uploads_dir / file_id

    if not original_path.exists():
        raise HTTPException(status_code=404, detail="原始文件不存在，请重新上传")

    # 创建新的 Word 文档
    doc = Document(original_path)

    # 解析 HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # 清除所有现有段落（保留样式）
    for para in doc.paragraphs:
        para.clear()

    # 处理段落
    for element in soup.children:
        if element.name == 'p':
            # 添加段落
            para = doc.add_paragraph()
            for child in element.children:
                if hasattr(child, 'name') and child.name:
                    if child.name == 'b' or child.name == 'strong':
                        run = para.add_run(child.get_text())
                        run.bold = True
                    elif child.name == 'i':
                        run = para.add_run(child.get_text())
                        run.italic = True
                    else:
                        para.add_run(child.get_text())
                else:
                    para.add_run(str(child))

        elif element.name == 'table':
            # 处理表格 - 跳过（保留原表格）
            pass
        elif hasattr(element, 'get_text') and element.get_text().strip():
            # 其他文本
            para = doc.add_paragraph()
            para.add_run(element.get_text())

    # 保存
    temp_dir = tempfile.gettempdir()
    output_filename = f"saved_report_{uuid.uuid4().hex}.docx"
    output_path = Path(temp_dir) / output_filename
    doc.save(str(output_path))

    return FileResponse(
        path=str(output_path),
        filename=f"偏差报告_{uuid.uuid4().hex[:8]}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.get("/download", summary="下载原始Word文档")
async def download_word(
    file_id: str,
):
    """下载原始Word文档"""
    from fastapi.responses import FileResponse

    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    uploads_dir = backend_dir / "uploads" / "deviation_reports"
    file_path = uploads_dir / file_id

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    return FileResponse(
        path=str(file_path),
        filename=file_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ============ AI 优化接口 ============

class OptimizeTextRequest(BaseModel):
    """AI优化文本请求"""
    text: str
    optimize_type: str = "polish"  # polish=润色, summarize=总结, expand=扩展, simplify=简化


@router.post("/ai/optimize", summary="AI优化文本")
async def ai_optimize_text(
    request: OptimizeTextRequest,
):
    """使用AI优化选中的文本"""
    from app.platform.ai.minimax_util import get_vision_util

    # 根据优化类型构建提示词
    prompts = {
        "polish": f"""请对以下文本进行专业润色，保持原意，使语言更加流畅、专业、规范。

原文：
{request.text}

要求：
1. 保持原文的核心意思不变
2. 修正语法错误和表达不当之处
3. 使语言更加流畅、专业
4. 保持原有的格式和结构
5. 直接返回润色后的文本，不要添加任何解释说明

润色结果：""",

        "summarize": f"""请对以下文本进行简要总结。

原文：
{request.text}

要求：
1. 提取核心要点
2. 语言简洁明了
3. 直接返回总结内容，不要添加任何解释说明

总结结果：""",

        "expand": f"""请对以下文本进行适当扩展和详细说明。

原文：
{request.text}

要求：
1. 在保持原意的基础上进行扩展
2. 增加必要的细节和说明
3. 使内容更加完整和详尽
4. 直接返回扩展后的文本，不要添加任何解释说明

扩展结果：""",

        "simplify": f"""请对以下文本进行简化，使其更加简洁易懂。

原文：
{request.text}

要求：
1. 去除冗余表达
2. 使用简洁清晰的语言
3. 保持核心意思不变
4. 直接返回简化后的文本，不要添加任何解释说明

简化结果：""",
    }

    prompt = prompts.get(request.optimize_type, prompts["polish"])

    try:
        vision_util = get_vision_util()
        result = await vision_util.recognize_text(
            prompt=prompt,
            max_tokens=4096,
        )

        # 清理返回结果
        result = result.strip()
        if result.startswith("润色结果："):
            result = result[5:]
        elif result.startswith("总结结果："):
            result = result[5:]
        elif result.startswith("扩展结果："):
            result = result[5:]
        elif result.startswith("简化结果："):
            result = result[5:]

        return ApiResponse(
            message="优化成功",
            data={"optimized_text": result}
        )
    except Exception as e:
        logger.error(f"AI优化失败: {e}")
        raise HTTPException(status_code=500, detail=f"AI优化失败: {str(e)}")
