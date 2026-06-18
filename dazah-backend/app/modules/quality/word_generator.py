"""Word文档生成器 - 支持表格扫描解析"""

import re
from io import BytesIO
from pathlib import Path
from typing import Optional, TypedDict
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


class TableColumnConfig(TypedDict):
    """表格列配置"""
    key: str
    label: str
    type: str


class TemplateFields(TypedDict):
    """模板字段配置"""
    static: dict
    table_fields: dict


class WordGenerator:
    """Word文档生成器"""

    # 占位符匹配模式：{{field_key}}
    PLACEHOLDER_PATTERN = re.compile(r"\{\{([^}]+)\}\}")
    # 表格占位符前缀
    TABLE_PLACEHOLDER_PREFIX = "table:"

    def __init__(self, template_path: str):
        self.template_path = template_path
        self.document: Optional[Document] = None

    def load_template(self) -> Document:
        """加载模板文件"""
        self.document = Document(self.template_path)
        return self.document

    def scan_table_structure(self) -> dict:
        """扫描Word文档中的表格结构，自动提取字段配置"""
        if not self.document:
            self.load_template()

        result: TemplateFields = {
            "static": {},
            "table_fields": {"columns": [], "rows": 10}
        }

        # 扫描所有表格
        for table_idx, table in enumerate(self.document.tables):
            # 检查是否是数据表格（通常有表头行）
            if len(table.rows) < 2:
                continue

            # 获取第一行作为表头
            header_row = table.rows[0]
            columns: list[TableColumnConfig] = []

            for col_idx, cell in enumerate(header_row.cells):
                header_text = cell.text.strip()

                # 跳过空表头或明显的占位符
                if not header_text or header_text.startswith("{{"):
                    continue

                # 生成字段key（将中文转为拼音首字母或英文）
                key = self._text_to_key(header_text)

                columns.append({
                    "key": key,
                    "label": header_text,
                    "type": "text"
                })

            # 如果找到多个列，认为是数据表格
            if len(columns) >= 2:
                result["table_fields"]["columns"] = columns
                result["table_fields"]["rows"] = max(len(table.rows) - 1, 10)  # 至少10行

        # 扫描段落中的静态字段占位符
        for paragraph in self.document.paragraphs:
            text = paragraph.text
            matches = self.PLACEHOLDER_PATTERN.findall(text)
            for match in matches:
                if not match.startswith(self.TABLE_PLACEHOLDER_PREFIX):
                    result["static"][match] = {
                        "placeholder": f"{{{{{match}}}}}",
                        "label": match,
                    }

        return result

    def _text_to_key(self, text: str) -> str:
        """将文本转换为字段key"""
        # 移除特殊字符，只保留字母、数字、下划线
        import re
        key = re.sub(r'[^\w]', '', text)
        if not key:
            # 如果为空，使用序号
            return "field"
        return key.lower()

    def parse_template(self) -> dict:
        """解析模板，提取字段配置（兼容旧版本）"""
        if not self.document:
            self.load_template()

        fields = {
            "static": {},
            "table_fields": {"columns": [], "rows": 10},
        }

        # 解析段落中的占位符
        for paragraph in self.document.paragraphs:
            text = paragraph.text
            matches = self.PLACEHOLDER_PATTERN.findall(text)
            for match in matches:
                if not match.startswith(self.TABLE_PLACEHOLDER_PREFIX):
                    fields["static"][match] = {
                        "placeholder": f"{{{{{match}}}}}",
                        "label": match,
                    }

        # 扫描表格结构
        table_config = self.scan_table_structure()
        if table_config["table_fields"]["columns"]:
            fields["table_fields"] = table_config["table_fields"]

        return fields

    def generate_report(
        self, static_data: dict, table_data: list[dict]
    ) -> bytes:
        """生成报告单Word文档"""
        if not self.document:
            self.load_template()

        # 创建文档副本
        doc = Document(self.template_path)

        # 替换静态字段
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, static_data)

        # 替换表格内容
        for table in doc.tables:
            self._replace_in_table(table, table_data)

        # 保存到内存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _replace_in_paragraph(self, paragraph, data: dict):
        """替换段落中的占位符"""
        for run in paragraph.runs:
            text = run.text
            matches = self.PLACEHOLDER_PATTERN.findall(text)
            for key in matches:
                if key in data and data[key] is not None:
                    value = str(data[key])
                    text = text.replace(f"{{{{{key}}}}}", value)
            run.text = text

    def _replace_in_table(self, table, rows: list[dict]):
        """替换表格内容"""
        # 查找包含表格占位符的行
        placeholder_row_idx = None
        for idx, row in enumerate(table.rows):
            row_text = "".join(cell.text for cell in row.cells)
            if self.TABLE_PLACEHOLDER_PREFIX in row_text:
                placeholder_row_idx = idx
                break

        if placeholder_row_idx is None:
            return

        # 获取模板行的单元格数量
        template_row = table.rows[placeholder_row_idx]
        num_cols = len(template_row.cells)

        # 移除占位符行
        row_elem = template_row._element
        row_elem.getparent().remove(row_elem)

        # 添加数据行
        for row_data in rows:
            row = table.add_row()
            for col_idx, cell in enumerate(row.cells):
                if col_idx < num_cols:
                    # 获取对应列的key
                    template_cell = template_row.cells[col_idx]
                    cell_text = template_cell.text
                    matches = self.PLACEHOLDER_PATTERN.findall(cell_text)

                    if matches:
                        key = matches[0]
                        value = row_data.get(key, "")
                        cell.text = str(value) if value is not None else ""
                    else:
                        # 如果是普通文本，复制模板内容
                        cell.text = template_cell.text

    def save_to_file(self, output_path: str, static_data: dict, table_data: list[dict]):
        """保存到文件"""
        content = self.generate_report(static_data, table_data)
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "wb") as f:
            f.write(content)


def scan_word_table_structure(template_path: str) -> dict:
    """扫描Word文档表格结构，返回字段配置"""
    generator = WordGenerator(template_path)
    return generator.scan_table_structure()


def get_template_fields(template_path: str) -> dict:
    """解析模板文件，返回字段配置"""
    generator = WordGenerator(template_path)
    return generator.parse_template()


def generate_report_file(
    template_path: str,
    output_path: str,
    static_data: dict,
    table_data: list[dict],
) -> str:
    """生成报告文件并返回保存路径"""
    generator = WordGenerator(template_path)
    generator.save_to_file(output_path, static_data, table_data)
    return output_path


def generate_report_bytes(
    template_path: str,
    static_data: dict,
    table_data: list[dict],
) -> bytes:
    """生成报告文件并返回字节内容"""
    generator = WordGenerator(template_path)
    return generator.generate_report(static_data, table_data)