from __future__ import annotations

from datetime import date
from decimal import ROUND_HALF_UP, Decimal
from io import BytesIO

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.modules.procurement.contract_generator import (
    generate_contract,
    get_contract_template_metadata,
    rmb_upper,
)
from app.modules.procurement.schemas import (
    ContractCategory,
    ContractGenerateRequest,
    ContractItemInput,
    ContractPartyInfo,
)

SELLER = ContractPartyInfo(
    name="宁夏自动化测试供应商有限公司",
    representative="李四法人",
    address="宁夏银川市金凤区测试路88号",
    postal_code="750002",
    contact_person="张三测试",
    contact_address="宁夏银川市金凤区联系人地址99号",
    contact_phone="0951-77779999",
    mobile="13900001111",
    phone="0951-66668888",
    bank_name="中国银行银川测试支行",
    bank_account="6400123456789012345",
    tax_id="91640100TEST123456",
    bank_line_number="104871000001",
    email="contract-test@example.com",
)

ITEMS = [
    ContractItemInput(
        item_code="MAT-A001",
        name="自动化测试物料A",
        specification="25kg/袋",
        quality_standard="企业内控标准A",
        manufacturer="测试制造一厂",
        department="生产一部",
        quantity=Decimal("12.5"),
        unit="袋",
        unit_price=Decimal("80.40"),
        amount=None,
        remarks="首批送货",
    ),
    ContractItemInput(
        item_code="MAT-B002",
        name="自动化测试物料B",
        specification="500ml/瓶",
        quality_standard="企业内控标准B",
        manufacturer="测试制造二厂",
        department="质检部",
        quantity=Decimal("3"),
        unit="瓶",
        unit_price=Decimal("199.99"),
        amount=Decimal("600.00"),
        remarks="需附 COA",
    ),
]

OLD_TEMPLATE_SAMPLES = {
    ContractCategory.raw_materials: ["活性炭", "上海兴长活性炭有限公司", "254000元"],
    ContractCategory.consumables: ["单头微量称量勺", "R2A琼脂培养基", "123360.90"],
    ContractCategory.hardware: [
        "不锈钢法兰,DN20-PN10/304",
        "宁夏诚弘胜机电泵阀有限公司",
        "45210.00",
    ],
    ContractCategory.fixed_assets: [
        "LBC300螺杆泵",
        "浙江方远力鑫真空设备有限公司",
        "236000",
    ],
}

EXPECTED_FIRST_TABLE_ROWS = {
    ContractCategory.raw_materials: len(ITEMS) + 4,
    ContractCategory.consumables: len(ITEMS) + 2,
    ContractCategory.hardware: len(ITEMS) + 2,
    ContractCategory.fixed_assets: len(ITEMS) + 2,
}


def _payload(category: ContractCategory) -> ContractGenerateRequest:
    common = dict(
        contract_date=date(2026, 6, 27),
        delivery_date=date(2026, 7, 15),
        delivery_terms="2026年07月15日前一次性交付至买方仓库",
        payment_terms="货到验收合格并收到全额增值税专用发票后45日内银行承兑支付",
        tax_rate=Decimal("13"),
        seller=SELLER,
        items=ITEMS,
    )
    if category is ContractCategory.fixed_assets:
        return ContractGenerateRequest(
            category=category,
            contract_number="QA-FA-20260627-001",
            attached_documents="合格证、说明书、装箱单、校验证书",
            installation_days=5,
            warranty_months=18,
            response_hours=8,
            onsite_hours=24,
            maintenance_response_hours=36,
            overdue_days=10,
            jurisdiction="宁夏回族自治区银川市金凤区人民法院",
            attachment_note="附件一：设备技术协议；附件二：验收标准",
            copies=4,
            buyer_copies=2,
            arrival_payment_condition="货到安装调试验收合格后30",
            arrival_payment_method="银行承兑汇票",
            arrival_payment_ratio=Decimal("85"),
            warranty_payment_ratio=Decimal("15"),
            warranty_payment_method="电汇",
            **common,
        )
    if category is ContractCategory.consumables:
        return ContractGenerateRequest(
            category=category,
            contract_number=f"QA-{category.value}-20260627-001",
            overdue_days=9,
            buyer_invoice_recipient="耗材发票接收人-赵一",
            buyer_invoice_recipient_mobile="13800000001",
            buyer_receiver="耗材收货人-钱二",
            buyer_receiver_mobile="13800000002",
            buyer_receiver_phone="0952-TEST-0002",
            **common,
        )
    return ContractGenerateRequest(
        category=category,
        contract_number=f"QA-{category.value}-20260627-001",
        **common,
    )


def _is_red_run(run) -> bool:
    rpr = run._element.rPr
    if rpr is None:
        return False
    color = rpr.find(qn("w:color"))
    if color is None:
        return False
    value = (color.get(qn("w:val")) or "").upper()
    if len(value) == 6:
        try:
            red = int(value[0:2], 16)
            green = int(value[2:4], 16)
            blue = int(value[4:6], 16)
            return red >= 180 and green <= 100 and blue <= 100
        except ValueError:
            return False
    return value in {"FF0000", "C00000"}


def _red_groups(doc: Document) -> list[str]:
    groups = []
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for paragraph in paragraphs:
        current = []
        for run in paragraph.runs:
            if _is_red_run(run):
                current.append(run.text)
            elif current:
                groups.append("".join(current))
                current = []
        if current:
            groups.append("".join(current))
    return [group for group in groups if group.strip()]


def _all_text(doc: Document) -> str:
    chunks = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(chunks)


def _money(value: Decimal) -> str:
    return f"{value.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP):.2f}"


def _item_amount(item: ContractItemInput) -> Decimal:
    if item.amount is not None:
        return item.amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return (item.quantity * item.unit_price).quantize(
        Decimal("0.01"),
        rounding=ROUND_HALF_UP,
    )


def _summary_values(payload: ContractGenerateRequest) -> list[str]:
    total = sum((_item_amount(item) for item in payload.items), Decimal("0"))
    tax_base = Decimal("1") + payload.tax_rate / Decimal("100")
    untaxed = (total / tax_base).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    tax = (total - untaxed).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return [
        _money(total),
        _money(untaxed),
        _money(tax),
        rmb_upper(total),
        rmb_upper(untaxed),
        rmb_upper(tax),
    ]


def _expected_values(payload: ContractGenerateRequest) -> list[str]:
    values = [
        payload.contract_number,
        SELLER.name,
        SELLER.address,
        SELLER.phone,
        SELLER.bank_name,
        SELLER.bank_account,
        SELLER.tax_id,
        *_summary_values(payload),
    ]
    if payload.category is not ContractCategory.fixed_assets:
        values.append(payload.payment_terms)

    if payload.category is ContractCategory.raw_materials:
        values += [
            "2026 年06月27日",
            payload.delivery_terms,
            SELLER.postal_code,
            "合同金额（含13%增值税）：",
        ]
        for item in payload.items:
            values += [
                item.item_code,
                item.name,
                item.specification,
                item.manufacturer,
                item.unit,
                _money(item.unit_price),
                _money(_item_amount(item)),
                item.remarks,
            ]
    elif payload.category is ContractCategory.consumables:
        values += [
            "2026",
            "07",
            "15",
            f"{payload.overdue_days}天",
            SELLER.representative,
            SELLER.contact_person,
            SELLER.contact_address,
            SELLER.mobile,
            SELLER.email,
            payload.buyer_invoice_recipient,
            payload.buyer_invoice_recipient_mobile,
            payload.buyer_receiver,
            payload.buyer_receiver_mobile,
            payload.buyer_receiver_phone,
            "税务登记号",
            "代表人姓名（当货品有质量问题等联络用）",
        ]
        for item in payload.items:
            values += [
                item.name,
                item.specification,
                str(item.quantity.normalize()),
                item.unit,
                _money(item.unit_price),
                _money(_item_amount(item)),
                item.department,
            ]
    elif payload.category is ContractCategory.hardware:
        values += [
            "2026年07月15日",
            SELLER.representative,
            SELLER.postal_code,
            SELLER.contact_person,
            SELLER.contact_address,
            SELLER.contact_phone,
            SELLER.mobile,
            SELLER.email,
        ]
        for item in payload.items:
            values += [
                item.item_code,
                item.name,
                item.department,
                str(item.quantity.normalize()),
                _money(item.unit_price),
                _money(_item_amount(item)),
                item.unit,
                item.remarks,
            ]
    elif payload.category is ContractCategory.fixed_assets:
        values += [
            payload.attached_documents,
            "2026 年07月15日",
            payload.delivery_terms,
            str(payload.installation_days),
            str(payload.warranty_months),
            str(payload.response_hours),
            str(payload.onsite_hours),
            str(payload.maintenance_response_hours),
            str(payload.overdue_days),
            payload.jurisdiction,
            payload.attachment_note,
            str(payload.copies),
            str(payload.buyer_copies),
            payload.arrival_payment_condition,
            payload.arrival_payment_method,
            "85%",
            "15%",
            payload.warranty_payment_method,
            SELLER.bank_line_number,
            SELLER.contact_person,
        ]
        for item in payload.items:
            values += [
                item.name,
                item.specification,
                item.manufacturer,
                f"{str(item.quantity.normalize())}{item.unit}",
                _money(item.unit_price),
                _money(_item_amount(item)),
                item.remarks,
            ]
    return [value for value in values if value]


@pytest.mark.parametrize("category", list(ContractCategory))
def test_generate_contract_fills_template_without_leftovers(
    category: ContractCategory,
) -> None:
    payload = _payload(category)

    buffer, filename, media_type = generate_contract(payload)
    doc = Document(BytesIO(buffer.getvalue()))
    text = _all_text(doc)

    assert filename.endswith(".docx")
    assert media_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert _red_groups(doc) == []
    assert len(doc.tables[0].rows) == EXPECTED_FIRST_TABLE_ROWS[category]
    assert [value for value in _expected_values(payload) if value not in text] == []
    assert [value for value in OLD_TEMPLATE_SAMPLES[category] if value in text] == []

    if category is ContractCategory.raw_materials:
        summary_start = len(payload.items) + 1
        for row_index in range(summary_start, summary_start + 3):
            for cell in doc.tables[0].rows[row_index].cells[:3]:
                assert all(
                    paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
                    for paragraph in cell.paragraphs
                )

    if category is ContractCategory.consumables:
        seller_contact_text = doc.tables[1].rows[3].cells[0].text
        buyer_delivery_text = doc.tables[1].rows[3].cells[2].text
        assert "\n电话：" not in seller_contact_text
        assert f"手机：{SELLER.mobile}" in seller_contact_text
        assert f"接收人：{payload.buyer_invoice_recipient}" in buyer_delivery_text
        assert f"手机：{payload.buyer_invoice_recipient_mobile}" in buyer_delivery_text
        assert f"收货人：{payload.buyer_receiver}" in buyer_delivery_text
        assert f"手机：{payload.buyer_receiver_mobile}" in buyer_delivery_text
        assert f"电话：{payload.buyer_receiver_phone}" in buyer_delivery_text

    if category is ContractCategory.fixed_assets:
        assert f"后{payload.maintenance_response_hours}小时内" in text
        assert f"_{payload.maintenance_response_hours}__小时" not in text


def test_fixed_asset_template_metadata_exposes_red_marker_fields() -> None:
    metadata = get_contract_template_metadata(ContractCategory.fixed_assets)
    field_names = {field.name for field in metadata.fields}

    assert {
        "attached_documents",
        "seller.bank_line_number",
        "installation_days",
        "warranty_months",
        "response_hours",
        "onsite_hours",
        "maintenance_response_hours",
        "overdue_days",
        "jurisdiction",
        "attachment_note",
        "copies",
        "buyer_copies",
        "arrival_payment_condition",
        "arrival_payment_method",
        "arrival_payment_ratio",
        "warranty_payment_ratio",
        "warranty_payment_method",
    } <= field_names


def test_consumables_template_metadata_exposes_buyer_delivery_fields() -> None:
    metadata = get_contract_template_metadata(ContractCategory.consumables)
    field_names = {field.name for field in metadata.fields}

    assert {
        "buyer_invoice_recipient",
        "buyer_invoice_recipient_mobile",
        "buyer_receiver",
        "buyer_receiver_mobile",
        "buyer_receiver_phone",
    } <= field_names
