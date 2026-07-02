import json
import uuid
from datetime import UTC, datetime
from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document

from app.modules.agent.schemas import AgentChatRequest, AgentToolExecuteRequest
from app.modules.agent.service import (
    HUMAN_DECISION_REQUIRED_MESSAGE,
    AgentService,
)
from app.modules.procurement.contract_generator import generate_contract
from app.modules.procurement.schemas import ContractGenerateRequest


class FakeDb:
    committed = False

    async def commit(self) -> None:
        self.committed = True


class FakeAgentRepository:
    def __init__(self) -> None:
        self.session = SimpleNamespace(
            id=uuid.uuid4(),
            user_id=None,
            context={},
            title=None,
        )
        self.messages = []
        self.tool_calls = []

    async def create_session(self, db, *, user_id, context, title):
        self.session.user_id = user_id
        self.session.context = context
        self.session.title = title
        return self.session

    async def add_message(
        self,
        db,
        *,
        session_id,
        role,
        content,
        metadata=None,
        user_id=None,
    ):
        message = SimpleNamespace(
            id=uuid.uuid4(),
            session_id=session_id,
            role=role,
            content=content,
            message_metadata=metadata or {},
            created_at=datetime.now(UTC),
            user_id=user_id,
        )
        self.messages.append(message)
        return message

    async def list_messages(self, db, *, session_id, limit=20):
        return self.messages[-limit:]

    async def create_tool_call(
        self,
        db,
        *,
        session_id,
        operation,
        request_payload,
    ):
        call = SimpleNamespace(
            session_id=session_id,
            operation=operation,
            request_payload=request_payload,
            status="started",
            response_payload=None,
            error_message=None,
        )
        self.tool_calls.append(call)
        return call

    async def finish_tool_call(
        self,
        db,
        call,
        *,
        status,
        response_payload=None,
        error_message=None,
    ):
        call.status = status
        call.response_payload = response_payload
        call.error_message = error_message
        return call

    async def create_confirmation(
        self,
        db,
        *,
        session_id,
        user_id,
        operation,
        summary,
        risk_level,
        request_payload,
        expires_at,
    ):
        confirmation = SimpleNamespace(
            id=uuid.uuid4(),
            session_id=session_id,
            user_id=user_id,
            operation=operation,
            summary=summary,
            risk_level=risk_level,
            status="pending",
            request_payload=request_payload,
            expires_at=expires_at,
        )
        return confirmation


class PolicyOnlyAgentService(AgentService):
    async def _call_hermes(self, **kwargs):
        raise AssertionError("policy-blocked messages must not reach Hermes")


class StreamingAgentService(AgentService):
    async def _call_hermes_stream(self, **kwargs):
        yield "delta", {"text": "你好，"}
        yield "ping", {"ts": 1}
        yield "delta", {"text": "我正在查询。"}
        yield "done", {
            "message": "你好，我正在查询。",
            "pending_confirmations": [],
            "tool_trace": [{"tool": "search"}],
        }


class ErrorStreamingAgentService(AgentService):
    async def _call_hermes_stream(self, **kwargs):
        yield "delta", {"text": "处理中"}
        yield "error", {"message": "上游中断"}


def parse_sse_events(frames: list[str]) -> list[tuple[str, dict]]:
    events = []
    for frame in frames:
        event = "message"
        data_lines = []
        for line in frame.strip().splitlines():
            if line.startswith("event:"):
                event = line.removeprefix("event:").strip()
            elif line.startswith("data:"):
                data_lines.append(line.removeprefix("data:").strip())
        if data_lines:
            events.append((event, json.loads("\n".join(data_lines))))
    return events


def test_human_decision_policy_detects_delegated_approval() -> None:
    assert AgentService._is_human_decision_required_message("帮我审批通过这条采购申请")
    assert AgentService._is_human_decision_required_message("直接驳回这个申请")
    assert AgentService._is_human_decision_required_message("请重启仓储飞书 WebSocket")


def test_human_decision_policy_allows_read_only_approval_queries() -> None:
    assert not AgentService._is_human_decision_required_message("查看待审批采购申请")
    assert not AgentService._is_human_decision_required_message(
        "给我查询待审批采购申请明细"
    )


@pytest.mark.anyio
async def test_chat_returns_policy_refusal_for_delegated_approval() -> None:
    repo = FakeAgentRepository()
    service = PolicyOnlyAgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4(), name="测试用户")
    db = FakeDb()

    response = await service.chat(
        db,
        request=AgentChatRequest(message="帮我审批通过这条采购申请"),
        current_user=user,
    )

    assert response.message.content == HUMAN_DECISION_REQUIRED_MESSAGE
    assert response.message.metadata["policy"] == "human_decision_required"
    assert response.pending_confirmations == []
    assert db.committed is True


@pytest.mark.anyio
async def test_stream_chat_emits_start_delta_ping_and_done() -> None:
    repo = FakeAgentRepository()
    service = StreamingAgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4(), name="测试用户")
    db = FakeDb()

    frames = [
        frame
        async for frame in service.stream_chat(
            db,
            request=AgentChatRequest(message="查一下库存"),
            current_user=user,
        )
    ]
    events = parse_sse_events(frames)

    assert [event for event, _ in events] == ["start", "delta", "ping", "delta", "done"]
    assert events[1][1]["text"] == "你好，"
    assert events[3][1]["text"] == "我正在查询。"
    assert events[-1][1]["message"]["content"] == "你好，我正在查询。"
    assert events[-1][1]["tool_trace"] == [{"tool": "search"}]
    assert db.committed is True


@pytest.mark.anyio
async def test_stream_chat_error_does_not_store_assistant_message() -> None:
    repo = FakeAgentRepository()
    service = ErrorStreamingAgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4(), name="测试用户")
    db = FakeDb()

    frames = [
        frame
        async for frame in service.stream_chat(
            db,
            request=AgentChatRequest(message="查一下库存"),
            current_user=user,
        )
    ]
    events = parse_sse_events(frames)

    assert [event for event, _ in events] == ["start", "delta", "error"]
    assert events[-1][1]["message"] == "上游中断"
    assert [message.role for message in repo.messages] == ["user"]
    assert db.committed is False


@pytest.mark.anyio
async def test_stream_chat_policy_refusal_returns_done_without_hermes() -> None:
    repo = FakeAgentRepository()
    service = PolicyOnlyAgentService(settings=SimpleNamespace(), repo=repo)
    user = SimpleNamespace(id=uuid.uuid4(), name="测试用户")
    db = FakeDb()

    frames = [
        frame
        async for frame in service.stream_chat(
            db,
            request=AgentChatRequest(message="帮我审批通过这条采购申请"),
            current_user=user,
        )
    ]
    events = parse_sse_events(frames)

    assert [event for event, _ in events] == ["start", "done"]
    assert events[-1][1]["message"]["content"] == HUMAN_DECISION_REQUIRED_MESSAGE
    assert events[-1][1]["pending_confirmations"] == []
    assert db.committed is True


@pytest.mark.anyio
async def test_high_risk_tool_execution_returns_policy_refusal() -> None:
    repo = FakeAgentRepository()
    service = AgentService(settings=SimpleNamespace(), repo=repo)

    response = await service.execute_tool(
        object(),
        request=AgentToolExecuteRequest(
            operation="procurement.approve_purchase_request",
            params={"request_id": str(uuid.uuid4())},
        ),
    )

    assert response.ok is False
    assert response.requires_confirmation is False
    assert response.confirmation is None
    assert response.meta["policy"] == "human_decision_required"
    assert response.data["message"] == HUMAN_DECISION_REQUIRED_MESSAGE
    assert repo.tool_calls[0].status == "rejected_by_policy"


def test_contract_generation_sample_request_is_normalized() -> None:
    service = AgentService(settings=SimpleNamespace())

    request = service._normalize_tool_request(
        AgentToolExecuteRequest(
            operation="procurement.generate_contract",
            params={"department": "行政部", "month": "2026年7月"},
            reason="生成办公用品耗材采购合同示例（行政部 2026年7月）",
        )
    )

    assert request.params == {}
    assert request.body["category"] == "consumables"
    assert request.body["contract_number"].startswith("AI-CONSUMABLES-")
    assert request.body["items"][0]["name"] == "办公用品耗材"
    assert request.body["items"][0]["department"] == "行政部"


@pytest.mark.anyio
async def test_agent_lists_all_contract_templates_with_fields() -> None:
    repo = FakeAgentRepository()
    service = AgentService(settings=SimpleNamespace(), repo=repo)

    response = await service.execute_tool(
        object(),
        request=AgentToolExecuteRequest(
            operation="procurement.list_contract_templates",
        ),
    )

    templates = response.data["templates"]
    by_category = {template["category"]: template for template in templates}

    assert response.ok is True
    assert set(by_category) == {
        "fixed-assets",
        "consumables",
        "hardware",
        "raw-materials",
    }
    assert by_category["fixed-assets"]["template"]["file"] == "fixed-assets.docx"
    assert by_category["fixed-assets"]["template"]["exists"] is True
    assert by_category["fixed-assets"]["template"]["size"] > 0
    assert "contract_number" in by_category["consumables"]["required_fields"]
    assert "seller.name" in by_category["raw-materials"]["required_fields"]
    assert by_category["hardware"]["item_required_fields"] == [
        "name",
        "quantity",
        "unit_price",
    ]
    assert repo.tool_calls[0].status == "succeeded"


@pytest.mark.anyio
async def test_agent_get_contract_template_normalizes_chinese_category() -> None:
    repo = FakeAgentRepository()
    service = AgentService(settings=SimpleNamespace(), repo=repo)

    response = await service.execute_tool(
        object(),
        request=AgentToolExecuteRequest(
            operation="procurement.get_contract_template",
            params={"category": "办公用品耗材"},
        ),
    )

    field_names = {field["name"] for field in response.data["fields"]}

    assert response.ok is True
    assert response.data["category"] == "consumables"
    assert response.data["template"]["file"] == "consumables.docx"
    assert "buyer_invoice_recipient" in field_names


def test_contract_generation_sample_matches_template_and_exports_docx() -> None:
    service = AgentService(settings=SimpleNamespace())
    request = service._normalize_tool_request(
        AgentToolExecuteRequest(
            operation="procurement.generate_contract",
            params={"department": "行政部", "month": "2026年7月"},
            reason="生成办公用品耗材采购合同示例（行政部 2026年7月）",
        )
    )
    payload = ContractGenerateRequest.model_validate(request.body)

    buffer, filename, media_type = generate_contract(payload)
    document = Document(BytesIO(buffer.getvalue()))
    text = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )

    assert filename.startswith("耗材合同_")
    assert filename.endswith(".docx")
    assert media_type.endswith(
        "vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(buffer.getvalue()) > 1000
    assert "办公用品耗材" in text
    assert "行政部" in text


@pytest.mark.anyio
async def test_contract_generation_without_items_returns_validation_error() -> None:
    repo = FakeAgentRepository()
    service = AgentService(
        settings=SimpleNamespace(AGENT_WRITE_CONFIRM_TTL_SECONDS=300),
        repo=repo,
    )

    response = await service.execute_tool(
        object(),
        request=AgentToolExecuteRequest(
            operation="procurement.generate_contract",
            params={"category": "耗材"},
            reason="生成耗材采购合同",
        ),
    )

    assert response.ok is False
    assert response.requires_confirmation is False
    assert response.meta["validation"] == "failed"
    assert "合同明细" in response.data["message"]
    assert repo.tool_calls[0].status == "invalid_request"
